"""
综合测试脚本：验证 8 个新功能在南京公报上的运行效果
功能：斜体、下划线、段落间距、项目符号、编号列表、段落底纹、段落边框、超链接
"""

import os
import sys
import shutil
from pathlib import Path
from docx import Document

# 必须从项目根目录运行此脚本
from core.agents.agent_a import AgentA
from core.orchestrator.task_spec import TaskSpec, TaskType, FileInfo, FileType


def test_all_8_features():
    """
    测试用自然语言指令，包含所有 8 个新功能：
    1. set_italic（斜体）
    2. set_underline（下划线）
    3. set_paragraph_spacing（段落间距）
    4. set_bullet_list（项目符号）
    5. set_numbered_list（编号列表）
    6. set_paragraph_shading（段落底纹）
    7. set_paragraph_border（段落边框）
    8. add_hyperlink（超链接）
    """
    
    # 测试文件路径
    data_dir = Path('tests/test_a/data')
    test_file = data_dir / "南京市2021~2024年国民经济和社会发展统计公报汇总.docx"
    output_file = data_dir / "南京公报_8功能测试结果.docx"
    
    if not test_file.exists():
        print(f"✗ 测试文件不存在: {test_file}")
        return
    
    # 综合测试指令，包含所有 8 个新功能
    instruction = (
        "请对文档进行以下格式编辑：\n"
        "1. 在第一段的'2021年'应用斜体格式\n"
        "2. 在第一段的'国民经济'应用下划线格式\n"
        "3. 将第二段的段落间距设置为段前6磅、段后12磅\n"
        "4. 将'主要成就'后面的内容改为项目符号列表\n"
        "5. 将'关键指标'后面的内容改为编号列表\n"
        "6. 为'发展目标'所在段落添加黄色底纹\n"
        "7. 为'统计数据'所在段落添加边框\n"
        "8. 为文档末尾添加一个超链接，文本为'更多统计数据'，指向'https://www.nanjing.gov.cn/statistics'"
    )
    
    print("=" * 80)
    print("开始综合测试：8 个新功能")
    print("=" * 80)
    print(f"测试文件: {test_file}")
    print(f"输出文件: {output_file}")
    print(f"\n指令:\n{instruction}\n")
    
    try:
        # 创建任务规范
        task = TaskSpec(
            task_type=TaskType.DOCUMENT_EDITING,
            instruction=instruction,
            source_files=[FileInfo(path=str(test_file), file_type=FileType.DOCX, name=test_file.name)],
            output_file=str(output_file),
        )
        
        # 执行任务
        print("执行指令...\n")
        agent = AgentA()
        result = agent.execute(task)
        
        print("=" * 80)
        print("执行结果:")
        print("=" * 80)
        print(f"✓ 成功: {result.success}")
        print(f"✓ 输出文件: {output_file}")
        
        # 获取生成的文档
        if output_file.exists():
            doc = Document(str(output_file))
            
            # 检查各种格式是否应用
            print("\n格式检查结果:")
            
            all_xml = '\n'.join(p._p.xml for p in doc.paragraphs)
            
            # 检查斜体
            has_italic = '<w:i/>' in all_xml or '<w:i ' in all_xml
            print(f"  斜体 (set_italic): {'✓' if has_italic else '✗'}")
            
            # 检查下划线
            has_underline = '<w:u ' in all_xml or '<w:u/>' in all_xml
            print(f"  下划线 (set_underline): {'✓' if has_underline else '✗'}")
            
            # 检查段落间距（空格元素）
            has_spacing = '<w:spacing ' in all_xml
            print(f"  段落间距 (set_paragraph_spacing): {'✓' if has_spacing else '✗'}")
            
            # 检查列表项目符号
            has_bullet = 'List Bullet' in str([p.style.name for p in doc.paragraphs])
            print(f"  项目符号列表 (set_bullet_list): {'✓' if has_bullet else '✗'}")
            
            # 检查编号列表
            has_numbered = 'List Number' in str([p.style.name for p in doc.paragraphs])
            print(f"  编号列表 (set_numbered_list): {'✓' if has_numbered else '✗'}")
            
            # 检查段落底纹（w:shd）
            has_shading = '<w:shd ' in all_xml
            print(f"  段落底纹 (set_paragraph_shading): {'✓' if has_shading else '✗'}")
            
            # 检查段落边框（w:pBdr）
            has_border = '<w:pBdr>' in all_xml or '<w:pBdr' in all_xml
            print(f"  段落边框 (set_paragraph_border): {'✓' if has_border else '✗'}")
            
            # 检查超链接
            has_hyperlink = '<w:hyperlink' in all_xml
            print(f"  超链接 (add_hyperlink): {'✓' if has_hyperlink else '✗'}")
            
            print(f"\n✓ 所有格式检查完毕，文档已保存到: {output_file}")
        else:
            print(f"✗ 输出文件未生成: {output_file}")
        
        # 打印执行报告
        if result.data:
            actions = result.data.get('actions', [])
            execution_report = result.data.get('execution_report', [])
            
            if actions:
                print(f"\n执行的动作数: {len(actions)}")
                for i, action in enumerate(actions, 1):
                    print(f"  {i}. {action}")
            
            if execution_report:
                print(f"\n执行报告数: {len(execution_report)}")
                for i, report in enumerate(execution_report, 1):
                    print(f"  {i}. {report}")
        
    except Exception as e:
        print(f"✗ 错误: {e}")
        import traceback
        traceback.print_exc()




if __name__ == "__main__":
    test_all_8_features()
