"""Tests for clickable TOC navigation and idempotent TOC insertion."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.shared import Pt

from core.agents.agent_a import DocxAdapter


class _FakeTOCLLM:
    def __init__(self, response: str):
        self._response = response
        self.config = SimpleNamespace(streaming=True)

    def is_available(self) -> bool:
        return True

    def chat(self, messages, temperature=0):
        return self._response


def _build_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("2024年1月")
    doc.add_paragraph("主要指标")
    doc.add_paragraph("这里是解释内容。")
    doc.add_paragraph("2024年2月")
    doc.add_paragraph("主要指标")
    doc.add_paragraph("这里是第二组解释内容。")
    doc.save(str(path))
    return path


def _build_doc_with_captions(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("主要指标")
    doc.add_paragraph("图1 南京市GDP趋势")
    doc.add_paragraph("表2 财政收入分项")
    doc.add_paragraph("南京市统计局、国家统计局南京调查队")
    doc.add_paragraph("政策建议")
    doc.save(str(path))
    return path


def test_insert_toc_creates_clickable_navigation_entries(tmp_path):
    source = _build_doc(tmp_path / "toc_nav.docx")
    llm = _FakeTOCLLM(
        '{"headings": ['
        '{"index": 1, "title": "主要指标", "level": 1}, '
        '{"index": 4, "title": "主要指标", "level": 1}'
        '], "toc_entries": ["主要指标", "主要指标"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})

    assert result.success is True
    assert result.details.get("linked_entries") == 2
    assert result.details.get("skipped") is False

    output = tmp_path / "toc_nav_out.docx"
    adapter.save(str(output))

    out_doc = Document(str(output))
    xml_all = "\n".join(p._p.xml for p in out_doc.paragraphs)
    assert "w:hyperlink" in xml_all
    assert "w:bookmarkStart" in xml_all


def test_insert_toc_should_skip_when_existing_toc_found(tmp_path):
    source = _build_doc(tmp_path / "toc_skip.docx")
    adapter = DocxAdapter(str(source))

    first = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})
    second = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})

    assert first.success is True
    assert second.success is True
    assert second.details.get("skipped") is True
    assert "目录已存在" in second.details.get("reason", "")


def test_insert_toc_should_exclude_figure_table_captions(tmp_path):
    source = _build_doc_with_captions(tmp_path / "toc_caption_filter.docx")
    llm = _FakeTOCLLM(
        '{"headings": ['
        '{"index": 0, "title": "主要指标", "level": 1}, '
        '{"index": 1, "title": "图1 南京市GDP趋势", "level": 1}, '
        '{"index": 2, "title": "表2 财政收入分项", "level": 1}, '
        '{"index": 3, "title": "南京市统计局、国家统计局南京调查队", "level": 1}, '
        '{"index": 4, "title": "政策建议", "level": 1}'
        '], "toc_entries": ["主要指标", "图1 南京市GDP趋势", "表2 财政收入分项", "南京市统计局、国家统计局南京调查队", "政策建议"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})
    assert result.success is True
    assert result.details.get("toc_entries") == 2

    output = tmp_path / "toc_caption_filter_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))
    texts = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    toc_lines = [t for t in texts if t.startswith("-")]
    assert "- 主要指标" in toc_lines
    assert "- 政策建议" in toc_lines
    assert not any("图1" in t for t in toc_lines)
    assert not any("表2" in t for t in toc_lines)
    assert not any("统计局" in t or "调查队" in t for t in toc_lines)


def test_insert_toc_hyperlink_text_should_not_be_explicit_blue(tmp_path):
    source = _build_doc(tmp_path / "toc_color.docx")
    llm = _FakeTOCLLM(
        '{"headings": ['
        '{"index": 1, "title": "主要指标", "level": 1}, '
        '{"index": 4, "title": "主要指标", "level": 1}'
        '], "toc_entries": ["主要指标", "主要指标"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})
    assert result.success is True

    output = tmp_path / "toc_color_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))

    toc_lines = [p for p in out_doc.paragraphs if (p.text or "").strip().startswith("-")]
    assert toc_lines
    for p in toc_lines:
        # 目录项不应带显式蓝色，避免与“标题深蓝色”混淆。
        for run in p.runs:
            if run.font.color.rgb is not None:
                assert str(run.font.color.rgb).upper() != "0000FF"


def test_toc_hyperlinks_should_survive_body_formatting_actions(tmp_path):
    source = _build_doc(tmp_path / "toc_keep_link.docx")
    llm = _FakeTOCLLM(
        '{"headings": ['
        '{"index": 1, "title": "主要指标", "level": 1}, '
        '{"index": 4, "title": "主要指标", "level": 1}'
        '], "toc_entries": ["主要指标", "主要指标"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    assert adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}}).success
    assert adapter.apply_action(
        {"action_type": "set_font_family", "target": {"scope": "body"}, "params": {"font_name": "宋体"}}
    ).success
    assert adapter.apply_action(
        {"action_type": "set_font_size", "target": {"scope": "body"}, "params": {"size_pt": 12}}
    ).success

    output = tmp_path / "toc_keep_link_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))

    xml_all = "\n".join(p._p.xml for p in out_doc.paragraphs)
    assert "w:hyperlink" in xml_all


def test_document_main_title_should_be_included_in_toc_and_colored(tmp_path):
    doc = Document()
    doc.add_paragraph("南京全市2024年国民经济和社会发展统计公报")
    doc.add_paragraph("2024年3月")
    doc.add_paragraph("一、综合")
    doc.add_paragraph("这是正文内容。")
    source = tmp_path / "toc_with_doc_title.docx"
    doc.save(str(source))

    llm = _FakeTOCLLM(
        '{"document_title": {"index": 0, "title": "南京全市2024年国民经济和社会发展统计公报", "confidence": 0.99}, '
        '"headings": [{"index": 1, "title": "2024年3月", "level": 1}, {"index": 2, "title": "一、综合", "level": 1}], '
        '"toc_entries": ["2024年3月", "一、综合"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    toc_result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})
    assert toc_result.success is True
    assert toc_result.details.get("document_title_included") is True

    adapter.apply_action(
        {
            "action_type": "set_font_color",
            "target": {"scope": "heading"},
            "params": {"color": "0000FF"},
        }
    )

    output = tmp_path / "toc_with_doc_title_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))

    texts = [(p.text or "").strip() for p in out_doc.paragraphs if (p.text or "").strip()]
    toc_lines = [t for t in texts if t.startswith("-")]
    assert toc_lines and toc_lines[0] == "- 南京全市2024年国民经济和社会发展统计公报"

    title_para = next(p for p in out_doc.paragraphs if (p.text or "").strip() == "南京全市2024年国民经济和社会发展统计公报")
    assert any(run.font.color.rgb is not None and str(run.font.color.rgb).upper() == "0000FF" for run in title_para.runs)


def test_heading_scope_should_always_include_document_title_with_toc_present(tmp_path):
    doc = Document()
    doc.add_paragraph("南京全市2024年国民经济和社会发展统计公报")
    doc.add_paragraph("2024年3月")
    doc.add_paragraph("一、综合")
    doc.add_paragraph("正文段落")
    source = tmp_path / "title_scope.docx"
    doc.save(str(source))

    llm = _FakeTOCLLM(
        '{"document_title": {"index": 0, "title": "南京全市2024年国民经济和社会发展统计公报", "confidence": 0.99}, '
        '"headings": [{"index": 1, "title": "2024年3月", "level": 1}, {"index": 2, "title": "一、综合", "level": 1}], '
        '"toc_entries": ["南京全市2024年国民经济和社会发展统计公报", "2024年3月", "一、综合"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    assert adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}}).success
    assert adapter.apply_action(
        {"action_type": "set_font_color", "target": {"scope": "heading"}, "params": {"color": "0000FF"}}
    ).success
    assert adapter.apply_action(
        {"action_type": "bold_heading", "target": {"level": 1}, "params": {"bold": True}}
    ).success

    output = tmp_path / "title_scope_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))
    title_para = next(p for p in out_doc.paragraphs if (p.text or "").strip() == "南京全市2024年国民经济和社会发展统计公报")
    assert any(run.font.color.rgb is not None and str(run.font.color.rgb).upper() == "0000FF" for run in title_para.runs)
    assert any(run.bold for run in title_para.runs)


def test_insert_toc_should_link_entries_with_normalized_text_match(tmp_path):
    doc = Document()
    doc.add_paragraph("2024年3月")
    doc.add_paragraph("注释：")
    doc.add_paragraph("资料来源：")
    source = tmp_path / "toc_normalized_match.docx"
    doc.save(str(source))

    llm = _FakeTOCLLM(
        '{"headings": ['
        '{"index": 0, "title": "2024年3月", "level": 1}, '
        '{"index": 1, "title": "注释：", "level": 1}, '
        '{"index": 2, "title": "资料来源：", "level": 1}'
        '], "toc_entries": ["2024年3月", "注释", "资料来源"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})

    assert result.success is True
    assert result.details.get("toc_entries") == 3
    assert result.details.get("linked_entries") == 3


def test_insert_toc_should_track_document_title_after_replace(tmp_path):
    doc = Document()
    doc.add_paragraph("南京市2024年国民经济和社会发展统计公报")
    doc.add_paragraph("2024年3月")
    doc.add_paragraph("一、综合")
    source = tmp_path / "toc_title_replace.docx"
    doc.save(str(source))

    llm = _FakeTOCLLM(
        '{"document_title": {"index": 0, "title": "南京市2024年国民经济和社会发展统计公报", "confidence": 0.99}, '
        '"headings": [{"index": 1, "title": "2024年3月", "level": 1}, {"index": 2, "title": "一、综合", "level": 1}], '
        '"toc_entries": ["2024年3月", "一、综合"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    assert adapter.apply_action(
        {"action_type": "replace_text", "target": {}, "params": {"find": "南京市", "replace": "南京全市"}}
    ).success

    toc_result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})
    assert toc_result.success is True
    assert toc_result.details.get("document_title_included") is True
    assert toc_result.details.get("linked_entries") == toc_result.details.get("toc_entries")

    color_result = adapter.apply_action(
        {
            "action_type": "set_font_color",
            "target": {"scope": "heading"},
            "params": {"color": "0000FF"},
        }
    )
    assert color_result.success is True

    title_para = next(
        p for p in adapter.document.paragraphs if (p.text or "").strip() == "南京全市2024年国民经济和社会发展统计公报"
    )
    assert any(run.font.color.rgb is not None and str(run.font.color.rgb).upper() == "0000FF" for run in title_para.runs)


def test_document_title_should_remain_bold_after_replace_text(tmp_path):
    doc = Document()
    doc.add_paragraph("南京市2024年国民经济和社会发展统计公报")
    doc.add_paragraph("2024年3月")
    doc.add_paragraph("一、综合")
    source = tmp_path / "title_bold_keep_after_replace.docx"
    doc.save(str(source))

    llm = _FakeTOCLLM(
        '{"document_title": {"index": 0, "title": "南京市2024年国民经济和社会发展统计公报", "confidence": 0.99}, '
        '"headings": [{"index": 1, "title": "2024年3月", "level": 1}, {"index": 2, "title": "一、综合", "level": 1}], '
        '"toc_entries": ["2024年3月", "一、综合"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    assert adapter.apply_action(
        {"action_type": "bold_heading", "target": {"level": 1}, "params": {"bold": True}}
    ).success
    assert adapter.apply_action(
        {"action_type": "replace_text", "target": {}, "params": {"find": "南京市", "replace": "南京全市"}}
    ).success

    title_para = next(
        p for p in adapter.document.paragraphs if (p.text or "").strip() == "南京全市2024年国民经济和社会发展统计公报"
    )
    assert any(run.bold for run in title_para.runs)


def test_set_font_size_body_should_not_change_document_title_size(tmp_path):
    doc = Document()
    title_para = doc.add_paragraph("南京全市2024年国民经济和社会发展统计公报")
    title_para.runs[0].font.size = Pt(20)
    doc.add_paragraph("2024年3月")
    doc.add_paragraph("一、综合")
    doc.add_paragraph("这是正文。")
    source = tmp_path / "title_size_should_keep.docx"
    doc.save(str(source))

    llm = _FakeTOCLLM(
        '{"document_title": {"index": 0, "title": "南京全市2024年国民经济和社会发展统计公报", "confidence": 0.99}, '
        '"headings": [{"index": 1, "title": "2024年3月", "level": 1}, {"index": 2, "title": "一、综合", "level": 1}], '
        '"toc_entries": ["南京全市2024年国民经济和社会发展统计公报", "2024年3月", "一、综合"], "field_values": {}}'
    )

    adapter = DocxAdapter(str(source), llm_service=llm)
    assert adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}}).success
    assert adapter.apply_action(
        {"action_type": "set_font_size", "target": {"scope": "body"}, "params": {"size_pt": 12}}
    ).success

    title_para = next(
        p for p in adapter.document.paragraphs if (p.text or "").strip() == "南京全市2024年国民经济和社会发展统计公报"
    )
    title_sizes = [run.font.size.pt for run in title_para.runs if run.font.size is not None]
    assert title_sizes and all(abs(s - 20.0) < 1e-6 for s in title_sizes)
