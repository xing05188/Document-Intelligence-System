from __future__ import annotations

from docx import Document

from core.agents.agent_a.docx_adapter import DocxAdapter
from core.agents.agent_a.instruction_parser import _extract_formatting_target


class _FakeLLM:
    def is_available(self) -> bool:
        return True

    def chat_with_system(self, system_prompt: str, user_input: str) -> str:
        # 故意返回错误泛化，验证解析器会用显式线索覆盖。
        return '{"paragraph_index": -1, "target_text": "ALL", "scope": "document"}'


def test_extract_formatting_target_prefers_explicit_paragraph_and_text_constraints():
    target = _extract_formatting_target("在第一段的'2021年'应用斜体格式", llm_service=_FakeLLM())

    assert target["scope"] == "selective"
    assert target["paragraph_index"] == 0
    assert target["target_text"] == "2021年"


def test_set_italic_only_styles_target_text_in_specific_paragraph(tmp_path):
    source = tmp_path / "precision.docx"
    doc = Document()
    doc.add_heading("南京市2021~2024年国民经济和社会发展统计公报", level=1)
    doc.add_paragraph("南京市统计局、国家统计局南京调查队")
    doc.add_paragraph("2021年南京市国民经济和社会发展统计公报")
    doc.add_paragraph("图1 2021年主要指标")
    doc.save(str(source))

    adapter = DocxAdapter(str(source))
    result = adapter.apply_action(
        {
            "action_type": "set_italic",
            "target": {"scope": "selective", "target_text": "2021年", "paragraph_index": 2},
            "params": {"italic": True},
        }
    )

    assert result.success

    target_para = adapter.document.paragraphs[2]
    assert any((run.text or "") == "2021年" and bool(run.italic) for run in target_para.runs)
    assert any("国民经济" in (run.text or "") and not bool(run.italic) for run in target_para.runs)

    caption_para = adapter.document.paragraphs[3]
    assert all(not bool(run.italic) for run in caption_para.runs if (run.text or "").strip())
