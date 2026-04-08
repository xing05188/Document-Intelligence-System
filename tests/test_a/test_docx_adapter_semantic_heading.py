"""DOCX adapter semantic heading tests with LLM-enhanced understanding."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from core.agents.agent_a import DocxAdapter


class _FakeLLM:
    def __init__(self, response: str):
        self._response = response
        self.config = SimpleNamespace(streaming=True)

    def is_available(self) -> bool:
        return True

    def chat(self, messages, temperature=0):
        return self._response


class _UnavailableLLM:
    def is_available(self) -> bool:
        return False


def _build_plain_docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("这里是正文内容。")
    doc.add_paragraph("第二章 发展目标")
    doc.add_paragraph("这里是第二段正文。")
    doc.save(str(path))
    return path


def _build_non_numbered_heading_docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("总体情况")
    doc.add_paragraph("这里是第一段正文内容。")
    doc.add_paragraph("主要问题")
    doc.add_paragraph("这里是第二段正文内容。")
    doc.save(str(path))
    return path


def test_bold_heading_uses_semantic_llm_when_no_heading_style(tmp_path):
    source = _build_plain_docx(tmp_path / "plain.docx")
    fake_llm = _FakeLLM('{"indices": [0, 2]}')
    adapter = DocxAdapter(str(source), llm_service=fake_llm)

    result = adapter.apply_action(
        {"action_type": "bold_heading", "target": {"level": 1}, "params": {"bold": True}}
    )

    assert result.success is True
    assert result.details.get("style_hits") == 0
    assert result.details.get("semantic_hits") == 2

    p0 = adapter.document.paragraphs[0]
    p2 = adapter.document.paragraphs[2]
    assert any(run.bold for run in p0.runs)
    assert any(run.bold for run in p2.runs)


def test_insert_toc_uses_semantic_llm_when_no_heading_style(tmp_path):
    source = _build_plain_docx(tmp_path / "plain_toc.docx")
    fake_llm = _FakeLLM('{"indices": [0, 2]}')
    adapter = DocxAdapter(str(source), llm_service=fake_llm)

    result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})

    assert result.success is True
    assert result.details.get("toc_entries") == 2
    assert result.details.get("semantic_heading_used") is True

    output = tmp_path / "plain_toc_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))
    texts = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert "目录" in texts
    assert "- 第一章 总则" in texts
    assert "- 第二章 发展目标" in texts


def test_extract_content_uses_semantic_llm_when_no_heading_style(tmp_path):
    source = _build_plain_docx(tmp_path / "plain_extract.docx")
    fake_llm = _FakeLLM('{"indices": [0, 2]}')
    adapter = DocxAdapter(str(source), llm_service=fake_llm)

    result = adapter.apply_action({"action_type": "extract_content", "target": {}, "params": {"fields": []}})

    assert result.success is True
    assert result.details.get("semantic_heading_used") is True
    headings = result.details.get("headings", [])
    assert "第一章 总则" in headings
    assert "第二章 发展目标" in headings


def test_non_numbered_heading_can_be_detected_without_llm(tmp_path):
    source = _build_non_numbered_heading_docx(tmp_path / "plain_non_numbered.docx")
    adapter = DocxAdapter(str(source), llm_service=_UnavailableLLM())

    result = adapter.apply_action(
        {"action_type": "bold_heading", "target": {"level": 1}, "params": {"bold": True}}
    )

    assert result.success is True
    assert result.details.get("semantic_hits", 0) >= 2


def test_bold_heading_should_follow_toc_strategy_style_first(tmp_path):
    doc = Document()
    doc.add_heading("一、已套样式标题", level=1)
    doc.add_paragraph("二、未套样式标题")
    doc.add_paragraph("这是正文内容。")
    source = tmp_path / "mixed_style_heading.docx"
    doc.save(str(source))

    fake_llm = _FakeLLM(
        '{"headings": [{"index": 0, "title": "一、已套样式标题", "level": 1}, {"index": 1, "title": "二、未套样式标题", "level": 1}], "toc_entries": [], "field_values": {}}'
    )
    adapter = DocxAdapter(str(source), llm_service=fake_llm)

    result = adapter.apply_action(
        {"action_type": "bold_heading", "target": {"level": 1}, "params": {"bold": True}}
    )

    assert result.success is True
    assert result.details.get("style_hits") == 1
    assert result.details.get("semantic_hits") == 0
    assert any(run.bold for run in adapter.document.paragraphs[0].runs)
    assert not any(run.bold for run in adapter.document.paragraphs[1].runs)


def test_heading_scope_style_should_not_color_non_heading_paragraph(tmp_path):
    doc = Document()
    doc.add_paragraph("主要指标")
    doc.add_paragraph("地区生产总值：18500亿元")
    source = tmp_path / "heading_scope_precision.docx"
    doc.save(str(source))

    fake_llm = _FakeLLM(
        '{"headings": [{"index": 0, "title": "主要指标", "level": 1}, {"index": 1, "title": "地区生产总值：18500亿元", "level": 1}], "toc_entries": [], "field_values": {}}'
    )
    adapter = DocxAdapter(str(source), llm_service=fake_llm)

    result = adapter.apply_action(
        {
            "action_type": "set_font_color",
            "target": {"scope": "heading"},
            "params": {"color": "003A6C"},
        }
    )

    assert result.success is True
    heading_runs = adapter.document.paragraphs[0].runs
    body_runs = adapter.document.paragraphs[1].runs
    assert heading_runs and heading_runs[0].font.color.rgb is not None
    assert all(run.font.color.rgb is None for run in body_runs)


def test_semantic_heading_cache_should_be_level_aware(tmp_path):
    doc = Document()
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("（一）工作目标")
    source = tmp_path / "level_cache.docx"
    doc.save(str(source))

    fake_llm = _FakeLLM(
        '{"headings": [{"index": 0, "title": "第一章 总则", "level": 1}, {"index": 1, "title": "（一）工作目标", "level": 2}], "toc_entries": [], "field_values": {}}'
    )
    adapter = DocxAdapter(str(source), llm_service=fake_llm)

    l1 = adapter._get_semantic_heading_indices(level=1)
    l2 = adapter._get_semantic_heading_indices(level=2)
    assert l1 == [0]
    assert l2 == [1]


def test_bold_heading_level1_should_include_document_main_title(tmp_path):
    doc = Document()
    doc.add_paragraph("南京全市2024年国民经济和社会发展统计公报")
    doc.add_paragraph("2024年3月")
    doc.add_paragraph("一、综合")
    source = tmp_path / "main_title_bold.docx"
    doc.save(str(source))

    fake_llm = _FakeLLM(
        '{"document_title": {"index": 0, "title": "南京全市2024年国民经济和社会发展统计公报", "confidence": 0.99}, '
        '"headings": [{"index": 1, "title": "2024年3月", "level": 1}, {"index": 2, "title": "一、综合", "level": 1}], '
        '"toc_entries": ["2024年3月", "一、综合"], "field_values": {}}'
    )
    adapter = DocxAdapter(str(source), llm_service=fake_llm)

    result = adapter.apply_action(
        {"action_type": "bold_heading", "target": {"level": 1}, "params": {"bold": True}}
    )

    assert result.success is True
    title_para = adapter.document.paragraphs[0]
    assert any(run.bold for run in title_para.runs)
