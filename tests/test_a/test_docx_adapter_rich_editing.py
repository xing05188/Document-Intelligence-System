"""Tests for richer DOCX editing actions: font/color/size/alignment/line spacing."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from core.agents.agent_a import DocxAdapter


def _build_doc(path: Path) -> Path:
    doc = Document()
    doc.add_heading("测试标题", level=1)
    doc.add_paragraph("这是正文第一段")
    doc.add_paragraph("这是正文第二段")
    doc.save(str(path))
    return path


def test_docx_adapter_rich_editing_actions(tmp_path):
    source = _build_doc(tmp_path / "rich_edit.docx")
    adapter = DocxAdapter(str(source))

    actions = [
        {"action_type": "set_font_family", "target": {"scope": "body"}, "params": {"font_name": "宋体"}},
        {"action_type": "set_font_color", "target": {"scope": "heading"}, "params": {"color": "FF0000"}},
        {"action_type": "set_font_size", "target": {"scope": "body"}, "params": {"size_pt": 12}},
        {"action_type": "set_paragraph_alignment", "target": {"scope": "all"}, "params": {"alignment": "center"}},
        {"action_type": "set_line_spacing", "target": {"scope": "body"}, "params": {"line_spacing": 1.5}},
    ]

    results = [adapter.apply_action(a) for a in actions]
    assert all(r.success for r in results)

    output = tmp_path / "rich_edit_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))

    # heading color
    heading = out_doc.paragraphs[0]
    assert heading.runs and heading.runs[0].font.color.rgb is not None

    # body font size / spacing / alignment
    body = [p for p in out_doc.paragraphs[1:] if p.text.strip()]
    assert body
    assert all(p.alignment is not None for p in out_doc.paragraphs if p.text.strip())
    assert all(p.paragraph_format.line_spacing == 1.5 for p in body)
    for p in body:
        for run in p.runs:
            if run.text.strip():
                assert run.font.size is not None


def test_docx_adapter_section_color_and_body_first_line_indent(tmp_path):
    doc = Document()
    doc.add_paragraph("一、综合")
    doc.add_paragraph("综合部分第一段")
    doc.add_paragraph("综合部分第二段")
    doc.add_paragraph("二、农业")
    doc.add_paragraph("农业部分正文")
    source = tmp_path / "section_scope.docx"
    doc.save(str(source))

    adapter = DocxAdapter(str(source))
    assert adapter.apply_action(
        {
            "action_type": "set_font_color",
            "target": {"scope": "section_content", "section_title": "综合"},
            "params": {"color": "FF0000"},
        }
    ).success
    assert adapter.apply_action(
        {
            "action_type": "set_first_line_indent",
            "target": {"scope": "body"},
            "params": {"indent_pt": 24},
        }
    ).success

    out_doc = adapter.document
    # 综合章节正文变红
    p_sec_1 = next(p for p in out_doc.paragraphs if (p.text or "").strip() == "综合部分第一段")
    p_sec_2 = next(p for p in out_doc.paragraphs if (p.text or "").strip() == "综合部分第二段")
    for p in (p_sec_1, p_sec_2):
        assert p.runs and any(run.font.color.rgb is not None and str(run.font.color.rgb).upper() == "FF0000" for run in p.runs)

    # 其他章节正文不应被误着色
    p_other = next(p for p in out_doc.paragraphs if (p.text or "").strip() == "农业部分正文")
    assert all(run.font.color.rgb is None for run in p_other.runs)

    # 正文首行缩进应生效（标题不属于 body）
    p_sec = next(p for p in out_doc.paragraphs if (p.text or "").strip() == "综合部分第一段")
    p_heading = next(p for p in out_doc.paragraphs if (p.text or "").strip() == "一、综合")

    assert p_sec.paragraph_format.first_line_indent is not None
    assert abs(p_sec.paragraph_format.first_line_indent.pt - 24.0) < 1e-6
    assert p_other.paragraph_format.first_line_indent is not None
    assert abs(p_other.paragraph_format.first_line_indent.pt - 24.0) < 1e-6
    assert p_heading.paragraph_format.first_line_indent is None
