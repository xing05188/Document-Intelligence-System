"""Tests for whole-document LLM understanding in DocxAdapter."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document

from core.agents.agent_a import DocxAdapter


class _FakeUnderstandingLLM:
    def __init__(self, response: str):
        self._response = response
        self.config = SimpleNamespace(streaming=True)

    def is_available(self) -> bool:
        return True

    def chat(self, messages, temperature=0):
        return self._response


def _build_plain_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("南京全市地区生产总值：18500亿元")
    doc.add_paragraph("第二章 主要指标")
    doc.add_paragraph("一般公共预算收入：1600亿元")
    doc.save(str(path))
    return path


def test_extract_content_prefers_llm_document_understanding(tmp_path):
    source = _build_plain_doc(tmp_path / "understanding.docx")
    llm = _FakeUnderstandingLLM(
        '{"headings": [{"index": 0, "title": "第一章 总则", "level": 1}, {"index": 2, "title": "第二章 主要指标", "level": 1}], '
        '"field_values": {"地区生产总值": "18500亿元", "一般公共预算收入": "1600亿元"}}'
    )
    adapter = DocxAdapter(str(source), llm_service=llm)

    result = adapter.apply_action(
        {
            "action_type": "extract_content",
            "target": {},
            "params": {"fields": ["地区生产总值", "一般公共预算收入"]},
        }
    )

    assert result.success is True
    assert result.details.get("llm_understanding_used") is True
    headings = result.details.get("headings", [])
    assert "第一章 总则" in headings
    assert "第二章 主要指标" in headings

    field_values = result.details.get("field_values", {})
    assert field_values.get("地区生产总值") == "18500亿元"
    assert field_values.get("一般公共预算收入") == "1600亿元"


def test_insert_toc_prefers_llm_document_understanding(tmp_path):
    source = _build_plain_doc(tmp_path / "understanding_toc.docx")
    llm = _FakeUnderstandingLLM(
        '{"headings": [{"index": 0, "title": "第一章 总则", "level": 1}, {"index": 2, "title": "第二章 主要指标", "level": 1}], "field_values": {}}'
    )
    adapter = DocxAdapter(str(source), llm_service=llm)

    result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})
    assert result.success is True
    assert result.details.get("llm_understanding_used") is True
    assert result.details.get("toc_entries") == 2


def test_insert_toc_with_repeated_titles_should_not_add_month_context_prefix(tmp_path):
    source = _build_plain_doc(tmp_path / "understanding_toc_context.docx")
    llm = _FakeUnderstandingLLM(
        '{"headings": ['
        '{"index": 0, "title": "主要指标", "level": 1}, '
        '{"index": 2, "title": "主要指标", "level": 1}'
        '], "toc_entries": ["主要指标", "主要指标"], "field_values": {}}'
    )
    adapter = DocxAdapter(str(source), llm_service=llm)

    result = adapter.apply_action({"action_type": "insert_toc", "target": {}, "params": {}})
    assert result.success is True
    assert result.details.get("contextualized") is False

    output = tmp_path / "understanding_toc_context_out.docx"
    adapter.save(str(output))
    out_doc = Document(str(output))
    texts = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert "- 主要指标" in texts
    assert not any("[" in t and "月" in t for t in texts)
