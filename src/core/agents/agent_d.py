"""core.agents.agent_d

Agent_D: Excel 数据筛选与填表 Agent（由原 Agent_B 改名而来）。

职责：
- 从 Excel（.xlsx）读取表格数据
- 根据用户自然语言/参数生成筛选计划并筛选行
- 输出筛选结果为 JSON
- 将结果按字段映射填入 Excel/Word 模板（支持单表/多表）
"""
from __future__ import annotations

import importlib.util
import json
import re
import os
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from openpyxl import load_workbook
from docx import Document

from .base_agent import BaseAgent, AgentResponse
from config import SystemConfig, get_config
from core.orchestrator.task_spec import TaskSpec, TaskType, FileInfo, FileType
from core.llm.llm_service import get_llm_service
from utils.logger import get_logger


class AgentB(BaseAgent):
    """
    Agent_D: Excel 数据筛选与填表（历史兼容类名：AgentB）

    能力：
    - 理解自然语言筛选条件
    - 从 Excel 中筛选出命中的数据行并输出 JSON
    - （可选）根据模板列进行字段映射并填充到 Word/Excel 模板
    - 支持模板格式: .xlsx / .docx（含多表填表）
    """

    def __init__(self, config: Optional[SystemConfig] = None):
        super().__init__(config)
        # 文件已迁移为 Agent_D，但为了尽量不破坏旧代码，保留类名 AgentB。
        # 对外展示名称使用 Agent_D。
        self.name = "Agent_D"
        self.agent_type = "extraction"
        self.logger = get_logger(__name__)
        try:
            self.llm = get_llm_service()
        except Exception:
            self.llm = None

        self._operator_map = {
            "eq": self._op_eq,
            "ne": self._op_ne,
            "gt": self._op_gt,
            "gte": self._op_gte,
            "lt": self._op_lt,
            "lte": self._op_lte,
            "contains": self._op_contains,
            "not_contains": self._op_not_contains,
            "in": self._op_in,
            "not_in": self._op_not_in,
            "between": self._op_between,
            "regex": self._op_regex,
            "is_null": self._op_is_null,
            "not_null": self._op_not_null,
        }

    def execute(self, task_spec: TaskSpec, **kwargs) -> AgentResponse:
        """
        执行 Excel 数据筛选与填表任务
        """
        # 验证输入
        is_valid, error_msg = self.validate_input(task_spec)
        if not is_valid:
            return AgentResponse(success=False, message=error_msg)

        try:
            return self._extract_entities(task_spec)
        except Exception as e:
            return AgentResponse(
                success=False,
                message=f"提取失败: {str(e)}"
            )

    def _extract_entities(self, task_spec: TaskSpec) -> AgentResponse:
        """
        根据自然语言条件筛选 Excel 行，并输出 JSON；如提供模板则进行填表。
        """
        excel_path = self._find_excel_source(task_spec)
        if not excel_path:
            return AgentResponse(success=False, message="未找到可用的Excel源文件(.xlsx/.xls)")

        table_targets = self._get_table_targets(task_spec.parameters)
        instruction = self._resolve_instruction(task_spec)
        if not instruction and not table_targets:
            return AgentResponse(success=False, message="缺少筛选条件描述，请在 instruction 或 parameters 中提供")

        rows, columns = self._read_excel_rows(excel_path, task_spec.parameters)
        if not rows:
            return AgentResponse(success=False, message="Excel中未读取到可用数据行")

        source_path = Path(excel_path).resolve()
        source_dir = source_path.parent
        default_json_output = str(source_dir / f"{source_path.stem}_filtered_rows.json")

        # 多表模式：每个目标独立筛选并写入对应表区
        if table_targets and task_spec.template_file and self._is_supported_template(task_spec.template_file.path):
            template_path = task_spec.template_file.path
            allow_rule_fallback = bool(task_spec.parameters.get("allow_rule_fallback", True))

            target_results = []
            union_rows: List[Dict[str, Any]] = []

            for idx, target in enumerate(table_targets, start=1):
                target_instruction = str(target.get("instruction") or target.get("condition") or "").strip()
                if not target_instruction:
                    return AgentResponse(success=False, message=f"table_targets[{idx}] 缺少 instruction/condition")

                field_candidates = self._extract_field_candidates(target_instruction, columns)
                plan = self._build_filter_plan_with_llm(target_instruction, columns, field_candidates)
                llm_has_conditions = bool(plan.get("conditions"))
                llm_has_groups = bool(plan.get("groups"))
                plan_source = "llm"

                if (not llm_has_conditions and not llm_has_groups) and allow_rule_fallback:
                    plan = self._build_filter_plan_fallback(target_instruction, columns)
                    plan_source = "rule_fallback"

                if not plan.get("conditions") and not plan.get("groups"):
                    return AgentResponse(success=False, message=f"table_targets[{idx}] 条件无法解析为可执行筛选计划")

                filtered_rows = self._apply_filter_plan(rows, plan)
                union_rows.extend(filtered_rows)

                target_params = self._merge_parameters_with_target(task_spec.parameters, target)
                template_columns = self._read_template_columns(template_path, target_params)
                mapping = self._build_template_column_mapping(
                    source_columns=columns,
                    template_columns=template_columns,
                    source_rows=rows,
                    parameters=target_params,
                )

                target_results.append({
                    "name": str(target.get("name") or f"表{idx}"),
                    "instruction": target_instruction,
                    "plan_source": plan_source,
                    "plan": plan,
                    "matched_rows": len(filtered_rows),
                    "filtered_rows": filtered_rows,
                    "mapping": mapping,
                    "sheet_name": str(target_params.get("template_sheet_name", "")).strip(),
                    "table_index": int(target_params.get("template_table_index", 0)),
                    "header_row": int(target_params.get("template_header_row", 1)),
                    "start_row": int(target_params.get("template_start_row", int(target_params.get("template_header_row", 1)) + 1)),
                })

            if self._is_excel_template(template_path):
                task_spec.parameters.setdefault(
                    "template_output_file",
                    str(source_dir / f"{source_path.stem}_filled{Path(template_path).suffix.lower() or '.xlsx'}"),
                )
                template_output_path = self._fill_excel_template_multi(
                    template_path=template_path,
                    target_results=target_results,
                    parameters=task_spec.parameters,
                )
            else:
                task_spec.parameters.setdefault(
                    "template_output_file",
                    str(source_dir / f"{source_path.stem}_filled{Path(template_path).suffix.lower() or '.docx'}"),
                )
                template_output_path = self._fill_docx_template_multi(
                    template_path=template_path,
                    target_results=target_results,
                    parameters=task_spec.parameters,
                )

            output_path = self._write_rows_to_json(union_rows, task_spec.output_file or default_json_output)

            return AgentResponse(
                success=True,
                message=f"多表填表完成，共处理 {len(target_results)} 个表目标",
                data={
                    "status": "completed",
                    "excel_path": excel_path,
                    "output_json": output_path,
                    "total_rows": len(rows),
                    "matched_rows": len(union_rows),
                    "template_filled": True,
                    "template_output": template_output_path,
                    "multi_table_results": [
                        {
                            "name": t["name"],
                            "instruction": t["instruction"],
                            "plan_source": t["plan_source"],
                            "matched_rows": t["matched_rows"],
                            "mapping": t["mapping"],
                            "sheet_name": t["sheet_name"],
                            "header_row": t["header_row"],
                            "start_row": t["start_row"],
                        }
                        for t in target_results
                    ],
                }
            )

        # 单表模式（原逻辑）
        field_candidates = self._extract_field_candidates(instruction, columns)
        plan = self._build_filter_plan_with_llm(instruction, columns, field_candidates)
        llm_has_conditions = bool(plan.get("conditions"))
        llm_has_groups = bool(plan.get("groups"))
        plan_source = "llm"

        allow_rule_fallback = bool(task_spec.parameters.get("allow_rule_fallback", True))
        if (not llm_has_conditions and not llm_has_groups) and allow_rule_fallback:
            plan = self._build_filter_plan_fallback(instruction, columns)
            plan_source = "rule_fallback"

        has_conditions = bool(plan.get("conditions"))
        has_groups = bool(plan.get("groups"))
        if not has_conditions and not has_groups:
            return AgentResponse(
                success=False,
                message="LLM未生成可执行筛选计划，请检查模型配置或补充更明确条件"
            )

        filtered_rows = self._apply_filter_plan(rows, plan)
        output_path = self._write_rows_to_json(filtered_rows, task_spec.output_file or default_json_output)

        template_output_path = None
        column_mapping: Dict[str, str] = {}
        if task_spec.template_file and self._is_supported_template(task_spec.template_file.path):
            template_path = task_spec.template_file.path
            template_columns = self._read_template_columns(template_path, task_spec.parameters)
            column_mapping = self._build_template_column_mapping(
                source_columns=columns,
                template_columns=template_columns,
                source_rows=rows,
                parameters=task_spec.parameters,
            )
            if self._is_excel_template(template_path):
                task_spec.parameters.setdefault(
                    "template_output_file",
                    str(source_dir / f"{source_path.stem}_filled{Path(template_path).suffix.lower() or '.xlsx'}"),
                )
                template_output_path = self._fill_excel_template(
                    filtered_rows=filtered_rows,
                    template_path=template_path,
                    mapping=column_mapping,
                    parameters=task_spec.parameters,
                )
            else:
                task_spec.parameters.setdefault(
                    "template_output_file",
                    str(source_dir / f"{source_path.stem}_filled{Path(template_path).suffix.lower() or '.docx'}"),
                )
                template_output_path = self._fill_docx_template(
                    filtered_rows=filtered_rows,
                    template_path=template_path,
                    mapping=column_mapping,
                    parameters=task_spec.parameters,
                )

        return AgentResponse(
            success=True,
            message=f"筛选完成，共命中 {len(filtered_rows)} 行",
            data={
                "status": "completed",
                "excel_path": excel_path,
                "output_json": output_path,
                "total_rows": len(rows),
                "matched_rows": len(filtered_rows),
                "used_plan": plan,
                "field_candidates": field_candidates,
                "plan_source": plan_source,
                "template_filled": bool(template_output_path),
                "template_output": template_output_path,
                "template_mapping": column_mapping,
            }
        )

    def _get_table_targets(self, parameters: Dict[str, Any]) -> List[Dict[str, Any]]:
        targets = parameters.get("table_targets", []) if parameters else []
        if isinstance(targets, list):
            return [t for t in targets if isinstance(t, dict)]
        return []

    def _merge_parameters_with_target(self, base_parameters: Dict[str, Any], target: Dict[str, Any]) -> Dict[str, Any]:
        merged = dict(base_parameters or {})

        if "sheet_name" in target and "template_sheet_name" not in target:
            merged["template_sheet_name"] = target.get("sheet_name")
        if "template_sheet_name" in target:
            merged["template_sheet_name"] = target.get("template_sheet_name")
        if "header_row" in target and "template_header_row" not in target:
            merged["template_header_row"] = target.get("header_row")
        if "template_header_row" in target:
            merged["template_header_row"] = target.get("template_header_row")
        if "start_row" in target and "template_start_row" not in target:
            merged["template_start_row"] = target.get("start_row")
        if "template_start_row" in target:
            merged["template_start_row"] = target.get("template_start_row")
        if "table_index" in target and "template_table_index" not in target:
            merged["template_table_index"] = target.get("table_index")
        if "template_table_index" in target:
            merged["template_table_index"] = target.get("template_table_index")

        return merged

    def _is_excel_template(self, template_path: str) -> bool:
        return str(template_path or "").lower().endswith(".xlsx")

    def _is_docx_template(self, template_path: str) -> bool:
        return str(template_path or "").lower().endswith(".docx")

    def _is_supported_template(self, template_path: str) -> bool:
        return self._is_excel_template(template_path) or self._is_docx_template(template_path)

    def _resolve_excel_sheet(self, workbook, sheet_name: str, sheet_index: int = 0):
        if sheet_name and sheet_name in workbook.sheetnames:
            return workbook[sheet_name]

        if 0 <= sheet_index < len(workbook.worksheets):
            return workbook.worksheets[sheet_index]

        return workbook.active

    def validate_input(self, task_spec: TaskSpec) -> tuple[bool, str]:
        """验证输入"""
        if not task_spec.source_files:
            return False, "缺少源文件"

        has_excel = any(f.path.lower().endswith(".xlsx") for f in task_spec.source_files)
        if not has_excel:
            return False, "当前阶段需要提供Excel源文件(.xlsx)"

        return True, ""

    def get_system_prompt(self) -> str:
        """获取系统提示词"""
        return get_config().agent.get_prompt(self.agent_type)

    def _find_excel_source(self, task_spec: TaskSpec) -> Optional[str]:
        for file_info in task_spec.source_files:
            path = (file_info.path or "").strip()
            if path.lower().endswith(".xlsx"):
                return path
        return None

    def _resolve_instruction(self, task_spec: TaskSpec) -> str:
        return (
            task_spec.instruction
            or str(task_spec.parameters.get("filter_condition", ""))
            or str(task_spec.parameters.get("query", ""))
        ).strip()

    def _read_excel_rows(self, excel_path: str, parameters: Dict[str, Any]) -> Tuple[List[Dict[str, Any]], List[str]]:
        sheet_name = str(parameters.get("sheet_name", "")).strip() if parameters else ""
        workbook = load_workbook(excel_path, data_only=True)
        try:
            ws = workbook[sheet_name] if sheet_name and sheet_name in workbook.sheetnames else workbook.active
            data = list(ws.iter_rows(values_only=True))
        finally:
            workbook.close()

        if not data:
            return [], []

        headers_raw = data[0]
        headers = [self._normalize_header(h, idx) for idx, h in enumerate(headers_raw)]

        rows: List[Dict[str, Any]] = []
        for row_values in data[1:]:
            if row_values is None:
                continue
            row_dict: Dict[str, Any] = {}
            for idx, header in enumerate(headers):
                value = row_values[idx] if idx < len(row_values) else None
                row_dict[header] = value

            if any(v is not None and str(v).strip() != "" for v in row_dict.values()):
                rows.append(row_dict)

        return rows, headers

    def _normalize_header(self, header: Any, idx: int) -> str:
        if header is None:
            return f"column_{idx + 1}"
        normalized = str(header).strip()
        return normalized if normalized else f"column_{idx + 1}"

    def _extract_field_candidates(self, instruction: str, columns: List[str]) -> List[str]:
        text = instruction.lower()
        candidates = [col for col in columns if col and str(col).lower() in text]
        return list(dict.fromkeys(candidates))

    def _build_filter_plan_with_llm(
        self,
        instruction: str,
        columns: List[str],
        field_candidates: List[str],
    ) -> Dict[str, Any]:
        if not self.llm or not self.llm.is_available():
            return {}

        system_prompt = (
            "你是Excel筛选条件编译器。你的任务是把用户的自然语言筛选要求，编译成可执行JSON筛选计划。"
            "\n你必须处理口语、省略、同义词、编号条目、多条件混合等表达。"
            "\n"
            "\n硬性规则（必须遵守）："
            "\n1) 只输出一个JSON对象，禁止输出解释、思考过程、Markdown代码块。"
            "\n2) field必须严格取自给定列名，禁止自造字段。"
            "\n3) operator只允许：eq, ne, gt, gte, lt, lte, contains, not_contains, in, not_in, between, regex, is_null, not_null。"
            "\n4) 返回结构固定为：{\"logic\":\"and|or\",\"conditions\":[],\"groups\":[]}。"
            "\n5) conditions用于扁平条件；groups用于分组条件。"
            "\n6) 编号表达（如1. 2. 3.）默认解释为：根logic=or，每个编号项一个group，group.logic=and。"
            "\n7) 时间、日期、字符串尽量保持原文，不要擅自截断、改写或格式化。"
            "\n8) BETWEEN 必须使用 value/value2 两个字段，禁止使用 values 数组。"
            "\n9) 不确定时优先选择语义最接近且可执行的字段；若完全无法判断，返回空计划：{\"logic\":\"and\",\"conditions\":[],\"groups\":[]}。"
            "\n"
            "\n语义映射指引（用于理解，不是输出字段名）："
            "\n- 等于: 是/为/等于/就是/必须是"
            "\n- 不等于: 不是/不为/不等于"
            "\n- 大于: 超过/高于/大于"
            "\n- 大于等于: 不小于/不少于/至少"
            "\n- 小于: 低于/小于/不到"
            "\n- 小于等于: 不大于/至多/不超过"
            "\n- 包含: 含有/包含/带有"
            "\n- 不包含: 不含/不包含/排除"
            "\n- 空值: 为空/空白/缺失"
            "\n- 非空: 不为空/有值"
            "\n- 范围: 在A到B之间/A-B/从A到B"
            "\n"
            "\n输出示例1（编号分组）："
            "\n{"
            "\"logic\":\"or\","
            "\"conditions\":[],"
            "\"groups\":["
            "{\"logic\":\"and\",\"conditions\":[{\"field\":\"监测时间\",\"operator\":\"eq\",\"value\":\"2025-11-25 09:00:00.0\"},{\"field\":\"城市\",\"operator\":\"eq\",\"value\":\"德州市\"}]},"
            "{\"logic\":\"and\",\"conditions\":[{\"field\":\"监测时间\",\"operator\":\"eq\",\"value\":\"2025-11-25 09:00:00.0\"},{\"field\":\"城市\",\"operator\":\"eq\",\"value\":\"潍坊市\"}]}"
            "]"
            "}"
            "\n"
            "\n输出示例2（扁平条件）："
            "\n{"
            "\"logic\":\"and\","
            "\"conditions\":["
            "{\"field\":\"城市\",\"operator\":\"in\",\"values\":[\"德州市\",\"潍坊市\"]},"
            "{\"field\":\"空气质量指数\",\"operator\":\"gt\",\"value\":50}"
            "],"
            "\"groups\":[]"
            "}"
        )
        user_prompt = (
            f"用户条件：{instruction}\n"
            f"可用列名：{json.dumps(columns, ensure_ascii=False)}\n"
            f"已命中候选列：{json.dumps(field_candidates, ensure_ascii=False)}\n"
            "任务要求：\n"
            "- 把用户条件编译为可执行JSON计划。\n"
            "- 若出现编号项，优先生成groups结构（or-of-and）。\n"
            "- 严格输出JSON对象本体，不要任何解释文本。\n"
        )

        try:
            response = self.llm.chat_with_system(system_prompt=system_prompt, user_input=user_prompt)
            plan = self._extract_json_object(response)
            return self._sanitize_filter_plan(plan, columns)
        except Exception as ex:
            self.logger.warning(f"LLM条件解析失败，回退规则解析: {ex}")
            return {}

    def _build_filter_plan_fallback(self, instruction: str, columns: List[str]) -> Dict[str, Any]:
        text = instruction.strip()
        if not text:
            return {}

        grouped_plan = self._build_grouped_plan_from_numbered_blocks(text, columns)
        if grouped_plan.get("groups"):
            return grouped_plan

        pair_conditions = self._extract_conditions_from_field_value_pairs(text, columns)
        if pair_conditions:
            logic = "or" if any(k in text for k in [" 或 ", "或者", "或是"]) else "and"
            return {"logic": logic, "conditions": pair_conditions}

        logic = "or" if any(k in text for k in [" 或 ", "或者", "或是"]) else "and"
        segments = re.split(r"并且|且|以及|同时| and | AND | 或者 | 或 ", text)
        conditions = []

        for seg in segments:
            segment = seg.strip()
            if not segment:
                continue

            field = self._best_match_column(segment, columns)
            if not field:
                continue

            condition = self._parse_segment_to_condition(segment, field)
            if condition:
                conditions.append(condition)

        return {"logic": logic, "conditions": conditions}

    def _read_template_columns(self, template_path: str, parameters: Dict[str, Any]) -> List[str]:
        if self._is_docx_template(template_path):
            return self._read_docx_template_columns(template_path, parameters)

        sheet_name = str(parameters.get("template_sheet_name", "")).strip() if parameters else ""
        header_row = int(parameters.get("template_header_row", 1)) if parameters else 1

        wb = load_workbook(template_path, data_only=True)
        try:
            ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
            row_values = [cell.value for cell in ws[header_row]]
        finally:
            wb.close()

        columns: List[str] = []
        for idx, value in enumerate(row_values):
            if value is None or str(value).strip() == "":
                columns.append(f"template_col_{idx + 1}")
            else:
                columns.append(str(value).strip())
        return columns

    def _read_docx_template_columns(self, template_path: str, parameters: Dict[str, Any]) -> List[str]:
        table_index = int(parameters.get("template_table_index", 0)) if parameters else 0
        header_row = int(parameters.get("template_header_row", 1)) if parameters else 1

        doc = Document(template_path)
        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Word模板中不存在 table_index={table_index} 的表格")

        table = doc.tables[table_index]
        h_idx = max(header_row - 1, 0)
        if h_idx >= len(table.rows):
            raise ValueError(f"Word模板表格表头行不存在: header_row={header_row}")

        row = table.rows[h_idx]
        columns: List[str] = []
        for idx, cell in enumerate(row.cells):
            text = str(cell.text).strip()
            columns.append(text if text else f"template_col_{idx + 1}")
        return columns

    def _build_template_column_mapping(
        self,
        source_columns: List[str],
        template_columns: List[str],
        source_rows: List[Dict[str, Any]],
        parameters: Dict[str, Any],
    ) -> Dict[str, str]:
        # 第一层：列名规则高置信匹配（早停）
        name_strong_threshold = int(parameters.get("mapping_name_strong_threshold", 85)) if parameters else 85
        fallback_threshold = int(parameters.get("mapping_fallback_threshold", 60)) if parameters else 60
        candidate_top_k = int(parameters.get("mapping_candidate_top_k", 5)) if parameters else 5

        mapping: Dict[str, str] = {}
        candidate_scores: Dict[str, List[Dict[str, Any]]] = {}

        for t_col in template_columns:
            ranked = self._rank_source_columns_for_template_column(t_col, source_columns, source_rows, parameters)
            candidate_scores[t_col] = ranked[:candidate_top_k]
            if ranked and int(ranked[0].get("score", 0)) >= name_strong_threshold:
                mapping[t_col] = str(ranked[0]["source"])

        unresolved = [t for t in template_columns if t not in mapping]
        if not unresolved:
            return mapping

        # 第二层：仅对失败列提取小样本值域 + 一次LLM重排
        llm_mapping = self._llm_rerank_unresolved_mapping(
            unresolved_template_columns=unresolved,
            candidate_scores=candidate_scores,
            source_rows=source_rows,
            parameters=parameters,
        )
        for t_col, s_col in llm_mapping.items():
            if t_col not in mapping:
                mapping[t_col] = s_col

        # 第三层：仍未匹配时使用分值回退
        unresolved_after_llm = [t for t in template_columns if t not in mapping]
        for t_col in unresolved_after_llm:
            ranked = candidate_scores.get(t_col, [])
            if ranked and int(ranked[0].get("score", 0)) >= fallback_threshold:
                mapping[t_col] = str(ranked[0]["source"])

        return mapping

    def _rank_source_columns_for_template_column(
        self,
        template_column: str,
        source_columns: List[str],
        source_rows: List[Dict[str, Any]],
        parameters: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        expected_type = self._guess_expected_type_by_column_name(template_column)
        profiles = self._build_source_column_profiles(source_columns, source_rows, parameters)

        ranked: List[Dict[str, Any]] = []
        for s_col in source_columns:
            name_score = self._name_similarity_score(template_column, s_col)
            type_bonus = self._type_compatibility_bonus(expected_type, profiles.get(s_col, {}))
            score = name_score + type_bonus
            ranked.append({
                "source": s_col,
                "score": score,
                "name_score": name_score,
                "type_bonus": type_bonus,
            })

        ranked.sort(key=lambda x: int(x.get("score", 0)), reverse=True)
        return ranked

    def _build_source_column_profiles(
        self,
        source_columns: List[str],
        source_rows: List[Dict[str, Any]],
        parameters: Dict[str, Any],
    ) -> Dict[str, Dict[str, Any]]:
        sample_size = int(parameters.get("mapping_sample_size", 60)) if parameters else 60
        preview_size = int(parameters.get("mapping_preview_size", 8)) if parameters else 8

        profiles: Dict[str, Dict[str, Any]] = {}
        for s_col in source_columns:
            values = []
            for row in source_rows:
                v = row.get(s_col)
                if v is None:
                    continue
                text = str(v).strip()
                if not text:
                    continue
                values.append(v)

            if not values:
                profiles[s_col] = {"detected_type": "unknown", "samples": []}
                continue

            sampled_values = values[:sample_size]
            samples = [str(v) for v in sampled_values[:preview_size]]
            detected_type = self._detect_value_type(sampled_values)

            profiles[s_col] = {
                "detected_type": detected_type,
                "samples": samples,
            }
        return profiles

    def _detect_value_type(self, values: List[Any]) -> str:
        if not values:
            return "unknown"

        numeric_count = 0
        datetime_like_count = 0
        for v in values:
            if isinstance(v, (datetime, date)):
                datetime_like_count += 1
                continue

            text = str(v).strip()
            if not text:
                continue

            if re.search(r"\d{4}[-/]\d{1,2}[-/]\d{1,2}", text):
                datetime_like_count += 1

            if self._extract_numeric_value(text) is not None:
                numeric_count += 1

        total = max(len(values), 1)
        if datetime_like_count / total >= 0.6:
            return "datetime"
        if numeric_count / total >= 0.8:
            return "numeric"

        unique_values = {str(v).strip() for v in values if str(v).strip()}
        if len(unique_values) <= min(20, int(total * 0.5) + 1):
            return "categorical"
        return "text"

    def _extract_numeric_value(self, value: Any) -> Optional[float]:
        if value is None:
            return None
        if isinstance(value, (int, float)) and not isinstance(value, bool):
            return float(value)

        text = str(value).strip()
        if not text:
            return None

        m = re.search(r"[-+]?\d+(?:\.\d+)?", text)
        if not m:
            return None
        try:
            return float(m.group(0))
        except Exception:
            return None

    def _guess_expected_type_by_column_name(self, column_name: str) -> str:
        normalized = self._normalize_column_name(column_name)
        if any(k in normalized for k in ["时间", "日期", "time", "date"]):
            return "datetime"
        if any(k in normalized for k in ["pm", "指数", "值", "浓度", "数量", "比", "率", "温度"]):
            return "numeric"
        if any(k in normalized for k in ["类型", "等级", "城市", "区", "名称", "污染物", "状态"]):
            return "categorical"
        return "text"

    def _type_compatibility_bonus(self, expected_type: str, profile: Dict[str, Any]) -> int:
        detected = str(profile.get("detected_type", "unknown"))
        if expected_type == "text":
            return 0
        if detected == expected_type:
            return 20
        if expected_type == "categorical" and detected in {"text", "categorical"}:
            return 10
        if expected_type == "numeric" and detected == "text":
            return -5
        if expected_type == "datetime" and detected == "text":
            return -8
        return 0

    def _name_similarity_score(self, template_column: str, source_column: str) -> int:
        t = self._normalize_column_name(template_column)
        s = self._normalize_column_name(source_column)
        if not t or not s:
            return 0

        if t == s:
            return 100

        score = 0
        if t in s or s in t:
            score += 65

        inter = len(set(t) & set(s))
        union = len(set(t) | set(s))
        if union > 0:
            score += int((inter / union) * 30)

        # 常见业务同义词加分
        synonym_groups = [
            ("城市", "城市名称", "地市"),
            ("监测时间", "时间", "监测时刻"),
            ("污染类型", "污染等级", "空气质量等级"),
            ("pm25", "pm2.5", "细颗粒物"),
            ("pm10", "颗粒物pm10"),
        ]
        for group in synonym_groups:
            hits_t = any(self._normalize_column_name(x) in t for x in group)
            hits_s = any(self._normalize_column_name(x) in s for x in group)
            if hits_t and hits_s:
                score += 15

        return min(score, 100)

    def _llm_rerank_unresolved_mapping(
        self,
        unresolved_template_columns: List[str],
        candidate_scores: Dict[str, List[Dict[str, Any]]],
        source_rows: List[Dict[str, Any]],
        parameters: Dict[str, Any],
    ) -> Dict[str, str]:
        if not unresolved_template_columns:
            return {}
        if not self.llm or not self.llm.is_available():
            return {}

        profiles = self._build_source_column_profiles(
            source_columns=list({c["source"] for items in candidate_scores.values() for c in items}),
            source_rows=source_rows,
            parameters=parameters,
        )

        llm_input: Dict[str, Any] = {}
        for t_col in unresolved_template_columns:
            candidates = candidate_scores.get(t_col, [])
            llm_input[t_col] = []
            for c in candidates:
                s_col = str(c.get("source"))
                llm_input[t_col].append({
                    "source": s_col,
                    "score": c.get("score", 0),
                    "name_score": c.get("name_score", 0),
                    "type_bonus": c.get("type_bonus", 0),
                    "detected_type": profiles.get(s_col, {}).get("detected_type", "unknown"),
                    "samples": profiles.get(s_col, {}).get("samples", []),
                })

        system_prompt = (
            "你是字段映射重排器。任务：仅对未匹配模板列，从给定候选源列中选择最佳映射。"
            "仅输出JSON对象，不要解释。"
            "输出格式：{\"mapping\": {\"模板列名\": \"候选源列名\"}}。"
            "规则："
            "1) 只能从每个模板列提供的候选列表里选。"
            "2) 重点结合候选样本值语义判断同义字段。"
            "3) 无把握可不输出该模板列。"
            "4) 不要输出候选列表以外的列名。"
        )
        user_prompt = (
            f"未匹配模板列及候选信息：{json.dumps(llm_input, ensure_ascii=False)}\n"
            "请返回mapping。"
        )

        try:
            response = self.llm.chat_with_system(system_prompt=system_prompt, user_input=user_prompt)
            obj = self._extract_json_object(response)
            raw_mapping = obj.get("mapping", {}) if isinstance(obj, dict) else {}
            return self._sanitize_llm_rerank_mapping(raw_mapping, llm_input)
        except Exception as ex:
            self.logger.warning(f"LLM重排映射失败，将使用回退映射: {ex}")
            return {}

    def _sanitize_llm_rerank_mapping(
        self,
        raw_mapping: Dict[str, Any],
        llm_input: Dict[str, List[Dict[str, Any]]],
    ) -> Dict[str, str]:
        if not isinstance(raw_mapping, dict):
            return {}

        result: Dict[str, str] = {}
        for t_col, s_col in raw_mapping.items():
            t_key = str(t_col).strip()
            s_val = str(s_col).strip()
            candidates = {str(x.get("source")) for x in llm_input.get(t_key, [])}
            if s_val in candidates:
                result[t_key] = s_val
        return result

    def _build_template_column_mapping_fallback(
        self,
        source_columns: List[str],
        template_columns: List[str],
    ) -> Dict[str, str]:
        result: Dict[str, str] = {}
        source_norm = {self._normalize_column_name(c): c for c in source_columns}

        for t_col in template_columns:
            t_norm = self._normalize_column_name(t_col)
            if t_norm in source_norm:
                result[t_col] = source_norm[t_norm]
                continue

            best_col = None
            best_score = 0
            for s_col in source_columns:
                s_norm = self._normalize_column_name(s_col)
                score = 0
                if t_norm and t_norm in s_norm:
                    score += 2
                if s_norm and s_norm in t_norm:
                    score += 2

                shared = len(set(t_norm) & set(s_norm))
                score += shared

                if score > best_score:
                    best_score = score
                    best_col = s_col

            if best_col and best_score >= 3:
                result[t_col] = best_col

        return result

    def _sanitize_template_mapping(
        self,
        raw_mapping: Dict[str, Any],
        source_columns: List[str],
        template_columns: List[str],
    ) -> Dict[str, str]:
        if not isinstance(raw_mapping, dict):
            return {}

        result: Dict[str, str] = {}
        source_set = set(source_columns)
        template_set = set(template_columns)

        for t_col, s_col in raw_mapping.items():
            t_key = str(t_col).strip()
            s_val = str(s_col).strip()
            if t_key in template_set and s_val in source_set:
                result[t_key] = s_val
        return result

    def _normalize_column_name(self, name: Any) -> str:
        if name is None:
            return ""
        text = str(name).strip().lower()
        text = re.sub(r"[\s_\-\(\)（）\[\]【】:：,.，。]", "", text)
        return text

    def _fill_excel_template(
        self,
        filtered_rows: List[Dict[str, Any]],
        template_path: str,
        mapping: Dict[str, str],
        parameters: Dict[str, Any],
    ) -> str:
        wb = load_workbook(template_path)
        try:
            sheet_name = str(parameters.get("template_sheet_name", "")).strip() if parameters else ""
            sheet_index = int(parameters.get("template_table_index", 0)) if parameters else 0
            ws = self._resolve_excel_sheet(wb, sheet_name, sheet_index)

            header_row = int(parameters.get("template_header_row", 1)) if parameters else 1
            start_row = int(parameters.get("template_start_row", header_row + 1)) if parameters else (header_row + 1)

            # 模板列名 -> 列索引
            template_col_index: Dict[str, int] = {}
            for cell in ws[header_row]:
                header = str(cell.value).strip() if cell.value is not None else ""
                if header:
                    template_col_index[header] = cell.column

            # 先清空旧数据区域（仅清空已映射列）
            if ws.max_row >= start_row:
                for r in range(start_row, ws.max_row + 1):
                    for t_col in mapping.keys():
                        c_idx = template_col_index.get(t_col)
                        if c_idx:
                            ws.cell(row=r, column=c_idx, value=None)

            # 写入筛选结果
            row_ptr = start_row
            for data_row in filtered_rows:
                for t_col, s_col in mapping.items():
                    c_idx = template_col_index.get(t_col)
                    if not c_idx:
                        continue
                    ws.cell(row=row_ptr, column=c_idx, value=self._json_serializable(data_row.get(s_col)))
                row_ptr += 1

            output_file = str(parameters.get("template_output_file", "")).strip() if parameters else ""
            if output_file:
                output_path = Path(output_file)
            else:
                output_dir = Path(get_config().output_dir)
                output_dir.mkdir(parents=True, exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = output_dir / f"filled_template_{timestamp}.xlsx"

            output_path.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(output_path))
            return str(output_path)
        finally:
            wb.close()

    def _fill_excel_template_multi(
        self,
        template_path: str,
        target_results: List[Dict[str, Any]],
        parameters: Dict[str, Any],
    ) -> str:
        wb = load_workbook(template_path)
        try:
            for target in target_results:
                sheet_name = str(target.get("sheet_name", "")).strip()
                sheet_index = int(target.get("table_index", 0))
                ws = self._resolve_excel_sheet(wb, sheet_name, sheet_index)

                header_row = int(target.get("header_row", 1))
                start_row = int(target.get("start_row", header_row + 1))
                mapping = target.get("mapping", {})
                filtered_rows = target.get("filtered_rows", [])

                template_col_index: Dict[str, int] = {}
                for cell in ws[header_row]:
                    header = str(cell.value).strip() if cell.value is not None else ""
                    if header:
                        template_col_index[header] = cell.column

                if ws.max_row >= start_row:
                    for r in range(start_row, ws.max_row + 1):
                        for t_col in mapping.keys():
                            c_idx = template_col_index.get(t_col)
                            if c_idx:
                                ws.cell(row=r, column=c_idx, value=None)

                row_ptr = start_row
                for data_row in filtered_rows:
                    for t_col, s_col in mapping.items():
                        c_idx = template_col_index.get(t_col)
                        if not c_idx:
                            continue
                        ws.cell(row=row_ptr, column=c_idx, value=self._json_serializable(data_row.get(s_col)))
                    row_ptr += 1

            output_file = str(parameters.get("template_output_file", "")).strip() if parameters else ""
            if output_file:
                output_path = Path(output_file)
            else:
                output_dir = Path(get_config().output_dir)
                output_dir.mkdir(parents=True, exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = output_dir / f"filled_template_multi_{timestamp}.xlsx"

            output_path.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(output_path))
            return str(output_path)
        finally:
            wb.close()

    def _fill_docx_template(
        self,
        filtered_rows: List[Dict[str, Any]],
        template_path: str,
        mapping: Dict[str, str],
        parameters: Dict[str, Any],
    ) -> str:
        doc = Document(template_path)
        table_index = int(parameters.get("template_table_index", 0)) if parameters else 0
        header_row = int(parameters.get("template_header_row", 1)) if parameters else 1
        start_row = int(parameters.get("template_start_row", header_row + 1)) if parameters else (header_row + 1)

        if table_index < 0 or table_index >= len(doc.tables):
            raise ValueError(f"Word模板中不存在 table_index={table_index} 的表格")

        table = doc.tables[table_index]
        h_idx = max(header_row - 1, 0)
        s_idx = max(start_row - 1, 0)

        if h_idx >= len(table.rows):
            raise ValueError(f"Word模板表格表头行不存在: header_row={header_row}")

        header_cells = table.rows[h_idx].cells
        template_col_index: Dict[str, int] = {}
        for i, cell in enumerate(header_cells):
            header = str(cell.text).strip()
            if header:
                template_col_index[header] = i

        col_count = len(header_cells)

        # 清空旧数据区域
        if len(table.rows) > s_idx:
            for r in range(s_idx, len(table.rows)):
                row_cells = table.rows[r].cells
                for t_col in mapping.keys():
                    c_idx = template_col_index.get(t_col)
                    if c_idx is not None and c_idx < len(row_cells):
                        row_cells[c_idx].text = ""

        # 确保有足够行
        need_rows = s_idx + len(filtered_rows)
        while len(table.rows) < need_rows:
            new_row = table.add_row()
            # 清空新增行默认内容
            for i in range(min(col_count, len(new_row.cells))):
                new_row.cells[i].text = ""

        # 写入数据
        for i, data_row in enumerate(filtered_rows):
            row_cells = table.rows[s_idx + i].cells
            for t_col, s_col in mapping.items():
                c_idx = template_col_index.get(t_col)
                if c_idx is None or c_idx >= len(row_cells):
                    continue
                value = self._json_serializable(data_row.get(s_col))
                row_cells[c_idx].text = "" if value is None else str(value)

        output_file = str(parameters.get("template_output_file", "")).strip() if parameters else ""
        if output_file:
            output_path = Path(output_file)
        else:
            output_dir = Path(get_config().output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"filled_template_{timestamp}.docx"

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return str(output_path)

    def _fill_docx_template_multi(
        self,
        template_path: str,
        target_results: List[Dict[str, Any]],
        parameters: Dict[str, Any],
    ) -> str:
        doc = Document(template_path)

        for target in target_results:
            table_index = int(target.get("table_index", 0))
            header_row = int(target.get("header_row", 1))
            start_row = int(target.get("start_row", header_row + 1))
            mapping = target.get("mapping", {})
            filtered_rows = target.get("filtered_rows", [])

            if table_index < 0 or table_index >= len(doc.tables):
                raise ValueError(f"Word模板中不存在 table_index={table_index} 的表格")

            table = doc.tables[table_index]
            h_idx = max(header_row - 1, 0)
            s_idx = max(start_row - 1, 0)

            if h_idx >= len(table.rows):
                raise ValueError(f"Word模板表格表头行不存在: header_row={header_row}")

            header_cells = table.rows[h_idx].cells
            template_col_index: Dict[str, int] = {}
            for i, cell in enumerate(header_cells):
                header = str(cell.text).strip()
                if header:
                    template_col_index[header] = i

            col_count = len(header_cells)

            if len(table.rows) > s_idx:
                for r in range(s_idx, len(table.rows)):
                    row_cells = table.rows[r].cells
                    for t_col in mapping.keys():
                        c_idx = template_col_index.get(t_col)
                        if c_idx is not None and c_idx < len(row_cells):
                            row_cells[c_idx].text = ""

            need_rows = s_idx + len(filtered_rows)
            while len(table.rows) < need_rows:
                new_row = table.add_row()
                for i in range(min(col_count, len(new_row.cells))):
                    new_row.cells[i].text = ""

            for i, data_row in enumerate(filtered_rows):
                row_cells = table.rows[s_idx + i].cells
                for t_col, s_col in mapping.items():
                    c_idx = template_col_index.get(t_col)
                    if c_idx is None or c_idx >= len(row_cells):
                        continue
                    value = self._json_serializable(data_row.get(s_col))
                    row_cells[c_idx].text = "" if value is None else str(value)

        output_file = str(parameters.get("template_output_file", "")).strip() if parameters else ""
        if output_file:
            output_path = Path(output_file)
        else:
            output_dir = Path(get_config().output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"filled_template_multi_{timestamp}.docx"

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return str(output_path)

    def _build_grouped_plan_from_numbered_blocks(self, instruction: str, columns: List[str]) -> Dict[str, Any]:
        marker = re.search(r"(?:^|\n|；|;)\s*\d+[\.、]\s*", instruction)
        if not marker:
            return {}

        numbered_text = instruction[marker.start():]
        blocks = re.split(r"(?:^|\n|；|;)\s*\d+[\.、]\s*", numbered_text, flags=re.M)
        groups = []
        for block in blocks:
            text = block.strip()
            if not text:
                continue

            conditions = self._extract_conditions_from_field_value_pairs(text, columns)
            if not conditions:
                parts = [p.strip() for p in re.split(r"[\n；;]+", text) if p.strip()]
                for part in parts:
                    field = self._best_match_column(part, columns)
                    if not field:
                        continue
                    condition = self._parse_segment_to_condition(part, field)
                    if condition:
                        conditions.append(condition)

            if conditions:
                groups.append({"logic": "and", "conditions": conditions})

        if not groups:
            return {}
        return {"logic": "or", "groups": groups, "conditions": []}

    def _extract_conditions_from_field_value_pairs(self, text: str, columns: List[str]) -> List[Dict[str, Any]]:
        conditions: List[Dict[str, Any]] = []

        valid_cols = [str(c).strip() for c in columns if str(c).strip()]
        if not valid_cols:
            return conditions

        # Match "字段:值" pairs and stop at the next known field label / line break / semicolon.
        field_union = "|".join(re.escape(c) for c in sorted(valid_cols, key=len, reverse=True))
        for col in columns:
            col_name = str(col).strip()
            if not col_name:
                continue

            pattern = rf"{re.escape(col_name)}\s*[：:]\s*(.+?)(?=(?:\s+(?:{field_union})\s*[：:])|[\n；;]|$)"
            matches = re.findall(pattern, text)
            for raw in matches:
                value = str(raw).strip().rstrip("。,.，")
                if value == "":
                    continue
                conditions.append({
                    "field": col_name,
                    "operator": "eq",
                    "value": value,
                })

        deduped: List[Dict[str, Any]] = []
        seen = set()
        for c in conditions:
            key = (c["field"], c["operator"], str(c.get("value", "")))
            if key in seen:
                continue
            seen.add(key)
            deduped.append(c)
        return deduped

    def _parse_segment_to_condition(self, segment: str, field: str) -> Optional[Dict[str, Any]]:
        if "不为空" in segment or "非空" in segment:
            return {"field": field, "operator": "not_null"}
        if "为空" in segment:
            return {"field": field, "operator": "is_null"}

        # 日期区间（优先于数值区间）：如 2020/7/1 到 2020/8/31
        m_date_between = re.search(
            r"(?:从|在)?\s*(\d{4}[/-]\d{1,2}[/-]\d{1,2}(?:\s+\d{1,2}:\d{1,2}(?::\d{1,2}(?:\.\d+)?)?)?)\s*(?:到|至|between)\s*(\d{4}[/-]\d{1,2}[/-]\d{1,2}(?:\s+\d{1,2}:\d{1,2}(?::\d{1,2}(?:\.\d+)?)?)?)",
            segment,
            re.I,
        )
        if m_date_between:
            return {
                "field": field,
                "operator": "between",
                "value": m_date_between.group(1).strip(),
                "value2": m_date_between.group(2).strip(),
            }

        m_between = re.search(r"(?:在)?\s*([\-]?[0-9]+(?:\.[0-9]+)?)\s*(?:到|至|between)\s*([\-]?[0-9]+(?:\.[0-9]+)?)", segment, re.I)
        if m_between:
            return {
                "field": field,
                "operator": "between",
                "value": self._to_number_if_possible(m_between.group(1)),
                "value2": self._to_number_if_possible(m_between.group(2)),
            }

        patterns = [
            (r"(>=|大于等于|不少于)", "gte"),
            (r"(<=|小于等于|不大于)", "lte"),
            (r"(>|大于|高于|超过)", "gt"),
            (r"(<|小于|低于)", "lt"),
            (r"(不等于|!=)", "ne"),
            (r"(等于|=)", "eq"),
            (r"(不包含)", "not_contains"),
            (r"(包含)", "contains"),
        ]

        for pattern, op in patterns:
            m = re.search(pattern, segment, re.I)
            if m:
                value = segment[m.end():].strip(" ：:，,。")
                if op in {"eq", "ne", "gt", "gte", "lt", "lte"}:
                    value = self._to_number_if_possible(value)
                return {"field": field, "operator": op, "value": value}

        m_in = re.search(r"(?:属于|是|为)\s*([\w\u4e00-\u9fa5,，、\s]+)$", segment)
        if m_in and any(sep in m_in.group(1) for sep in [",", "，", "、"]):
            values = [s.strip() for s in re.split(r"[,，、]", m_in.group(1)) if s.strip()]
            return {"field": field, "operator": "in", "values": values}

        m_default = re.search(r"[：: ]+(.+)$", segment)
        if m_default:
            value = m_default.group(1).strip()
            return {"field": field, "operator": "eq", "value": self._to_number_if_possible(value)}

        return None

    def _best_match_column(self, text: str, columns: List[str]) -> Optional[str]:
        lower_text = text.lower()
        for col in columns:
            if str(col).lower() in lower_text:
                return col

        text_tokens = [token for token in re.split(r"\W+", lower_text) if token]
        best_col = None
        best_score = 0
        for col in columns:
            col_l = str(col).lower()
            score = sum(1 for t in text_tokens if t and t in col_l)
            if score > best_score:
                best_score = score
                best_col = col
        return best_col if best_score > 0 else None

    def _extract_json_object(self, text: str) -> Dict[str, Any]:
        if not text:
            return {}
        cleaned = text.strip()

        fenced = re.search(r"```(?:json)?\s*(\{[\s\S]*\})\s*```", cleaned, re.I)
        if fenced:
            cleaned = fenced.group(1)

        try:
            obj = json.loads(cleaned)
            return obj if isinstance(obj, dict) else {}
        except json.JSONDecodeError:
            pass

        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                obj = json.loads(cleaned[start:end + 1])
                return obj if isinstance(obj, dict) else {}
            except json.JSONDecodeError:
                return {}
        return {}

    def _sanitize_filter_plan(self, plan: Dict[str, Any], columns: List[str]) -> Dict[str, Any]:
        if not isinstance(plan, dict):
            return {}

        logic = str(plan.get("logic", "and")).lower()
        if logic not in {"and", "or"}:
            logic = "and"

        conditions = plan.get("conditions", [])
        if not isinstance(conditions, list):
            conditions = []

        groups = plan.get("groups", [])
        if not isinstance(groups, list):
            groups = []

        sanitized = []
        for cond in conditions:
            if not isinstance(cond, dict):
                continue
            field = cond.get("field")
            operator = str(cond.get("operator", "")).lower().strip()
            if field not in columns or operator not in self._operator_map:
                continue

            item = {"field": field, "operator": operator}

            # 兼容 LLM 将 between 错写为 values=[low, high]
            if operator == "between" and isinstance(cond.get("values"), list):
                vv = cond.get("values", [])
                if len(vv) >= 2:
                    item["value"] = vv[0]
                    item["value2"] = vv[1]

            if "value" in cond:
                item["value"] = cond["value"]
            if "value2" in cond:
                item["value2"] = cond["value2"]
            if "values" in cond and isinstance(cond.get("values"), list):
                item["values"] = cond["values"]
            sanitized.append(item)

        sanitized_groups = []
        for group in groups:
            if not isinstance(group, dict):
                continue
            g_logic = str(group.get("logic", "and")).lower()
            if g_logic not in {"and", "or"}:
                g_logic = "and"

            g_conditions = group.get("conditions", [])
            if not isinstance(g_conditions, list):
                continue

            g_sanitized = []
            for cond in g_conditions:
                if not isinstance(cond, dict):
                    continue
                field = cond.get("field")
                operator = str(cond.get("operator", "")).lower().strip()
                if field not in columns or operator not in self._operator_map:
                    continue
                item = {"field": field, "operator": operator}
                if operator == "between" and isinstance(cond.get("values"), list):
                    vv = cond.get("values", [])
                    if len(vv) >= 2:
                        item["value"] = vv[0]
                        item["value2"] = vv[1]
                if "value" in cond:
                    item["value"] = cond["value"]
                if "value2" in cond:
                    item["value2"] = cond["value2"]
                if "values" in cond and isinstance(cond.get("values"), list):
                    item["values"] = cond["values"]
                g_sanitized.append(item)

            if g_sanitized:
                sanitized_groups.append({"logic": g_logic, "conditions": g_sanitized})

        return {"logic": logic, "conditions": sanitized, "groups": sanitized_groups}

    def _apply_filter_plan(self, rows: List[Dict[str, Any]], plan: Dict[str, Any]) -> List[Dict[str, Any]]:
        conditions = plan.get("conditions", [])
        groups = plan.get("groups", [])
        logic = plan.get("logic", "and")
        if not conditions and not groups:
            return rows

        filtered = []
        for row in rows:
            root_results = []

            if conditions:
                cond_results = [self._match_single_condition(row, cond) for cond in conditions]
                root_results.append(all(cond_results) if logic == "and" else any(cond_results))

            for group in groups:
                g_logic = group.get("logic", "and")
                g_conditions = group.get("conditions", [])
                if not g_conditions:
                    continue
                g_results = [self._match_single_condition(row, cond) for cond in g_conditions]
                group_matched = all(g_results) if g_logic == "and" else any(g_results)
                root_results.append(group_matched)

            matched = all(root_results) if logic == "and" else any(root_results)
            if matched:
                filtered.append(row)
        return filtered

    def _match_single_condition(self, row: Dict[str, Any], condition: Dict[str, Any]) -> bool:
        field = condition.get("field")
        operator = condition.get("operator")
        if field not in row or operator not in self._operator_map:
            return False

        row_value = row.get(field)
        return self._operator_map[operator](row_value, condition)

    def _op_eq(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        target = condition.get("value")
        n1, n2 = self._to_number_if_possible(row_value), self._to_number_if_possible(target)
        if isinstance(n1, (int, float)) and isinstance(n2, (int, float)):
            return n1 == n2
        return self._normalize_str(row_value) == self._normalize_str(target)

    def _op_ne(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return not self._op_eq(row_value, condition)

    def _op_gt(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return self._compare_numeric_or_text(row_value, condition.get("value"), op="gt")

    def _op_gte(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return self._compare_numeric_or_text(row_value, condition.get("value"), op="gte")

    def _op_lt(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return self._compare_numeric_or_text(row_value, condition.get("value"), op="lt")

    def _op_lte(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return self._compare_numeric_or_text(row_value, condition.get("value"), op="lte")

    def _op_contains(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        target = self._normalize_str(condition.get("value"))
        return target in self._normalize_str(row_value)

    def _op_not_contains(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return not self._op_contains(row_value, condition)

    def _op_in(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        values = condition.get("values")
        if not isinstance(values, list):
            values = [condition.get("value")]
        normalized_set = {self._normalize_str(v) for v in values}
        return self._normalize_str(row_value) in normalized_set

    def _op_not_in(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return not self._op_in(row_value, condition)

    def _op_between(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        low_dt = self._to_datetime_if_possible(condition.get("value"))
        high_dt = self._to_datetime_if_possible(condition.get("value2"))
        current_dt = self._to_datetime_if_possible(row_value)
        if low_dt is not None and high_dt is not None and current_dt is not None:
            return low_dt <= current_dt <= high_dt

        low = self._to_number_if_possible(condition.get("value"))
        high = self._to_number_if_possible(condition.get("value2"))
        current = self._to_number_if_possible(row_value)
        if not isinstance(low, (int, float)) or not isinstance(high, (int, float)):
            return False
        if not isinstance(current, (int, float)):
            return False
        return low <= current <= high

    def _op_regex(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        pattern = str(condition.get("value", ""))
        if not pattern:
            return False
        return re.search(pattern, self._normalize_str(row_value)) is not None

    def _op_is_null(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return self._is_empty_like(row_value)

    def _op_not_null(self, row_value: Any, condition: Dict[str, Any]) -> bool:
        return not self._op_is_null(row_value, condition)

    def _is_empty_like(self, value: Any) -> bool:
        if value is None:
            return True

        text = str(value).strip()
        if text == "":
            return True

        # Excel 中常用横线表示缺失值，也按空值处理。
        return text in {"-", "--", "—", "——"}

    def _compare_numeric_or_text(self, left: Any, right: Any, op: str) -> bool:
        ld = self._to_datetime_if_possible(left)
        rd = self._to_datetime_if_possible(right)
        if ld is not None and rd is not None:
            if op == "gt":
                return ld > rd
            if op == "gte":
                return ld >= rd
            if op == "lt":
                return ld < rd
            if op == "lte":
                return ld <= rd

        ln = self._to_number_if_possible(left)
        rn = self._to_number_if_possible(right)
        if isinstance(ln, (int, float)) and isinstance(rn, (int, float)):
            if op == "gt":
                return ln > rn
            if op == "gte":
                return ln >= rn
            if op == "lt":
                return ln < rn
            if op == "lte":
                return ln <= rn

        ls, rs = self._normalize_str(left), self._normalize_str(right)
        if op == "gt":
            return ls > rs
        if op == "gte":
            return ls >= rs
        if op == "lt":
            return ls < rs
        if op == "lte":
            return ls <= rs
        return False

    def _to_number_if_possible(self, value: Any) -> Any:
        if value is None:
            return value
        if isinstance(value, (int, float)):
            return value
        if isinstance(value, bool):
            return value

        text = str(value).strip()
        if text == "":
            return value

        if re.fullmatch(r"[\-]?\d+", text):
            try:
                return int(text)
            except Exception:
                return value
        if re.fullmatch(r"[\-]?\d+\.\d+", text):
            try:
                return float(text)
            except Exception:
                return value

        return value

    def _to_datetime_if_possible(self, value: Any) -> Optional[datetime]:
        if value is None:
            return None
        if isinstance(value, datetime):
            return value
        if isinstance(value, date):
            return datetime.combine(value, datetime.min.time())

        text = str(value).strip()
        if not text:
            return None

        normalized = text.replace("年", "-").replace("月", "-").replace("日", "")
        normalized = normalized.replace("/", "-")
        normalized = re.sub(r"\s+", " ", normalized).strip()
        # 常见数据里会出现秒后附加 .0，如 09:00:00.0
        normalized = re.sub(r"(\d{2}:\d{2}:\d{2})\.\d+$", r"\1", normalized)

        formats = [
            "%Y-%m-%d %H:%M:%S",
            "%Y-%m-%d %H:%M",
            "%Y-%m-%d",
        ]
        for fmt in formats:
            try:
                return datetime.strptime(normalized, fmt)
            except Exception:
                continue
        return None

    def _normalize_str(self, value: Any) -> str:
        if self._is_empty_like(value):
            return ""
        if isinstance(value, (datetime, date)):
            return value.isoformat().lower().strip()
        return str(value).lower().strip()

    def _json_serializable(self, value: Any) -> Any:
        if isinstance(value, (datetime, date)):
            return value.isoformat()
        return value

    def _write_rows_to_json(self, rows: List[Dict[str, Any]], output_file: Optional[str]) -> str:
        if output_file:
            output_path = Path(output_file)
        else:
            output_dir = Path(get_config().output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"filtered_rows_{timestamp}.json"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        json_rows = []
        for row in rows:
            json_rows.append({k: self._json_serializable(v) for k, v in row.items()})

        with output_path.open("w", encoding="utf-8") as f:
            json.dump(json_rows, f, ensure_ascii=False, indent=2)

        return str(output_path)


# 兼容：执行器/服务层默认从 core.agents.agent_d 导入 AgentD。
# 这里提供一个别名，不改变现有实现逻辑。
AgentD = AgentB


def _project_root() -> Path:
    """Return repository root directory.

    This file is located at: <root>/src/core/agents/agent_d.py
    """
    return Path(__file__).resolve().parents[3]


def _load_dotenv_file(dotenv_path: Path) -> None:
    """Load a simple .env file into os.environ.

    - Supports lines like KEY=VALUE
    - Supports optional leading `export `
    - Only fills missing env vars via os.environ.setdefault
    """

    if not dotenv_path.exists():
        return

    for raw_line in dotenv_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue

        if line.startswith("export "):
            line = line[len("export ") :].strip()

        if "=" not in line:
            continue

        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip()
        if not key:
            continue

        if (value.startswith('"') and value.endswith('"')) or (
            value.startswith("'") and value.endswith("'")
        ):
            value = value[1:-1]

        os.environ.setdefault(key, value)


def _load_data_module(path: Path):
    spec = importlib.util.spec_from_file_location("agent_d_api_data_module", str(path))
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


def _read_optional_str(module, name: str, default: str = "") -> str:
    value = getattr(module, name, default)
    if value is None:
        return default
    return str(value).strip()


def _resolve_input_path(value: str, base_dir: Path) -> str:
    path = Path(str(value).strip())
    if path.is_absolute():
        return str(path.resolve())
    return str((base_dir / path).resolve())


def _read_optional_int(module, name: str, default: int | None = None):
    value = getattr(module, name, default)
    if value is None:
        return None
    try:
        return int(value)
    except Exception:
        return default


def _read_optional_bool(module, name: str, default: bool):
    value = getattr(module, name, default)
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "on"}:
        return True
    if text in {"0", "false", "no", "n", "off"}:
        return False
    return default


def _read_optional_list(module, name: str):
    value = getattr(module, name, None)
    if isinstance(value, list):
        return value
    return None


def _infer_table_targets_from_prompt(prompt: str) -> List[Dict[str, Any]]:
    """Infer table_targets from a block-style prompt.

    Expected style:
      表一：
        监测时间：...
        城市：...
      表二：
        ...

    Returns empty list when prompt does not look like a multi-table instruction.
    """

    text = str(prompt or "").strip()
    if not text:
        return []

    inferred: List[Dict[str, Any]] = []

    # 形态1：块级样式（标题独占一行）
    block_markers = list(
        re.finditer(
            r"^\s*((?:表|目标|场景)[^\n：:]{0,60})\s*[：:]\s*$",
            text,
            flags=re.M,
        )
    )
    if len(block_markers) >= 2:
        for idx, marker in enumerate(block_markers):
            name = marker.group(1).strip() or f"表{idx + 1}"
            start = marker.end()
            end = block_markers[idx + 1].start() if idx + 1 < len(block_markers) else len(text)
            block = text[start:end].strip()

            lines = [line.strip() for line in block.splitlines() if line.strip()]
            # 用分号串联行内条件，避免字段值吞并下一字段（如“监测时间 ... 城市 ...”）。
            instruction = "；".join(lines).strip()
            if not instruction:
                continue

            inferred.append(
                {
                    "name": name,
                    "instruction": instruction,
                    "template_sheet_name": name,
                    "template_header_row": 1,
                    "template_start_row": 2,
                    "template_table_index": idx,
                }
            )

        if len(inferred) >= 2:
            return inferred

    # 形态2：行内样式（多个目标可能出现在同一行）
    inline_markers = list(
        re.finditer(
            r"((?:表|目标|场景)\s*[A-Za-z0-9一二三四五六七八九十]{1,8})\s*[：:]",
            text,
            flags=re.M,
        )
    )
    if len(inline_markers) < 2:
        return []

    for idx, marker in enumerate(inline_markers):
        name = str(marker.group(1)).strip() or f"表{idx + 1}"
        start = marker.end()
        end = inline_markers[idx + 1].start() if idx + 1 < len(inline_markers) else len(text)
        instruction = str(text[start:end]).strip().strip("。；;")
        if not instruction:
            continue
        inferred.append(
            {
                "name": name,
                "instruction": instruction,
                "template_sheet_name": name,
                "template_header_row": 1,
                "template_start_row": 2,
                "template_table_index": idx,
            }
        )

    return inferred if len(inferred) >= 2 else []


def _extract_json_object_from_text(text: str) -> Dict[str, Any]:
    content = str(text or "").strip()
    if not content:
        return {}

    try:
        obj = json.loads(content)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        pass

    start = content.find("{")
    end = content.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return {}

    fragment = content[start : end + 1]
    try:
        obj = json.loads(fragment)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        return {}


def _sanitize_inferred_table_targets(raw_targets: Any) -> List[Dict[str, Any]]:
    if not isinstance(raw_targets, list):
        return []

    result: List[Dict[str, Any]] = []
    for idx, item in enumerate(raw_targets):
        if not isinstance(item, dict):
            continue

        instruction = str(item.get("instruction") or item.get("condition") or "").strip()
        if not instruction:
            continue

        name = str(item.get("name") or f"表{idx + 1}").strip() or f"表{idx + 1}"
        sheet_name = str(item.get("template_sheet_name") or name).strip() or name

        try:
            table_index = int(item.get("template_table_index", idx))
        except Exception:
            table_index = idx

        try:
            header_row = int(item.get("template_header_row", 1))
        except Exception:
            header_row = 1

        try:
            start_row = int(item.get("template_start_row", header_row + 1))
        except Exception:
            start_row = header_row + 1

        result.append(
            {
                "name": name,
                "instruction": instruction,
                "template_sheet_name": sheet_name,
                "template_header_row": max(1, header_row),
                "template_start_row": max(1, start_row),
                "template_table_index": max(0, table_index),
            }
        )

    if len(result) < 2:
        return []

    # 按 table_index 排序并重新编号，避免 LLM 给出跳号/重复号。
    result.sort(key=lambda x: int(x.get("template_table_index", 0)))
    for idx, item in enumerate(result):
        item["template_table_index"] = idx
        if not str(item.get("template_sheet_name", "")).strip():
            item["template_sheet_name"] = f"表{idx + 1}"

    return result


_REFERENCE_TOKENS = ["同上", "同前", "同第一条", "同上一条"]


def _contains_reference_token(text: str) -> bool:
    content = str(text or "")
    if any(token in content for token in _REFERENCE_TOKENS):
        return True
    return bool(re.search(r"同目标\s*[A-Za-z一二三四五六七八九十]", content))


def _normalize_pair_field(field: str) -> str:
    name = str(field or "").strip()
    if name == "时间":
        return "监测时间"
    return name


def _extract_field_pairs_from_instruction(instruction: str) -> Dict[str, str]:
    text = str(instruction or "")
    pairs: Dict[str, str] = {}

    # 通用“字段:值”抽取
    for key, value in re.findall(r"([\u4e00-\u9fa5A-Za-z][\u4e00-\u9fa5A-Za-z0-9_\.（）()]{0,19})\s*[：:]\s*([^，。；;\n]+)", text):
        k = _normalize_pair_field(key)
        v = str(value).strip().strip("。；;")
        if k and v:
            pairs[k] = v

    # 常见口语表达补充
    simple_fields = ["监测时间", "时间", "城市", "区", "站点名称", "日期"]
    for field in simple_fields:
        pattern = rf"{re.escape(field)}\s*(?:是|为|取|锁定|=)\s*([^，。；;\n]+)"
        m = re.search(pattern, text)
        if m:
            k = _normalize_pair_field(field)
            v = str(m.group(1)).strip().strip("。；;")
            if k and v:
                pairs[k] = v

    for field in simple_fields:
        pattern = rf"{re.escape(field)}\s*(同上|同前|同第一条|同上一条|同目标\s*[A-Za-z一二三四五六七八九十])"
        m = re.search(pattern, text)
        if m:
            k = _normalize_pair_field(field)
            v = str(m.group(1)).replace(" ", "").strip()
            if k and v:
                pairs[k] = v

    return pairs


def _ref_label_to_index(label: str) -> Optional[int]:
    s = str(label or "").strip().replace(" ", "")
    if not s:
        return None

    if re.fullmatch(r"[A-Za-z]", s):
        return ord(s.upper()) - ord("A")

    cn_map = {
        "一": 0,
        "二": 1,
        "三": 2,
        "四": 3,
        "五": 4,
        "六": 5,
        "七": 6,
        "八": 7,
        "九": 8,
        "十": 9,
    }
    return cn_map.get(s)


def _select_reference_source_index(instruction: str, current_index: int) -> Optional[int]:
    text = str(instruction or "")
    if "同第一条" in text:
        return 0
    if any(token in text for token in ["同上一条", "同上", "同前"]):
        return current_index - 1 if current_index - 1 >= 0 else None

    m = re.search(r"同目标\s*([A-Za-z一二三四五六七八九十])", text)
    if m:
        return _ref_label_to_index(m.group(1))
    return None


def _format_instruction_from_pairs(pairs: Dict[str, str]) -> str:
    if not pairs:
        return ""

    ordered: List[Tuple[str, str]] = []
    preferred = ["监测时间", "城市", "区", "站点名称", "日期"]
    for key in preferred:
        if key in pairs:
            ordered.append((key, pairs[key]))
    for key, value in pairs.items():
        if key not in preferred:
            ordered.append((key, value))

    return "；".join([f"{k}：{v}" for k, v in ordered if str(v).strip()])


def _resolve_table_target_references_by_rule(targets: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], bool]:
    if not isinstance(targets, list) or len(targets) < 2:
        return targets, False

    resolved: List[Dict[str, Any]] = []
    history_pairs: List[Dict[str, str]] = []
    unresolved = False

    for idx, target in enumerate(targets):
        item = dict(target)
        instruction = str(item.get("instruction") or "").strip()
        pairs = _extract_field_pairs_from_instruction(instruction)

        src_idx = _select_reference_source_index(instruction, idx)
        if src_idx is not None and 0 <= src_idx < len(history_pairs):
            source_pairs = history_pairs[src_idx]
            for key, val in source_pairs.items():
                current_val = str(pairs.get(key, "")).strip()
                if not current_val or _contains_reference_token(current_val):
                    pairs[key] = val

        rebuilt = _format_instruction_from_pairs(pairs)
        if rebuilt:
            item["instruction"] = rebuilt

        if _contains_reference_token(str(item.get("instruction") or "")):
            unresolved = True

        history_pairs.append(pairs)
        resolved.append(item)

    return resolved, unresolved


def _resolve_table_target_references_with_llm(targets: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not isinstance(targets, list) or len(targets) < 2:
        return targets

    try:
        llm = get_llm_service()
    except Exception:
        return targets

    if not llm or not llm.is_available():
        return targets

    payload = []
    for idx, t in enumerate(targets):
        payload.append(
            {
                "index": idx,
                "name": str(t.get("name") or f"目标{idx + 1}"),
                "instruction": str(t.get("instruction") or "").strip(),
            }
        )

    system_prompt = (
        "你是多目标填表指代消解器。"
        "任务：将每个目标中的“同上/同前/同第一条/同目标A”等指代，改写成自包含条件。"
        "输出必须为 JSON 对象本体，禁止解释。"
        "固定格式：{\"table_targets\":[{\"index\":0,\"instruction\":\"...\"}] }。"
        "约束："
        "1) 不得改写已明确给出的城市/时间值；"
        "2) 仅补全缺失条件；"
        "3) 若无法确定，保留原 instruction。"
    )
    user_prompt = f"目标列表：{json.dumps(payload, ensure_ascii=False)}\n请返回 JSON。"

    try:
        response = llm.chat_with_system(system_prompt=system_prompt, user_input=user_prompt)
        obj = _extract_json_object_from_text(response)
    except Exception:
        return targets

    raw = obj.get("table_targets", []) if isinstance(obj, dict) else []
    if not isinstance(raw, list):
        return targets

    updated = [dict(t) for t in targets]
    for item in raw:
        if not isinstance(item, dict):
            continue
        try:
            idx = int(item.get("index", -1))
        except Exception:
            idx = -1
        if idx < 0 or idx >= len(updated):
            continue

        instruction = str(item.get("instruction") or "").strip()
        if instruction:
            updated[idx]["instruction"] = instruction

    return updated


def _resolve_table_target_references(targets: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], str]:
    resolved, unresolved = _resolve_table_target_references_by_rule(targets)
    if not unresolved:
        return resolved, "rule"

    llm_resolved = _resolve_table_target_references_with_llm(resolved)
    still_unresolved = any(_contains_reference_token(str(t.get("instruction") or "")) for t in llm_resolved)
    if still_unresolved:
        return llm_resolved, "llm_partial"
    return llm_resolved, "llm"


def _validate_inferred_table_targets(targets: List[Dict[str, Any]]) -> Tuple[bool, str]:
    if not isinstance(targets, list) or len(targets) < 2:
        return False, "table_targets 数量不足（需要至少2个）"

    required = {"instruction", "template_table_index", "template_header_row", "template_start_row"}
    for i, t in enumerate(targets):
        if not isinstance(t, dict):
            return False, f"table_targets[{i}] 不是对象"
        missing = [k for k in required if k not in t]
        if missing:
            return False, f"table_targets[{i}] 缺少字段: {missing}"
        if not str(t.get("instruction", "")).strip():
            return False, f"table_targets[{i}] instruction 为空"

    indices = [int(t.get("template_table_index", -1)) for t in targets]
    if sorted(indices) != list(range(len(targets))):
        return False, "template_table_index 必须从0开始连续"

    return True, ""


def _multi_target_signal_level(prompt: str) -> str:
    """Return multi-target signal level: strong / weak / none."""

    text = str(prompt or "")
    if not text.strip():
        return "none"

    # 强信号：显式“表一/目标A/场景1:”这类标题 >= 2
    marker_count = len(
        re.findall(
            r"(?:表|目标|场景)[^\n：:]{0,60}\s*[：:]",
            text,
        )
    )
    if marker_count >= 2:
        return "strong"

    # 弱信号：存在多个“时间+城市”条件对，但不一定真是多表。
    city_count = len(re.findall(r"城市\s*(?:是|为|：|:)", text))
    time_count = len(re.findall(r"(?:监测时间|时间)\s*(?:是|为|：|:)", text))
    if city_count >= 2 and time_count >= 2:
        return "weak"

    return "none"


def _infer_table_mode_with_llm(prompt: str) -> str:
    """Infer prompt mode: single / multi / unknown."""

    text = str(prompt or "").strip()
    if not text:
        return "unknown"

    try:
        llm = get_llm_service()
    except Exception:
        return "unknown"

    if not llm or not llm.is_available():
        return "unknown"

    system_prompt = (
        "你是填表任务模式分类器。"
        "只做一件事：判断用户输入是单表填表还是多表填表。"
        "输出必须是 JSON 对象本体，禁止解释、禁止 Markdown。"
        "固定输出格式：{\"mode\":\"single|multi|unknown\",\"reason\":\"...\"}。"
        "判定规则："
        "1) 若明确存在两个及以上独立目标（如表一/表二、目标A/B、场景1/2），mode=multi；"
        "2) 若只是一个筛选目标（即使含多个条件、多个城市列表），mode=single；"
        "3) 出现“两个文件/多个文件/两份数据”等仅表示数据来源数量，不等于多表；若未出现多个独立目标，仍应判为single；"
        "3) 无法可靠判断时，mode=unknown。"
    )
    user_prompt = f"用户输入:\n{text}\n\n请只输出 JSON。"

    try:
        response = llm.chat_with_system(system_prompt=system_prompt, user_input=user_prompt)
        obj = _extract_json_object_from_text(response)
        mode = str(obj.get("mode", "unknown")).strip().lower()
        if mode in {"single", "multi", "unknown"}:
            return mode
    except Exception:
        return "unknown"

    return "unknown"


def _infer_table_targets_from_prompt_with_llm(prompt: str) -> List[Dict[str, Any]]:
    text = str(prompt or "").strip()
    if not text:
        return []

    try:
        llm = get_llm_service()
    except Exception:
        return []

    if not llm or not llm.is_available():
        return []

    system_prompt = (
        "你是多表填表任务的结构化解析器。"
        "任务：仅当用户明确要填写多个独立目标时，才把整段自然语言拆分成 table_targets；"
        "若是单表任务或无法确认多表，必须返回空数组以保护单表流程。"
        "用户文风可能是：表一/表二，目标A/目标B，场景1/场景2，等等。"
        "输出必须为 JSON 对象本体，禁止解释、禁止 Markdown 代码块。"
        "固定输出格式："
        "{\"table_targets\":["
        "{\"name\":\"表一\",\"instruction\":\"监测时间：... 城市：...\","
        "\"template_sheet_name\":\"表一\",\"template_header_row\":1,\"template_start_row\":2,\"template_table_index\":0}"
        "]}。"
        "硬性约束："
        "1) 至少拆分2个目标才允许返回非空 table_targets；"
        "2) 每个目标 instruction 仅包含该目标条件，不要混合其它目标；"
        "3) template_table_index 从0开始连续递增；"
        "4) 单表任务（即使有多个筛选条件）必须返回 {\"table_targets\": []}；"
        "5) 无法可靠拆分时也返回 {\"table_targets\": []}。"
    )

    user_prompt = (
        f"用户输入:\n{text}\n\n"
        "请直接输出 JSON。"
    )

    try:
        response = llm.chat_with_system(system_prompt=system_prompt, user_input=user_prompt)
    except Exception:
        return []

    parsed = _extract_json_object_from_text(response)
    targets = _sanitize_inferred_table_targets(parsed.get("table_targets", []))
    ok, reason = _validate_inferred_table_targets(targets)
    if ok:
        return targets

    # 一次修复重试：把上次错误原因反馈给模型，让其仅修正结构。
    repair_prompt = (
        f"上一次输出不合规，原因：{reason}。\n"
        "请在不改变原始语义的前提下修正结构，并只输出 JSON 对象。\n"
        f"原始用户输入:\n{text}\n"
    )
    try:
        repaired = llm.chat_with_system(system_prompt=system_prompt, user_input=repair_prompt)
    except Exception:
        return []

    repaired_obj = _extract_json_object_from_text(repaired)
    repaired_targets = _sanitize_inferred_table_targets(repaired_obj.get("table_targets", []))
    ok2, _ = _validate_inferred_table_targets(repaired_targets)
    return repaired_targets if ok2 else []


def run_agent_d_api(
    *,
    src: str,
    prompt: str = "",
    template: str = "",
    output_json: str = "",
    output_template: str = "",
    allow_rule_fallback: bool = False,
    table_targets: Optional[List[Dict[str, Any]]] = None,
    template_sheet_name: str = "",
    template_header_row: Optional[int] = None,
    template_start_row: Optional[int] = None,
    template_table_index: Optional[int] = None,
) -> Dict[str, Any]:
    """AgentD 单一调用接口：筛选 Excel +（可选）填表。

    参数与 tests/test_b/data*/data.py 对齐。
    """

    root = _project_root()
    _load_dotenv_file(root / ".env")

    excel_path = Path(str(src)).resolve()

    if output_json:
        resolved_output_json = Path(str(output_json)).resolve()
    else:
        resolved_output_json = excel_path.parent / f"{excel_path.stem}_filtered_rows.json"

    template_file = None
    if template:
        template_path = Path(str(template)).resolve()
        suffix = template_path.suffix.lower()
        template_type = FileType.DOCX if suffix == ".docx" else FileType.XLSX
        template_file = FileInfo(path=str(template_path), file_type=template_type)

    normalized_prompt = str(prompt or "").strip()
    inferred_table_targets: List[Dict[str, Any]] = []
    inferred_source = ""
    reference_resolve_source = ""
    mode_by_llm = "unknown"
    signal_level = _multi_target_signal_level(normalized_prompt)

    # 优先使用外部传入的 table_targets（例如上游 Agent 自动生成）；为空时再从 prompt 推断。
    provided_table_targets = _sanitize_inferred_table_targets(table_targets) if table_targets else []
    if provided_table_targets:
        inferred_table_targets = provided_table_targets
        inferred_source = "external"
    elif normalized_prompt:
        inferred_table_targets = _infer_table_targets_from_prompt(normalized_prompt)
        if inferred_table_targets:
            inferred_source = "rule"
        if not inferred_table_targets:
            inferred_table_targets = _infer_table_targets_from_prompt_with_llm(normalized_prompt)
            if inferred_table_targets:
                inferred_source = "llm"

        if inferred_table_targets:
            inferred_table_targets, reference_resolve_source = _resolve_table_target_references(inferred_table_targets)

        mode_by_llm = _infer_table_mode_with_llm(normalized_prompt)

    # 多目标强信号，或“弱信号+LLM判多表”同时满足时才fail-fast，避免单表被误判。
    should_fail_multi = (
        not inferred_table_targets
        and (
            signal_level == "strong"
            or (signal_level == "weak" and mode_by_llm == "multi")
        )
    )
    if should_fail_multi:
        return {
            "success": False,
            "message": "检测到多目标填表意图，但未能可靠切分为 table_targets。请补充结构化 table_targets 或简化提示词。",
            "data": {
                "status": "failed",
                "reason": "multi_target_inference_failed",
                "multi_target_signal_level": signal_level,
                "multi_target_mode_by_llm": mode_by_llm,
            },
            "resolved_input": {
                "src": str(excel_path),
                "template": template_file.path if template_file else None,
                "output_json": str(resolved_output_json),
                "output_template": str(Path(str(output_template)).resolve()) if output_template else None,
                "table_targets_source": "",
            },
        }

    parameters: Dict[str, Any] = {"allow_rule_fallback": bool(allow_rule_fallback)}

    if inferred_table_targets:
        parameters["table_targets"] = inferred_table_targets
        source_label = inferred_source
        if reference_resolve_source:
            source_label = f"{source_label}+coref_{reference_resolve_source}"
        parameters["table_targets_source"] = source_label

    if output_template:
        parameters["template_output_file"] = str(Path(str(output_template)).resolve())

    if template_sheet_name:
        parameters["template_sheet_name"] = str(template_sheet_name).strip()

    if template_header_row is not None:
        parameters["template_header_row"] = int(template_header_row)

    if template_start_row is not None:
        parameters["template_start_row"] = int(template_start_row)

    if template_table_index is not None:
        parameters["template_table_index"] = int(template_table_index)

    task_spec = TaskSpec(
        task_type=TaskType.TABLE_FILLING,
        instruction=normalized_prompt,
        source_files=[FileInfo(path=str(excel_path), file_type=FileType.XLSX)],
        template_file=template_file,
        output_file=str(resolved_output_json),
        parameters=parameters,
    )

    agent = AgentD()
    result = agent.execute(task_spec)

    return {
        "success": result.success,
        "message": result.message,
        "data": result.data,
        "resolved_input": {
            "src": str(excel_path),
            "template": template_file.path if template_file else None,
            "output_json": str(resolved_output_json),
            "output_template": parameters.get("template_output_file"),
            "table_targets_source": parameters.get("table_targets_source", ""),
        },
    }


def run_agent_d_from_data_file(data_py: Path) -> Dict[str, Any]:
    """从 data.py 读取参数并调用 run_agent_d_api。"""

    m = _load_data_module(Path(data_py))
    project_root = _project_root()

    src = _read_optional_str(m, "src")
    template = _read_optional_str(m, "template")
    output_json = _read_optional_str(m, "output_json")
    output_template = _read_optional_str(m, "output_template")

    return run_agent_d_api(
        src=_resolve_input_path(src, project_root) if src else src,
        prompt=_read_optional_str(m, "prompt"),
        template=_resolve_input_path(template, project_root) if template else template,
        output_json=_resolve_input_path(output_json, project_root) if output_json else output_json,
        output_template=_resolve_input_path(output_template, project_root) if output_template else output_template,
        allow_rule_fallback=_read_optional_bool(m, "allow_rule_fallback", False),
        table_targets=_read_optional_list(m, "table_targets"),
        template_sheet_name=_read_optional_str(m, "template_sheet_name"),
        template_header_row=_read_optional_int(m, "template_header_row", None),
        template_start_row=_read_optional_int(m, "template_start_row", None),
        template_table_index=_read_optional_int(m, "template_table_index", None),
    )


def _normalize_entity_row_value(value: Any) -> Any:
    if isinstance(value, list):
        return value[0] if value else None
    return value


def run_agent_d_fill_from_entities(
    *,
    entities: List[Dict[str, Any]],
    template: str,
    output_json: str = "",
    output_template: str = "",
    template_sheet_name: str = "",
    template_header_row: Optional[int] = None,
    template_start_row: Optional[int] = None,
    template_table_index: Optional[int] = None,
) -> Dict[str, Any]:
    """Fill a single template from merged structured entities.

    This API is used by mixed mode to write the merged docx+xlsx results
    into one final template (xlsx/docx), reusing AgentD's mapping/filling logic.
    """

    root = _project_root()
    _load_dotenv_file(root / ".env")

    template_path = Path(str(template)).resolve()
    if not template_path.exists():
        return {
            "success": False,
            "message": f"模板文件不存在: {template_path}",
            "data": {"status": "failed"},
        }

    if not isinstance(entities, list) or not entities:
        return {
            "success": False,
            "message": "缺少可填表实体数据（entities 为空）",
            "data": {"status": "failed"},
        }

    normalized_rows: List[Dict[str, Any]] = []
    for row in entities:
        if not isinstance(row, dict):
            continue
        normalized_rows.append({k: _normalize_entity_row_value(v) for k, v in row.items()})

    if not normalized_rows:
        return {
            "success": False,
            "message": "entities 中无有效字典数据",
            "data": {"status": "failed"},
        }

    parameters: Dict[str, Any] = {}
    if output_template:
        parameters["template_output_file"] = str(Path(str(output_template)).resolve())
    if template_sheet_name:
        parameters["template_sheet_name"] = str(template_sheet_name).strip()
    if template_header_row is not None:
        parameters["template_header_row"] = int(template_header_row)
    if template_start_row is not None:
        parameters["template_start_row"] = int(template_start_row)
    if template_table_index is not None:
        parameters["template_table_index"] = int(template_table_index)

    agent = AgentD()
    source_columns = list({k for row in normalized_rows for k in row.keys()})
    template_columns = agent._read_template_columns(str(template_path), parameters)
    mapping = agent._build_template_column_mapping(
        source_columns=source_columns,
        template_columns=template_columns,
        source_rows=normalized_rows,
        parameters=parameters,
    )

    if agent._is_excel_template(str(template_path)):
        if "template_output_file" not in parameters:
            parameters["template_output_file"] = str(template_path.parent / f"{template_path.stem}_filled.xlsx")
        template_output = agent._fill_excel_template(
            filtered_rows=normalized_rows,
            template_path=str(template_path),
            mapping=mapping,
            parameters=parameters,
        )
    elif agent._is_docx_template(str(template_path)):
        if "template_output_file" not in parameters:
            parameters["template_output_file"] = str(template_path.parent / f"{template_path.stem}_filled.docx")
        template_output = agent._fill_docx_template(
            filtered_rows=normalized_rows,
            template_path=str(template_path),
            mapping=mapping,
            parameters=parameters,
        )
    else:
        return {
            "success": False,
            "message": f"不支持的模板类型: {template_path.suffix}",
            "data": {"status": "failed"},
        }

    if output_json:
        output_json_path = str(Path(str(output_json)).resolve())
    else:
        output_json_path = str(template_path.parent / f"{template_path.stem}_merged_rows.json")
    output_json_file = agent._write_rows_to_json(normalized_rows, output_json_path)

    return {
        "success": True,
        "message": f"统一填表完成，共写入 {len(normalized_rows)} 行",
        "data": {
            "status": "completed",
            "output_json": output_json_file,
            "template_filled": True,
            "template_output": template_output,
            "matched_rows": len(normalized_rows),
            "total_rows": len(normalized_rows),
            "template_mapping": mapping,
        },
    }
