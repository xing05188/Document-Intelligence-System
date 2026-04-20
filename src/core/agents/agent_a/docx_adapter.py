"""DOCX adapter for executing document actions with python-docx."""

from __future__ import annotations

import json
import re
import copy
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.run import Run

from core.llm.llm_service import get_llm_service

from .standard_style import get_standard_style_preset


@dataclass
class ActionExecutionResult:
    """Result for a single action execution."""

    action_type: str
    success: bool
    message: str
    details: Dict[str, Any]


class DocxAdapter:
    """Adapter that applies action plan items to docx files."""

    def __init__(self, file_path: str, llm_service=None):
        self.file_path = file_path
        self.document = Document(file_path)
        self.execution_log: List[ActionExecutionResult] = []
        self._llm_service = llm_service
        self._semantic_heading_cache: Dict[int, List[int]] = {}
        self._document_understanding_cache: Optional[Dict[str, Any]] = None
        self._document_title_cache: Optional[Dict[str, Any]] = None

    def apply_action(self, action: Dict[str, Any]) -> ActionExecutionResult:
        action_type = action.get("action_type", "")
        params = action.get("params", {}) or {}
        target = action.get("target", {}) or {}

        handler_map = {
            "bold_heading": self._apply_bold_heading,
            "insert_page_number": self._apply_insert_page_number,
            "unify_style": self._apply_unify_style,
            "reorder_paragraphs": self._apply_reorder_paragraphs,
            "batch_format": self._apply_batch_format,
            "extract_content": self._apply_extract_content,
            "replace_text": self._apply_replace_text,
            "insert_toc": self._apply_insert_toc,
            "remove_blank_lines": self._apply_remove_blank_lines,
            "set_font_family": self._apply_set_font_family,
            "set_font_color": self._apply_set_font_color,
            "set_font_size": self._apply_set_font_size,
            "set_paragraph_alignment": self._apply_set_paragraph_alignment,
            "set_line_spacing": self._apply_set_line_spacing,
            "set_first_line_indent": self._apply_set_first_line_indent,
            "set_highlight": self._apply_set_highlight,
            "insert_table": self._apply_insert_table,
            "insert_footer_text": self._apply_insert_footer_text,
            "set_heading_numbering": self._apply_set_heading_numbering,
            "set_italic": self._apply_set_italic,
            "set_underline": self._apply_set_underline,
            "set_paragraph_spacing": self._apply_set_paragraph_spacing,
            "set_bullet_list": self._apply_set_bullet_list,
            "set_numbered_list": self._apply_set_numbered_list,
            "set_paragraph_shading": self._apply_set_paragraph_shading,
            "set_paragraph_border": self._apply_set_paragraph_border,
            "add_hyperlink": self._apply_add_hyperlink,
        }

        handler = handler_map.get(action_type)
        if handler is None:
            result = ActionExecutionResult(
                action_type=action_type,
                success=False,
                message=f"DOCX 适配器暂不支持动作: {action_type}",
                details={},
            )
            self.execution_log.append(result)
            return result

        try:
            details = handler(target, params)
            result = ActionExecutionResult(
                action_type=action_type,
                success=True,
                message="执行成功",
                details=details,
            )
        except Exception as e:
            result = ActionExecutionResult(
                action_type=action_type,
                success=False,
                message=f"执行失败: {e}",
                details={},
            )

        self.execution_log.append(result)
        return result

    def save(self, output_path: str) -> str:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(path))
        return str(path)

    def _apply_bold_heading(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        level = int(target.get("level", 1))
        heading_indices = self._get_heading_indices_for_title_ops(level=level)

        # 一级标题场景下，将文档总标题纳入加粗，保证与标题样式一致。
        if level == 1:
            doc_title = self._get_document_main_title(self._get_document_understanding())
            if isinstance(doc_title, dict):
                try:
                    doc_idx = int(doc_title.get("index", -1))
                except Exception:
                    doc_idx = -1
                if 0 <= doc_idx < len(self.document.paragraphs) and doc_idx not in self._get_toc_block_indices():
                    heading_indices = sorted(set(list(heading_indices) + [doc_idx]))

        style_only_indices = set(self._get_style_heading_indices(level=level))

        count = 0
        for idx in heading_indices:
            if idx < 0 or idx >= len(self.document.paragraphs):
                continue
            paragraph = self.document.paragraphs[idx]
            if not paragraph.text.strip():
                continue
            if not paragraph.runs and paragraph.text:
                paragraph.add_run(paragraph.text)
                paragraph.text = ""
            for run in paragraph.runs:
                run.bold = True
            count += 1

        style_hits = sum(1 for idx in heading_indices if idx in style_only_indices)
        semantic_hits = max(0, count - style_hits)

        return {
            "level": level,
            "matched_headings": count,
            "style_hits": style_hits,
            "semantic_hits": semantic_hits,
            "bold": params.get("bold", True),
        }

    def _apply_insert_page_number(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        alignment = str(params.get("alignment") or params.get("align") or "center").lower()
        alignment_map = {
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        }

        updated_sections = 0
        for section in self.document.sections:
            footer = section.footer
            paragraph = None
            for p in footer.paragraphs:
                if "PAGE" in p._p.xml:
                    paragraph = p
                    break
            if paragraph is None:
                paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                fld_simple = OxmlElement("w:fldSimple")
                fld_simple.set(qn("w:instr"), "PAGE")
                run = OxmlElement("w:r")
                fld_simple.append(run)
                paragraph._p.append(fld_simple)
            paragraph.alignment = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)
            updated_sections += 1

        return {"position": params.get("position", "footer"), "align": alignment, "updated_sections": updated_sections}

    def _apply_unify_style(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        preset = get_standard_style_preset("docx") if str(params.get("style_preset", "")).lower() == "standard" else {}
        body_font_name = str(params.get("body_font_name") or preset.get("body_font_name") or "")
        body_font_size = float(params.get("body_font_size") or preset.get("body_font_size") or 11)
        heading_font_name = str(params.get("heading_font_name") or preset.get("heading_font_name") or body_font_name or "")
        heading_font_size = float(params.get("heading_font_size") or preset.get("heading_font_size") or 14)

        def _as_bool(value: Any, default: bool = False) -> bool:
            if value is None:
                return default
            if isinstance(value, bool):
                return value
            text = str(value).strip().lower()
            if text in {"1", "true", "yes", "y", "on"}:
                return True
            if text in {"0", "false", "no", "n", "off"}:
                return False
            return default

        heading_bold = _as_bool(params.get("heading_bold"), _as_bool(preset.get("heading_bold"), True))

        updated = 0
        for paragraph in self.document.paragraphs:
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            is_heading = "heading" in style_name or self._looks_like_structured_heading((paragraph.text or "").strip())
            if is_heading:
                for run in paragraph.runs:
                    if heading_font_name:
                        run.font.name = heading_font_name
                    run.bold = heading_bold
                    run.font.size = Pt(heading_font_size)
                updated += 1
                continue

            if paragraph.style and paragraph.style.name != "Normal":
                paragraph.style = self.document.styles["Normal"]
            for run in paragraph.runs:
                if body_font_name:
                    run.font.name = body_font_name
                if run.font.size is None or str(params.get("style_preset", "")).lower() == "standard":
                    run.font.size = Pt(body_font_size)
            updated += 1
        return {
            "updated_paragraphs": updated,
            "strategy": params.get("strategy", "standard"),
            "style_preset": params.get("style_preset", "standard"),
        }

    def _apply_reorder_paragraphs(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        from_idx = int(params.get("from", 0))
        to_idx = int(params.get("to", 0))

        candidates = self._get_content_paragraphs_for_indexing()
        if from_idx <= 0 or to_idx <= 0 or from_idx > len(candidates) or to_idx > len(candidates):
            return {"moved": False, "reason": "索引越界", "from": from_idx, "to": to_idx}

        if from_idx == to_idx:
            return {"moved": True, "from": from_idx, "to": to_idx, "reason": "源段落与目标段落相同"}

        src = candidates[from_idx - 1]
        dst = candidates[to_idx - 1]
        src_anchor = self._get_heading_anchor_index_for_paragraph(src)
        dst_anchor = self._get_heading_anchor_index_for_paragraph(dst)
        if src_anchor != dst_anchor:
            return {
                "moved": False,
                "reason": "禁止跨标题块移动段落",
                "from": from_idx,
                "to": to_idx,
            }
        src_elem = src._p
        dst_elem = dst._p

        src_elem.getparent().remove(src_elem)
        dst_elem.addnext(src_elem)
        return {"moved": True, "from": from_idx, "to": to_idx}

    @staticmethod
    def _looks_like_structured_heading(text: str) -> bool:
        s = (text or "").strip()
        if not s:
            return False
        patterns = [
            r"^[一二三四五六七八九十]+、",
            r"^[（(][一二三四五六七八九十]+[）)]",
            r"^第[一二三四五六七八九十\d]+[章节部分篇]",
            r"^\d+(?:\.\d+){0,3}[\.、]",
            r"^#{1,6}\s+",
        ]
        return any(re.match(p, s) for p in patterns)

    def _get_heading_anchor_index_for_paragraph(self, paragraph: Any) -> int:
        paragraphs = list(self.document.paragraphs)
        idx = -1
        for i, p in enumerate(paragraphs):
            if p._p is paragraph._p:
                idx = i
                break
        if idx < 0:
            return -1

        anchor = -1
        for i in range(0, idx + 1):
            p = paragraphs[i]
            style_name = (p.style.name or "").lower() if p.style else ""
            txt = (p.text or "").strip()
            if "heading" in style_name or self._looks_like_structured_heading(txt):
                anchor = i
        return anchor

    def _apply_batch_format(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        updated = 0
        for paragraph in self.document.paragraphs:
            if not paragraph.text.strip():
                continue
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.paragraph_format.line_spacing = 1.5
            for run in paragraph.runs:
                if run.font.size is None:
                    run.font.size = Pt(11)
            updated += 1
        return {"updated_paragraphs": updated, "scope": params.get("scope", "selection")}

    def _apply_extract_content(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        headings: List[str] = []
        body: List[str] = []
        requested_fields = params.get("fields", [])
        extracted_fields: Dict[str, str] = {}
        understanding = self._get_document_understanding()

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style_name:
                headings.append(text)
            else:
                body.append(text)

        semantic_used = False
        llm_understanding_used = False
        if not headings:
            if isinstance(understanding.get("headings"), list) and understanding.get("headings"):
                for item in understanding.get("headings", []):
                    if isinstance(item, dict):
                        title = str(item.get("title", "")).strip()
                        if title:
                            headings.append(title)
                if headings:
                    llm_understanding_used = True

            if not headings:
                semantic_indices = self._get_semantic_heading_indices(level=1)
            else:
                semantic_indices = []

            for idx in semantic_indices:
                if idx < 0 or idx >= len(self.document.paragraphs):
                    continue
                t = self.document.paragraphs[idx].text.strip()
                if t:
                    headings.append(t)
            semantic_used = len(headings) > 0

        if requested_fields and isinstance(understanding.get("field_values"), dict):
            fv = understanding.get("field_values", {})
            for f in requested_fields:
                key = str(f)
                if key in fv and str(fv.get(key, "")).strip():
                    extracted_fields[key] = str(fv.get(key))

        if requested_fields and not extracted_fields:
            extracted_fields = self._extract_fields_heuristically(body, requested_fields)

        return {
            "headings": headings,
            "paragraph_count": len(body),
            "table_count": len(self.document.tables),
            "fields": requested_fields,
            "field_values": extracted_fields,
            "semantic_heading_used": semantic_used,
            "llm_understanding_used": llm_understanding_used,
        }

    def _apply_replace_text(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        find_text = str(params.get("find", ""))
        replace_text = str(params.get("replace", ""))
        if not find_text:
            return {"replaced": 0}

        replaced = 0
        for paragraph in self.document.paragraphs:
            if self._replace_text_in_paragraph_runs(paragraph, find_text, replace_text):
                replaced += 1

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self._replace_text_in_paragraph_runs(paragraph, find_text, replace_text):
                            replaced += 1

        return {"find": find_text, "replace": replace_text, "replaced": replaced}

    @staticmethod
    def _replace_text_in_paragraph_runs(paragraph, find_text: str, replace_text: str) -> bool:
        text = paragraph.text or ""
        if find_text not in text:
            return False

        # 优先在 run 内替换，尽量保留既有粗体/颜色等样式。
        run_changed = False
        for run in paragraph.runs:
            if find_text in (run.text or ""):
                run.text = (run.text or "").replace(find_text, replace_text)
                run_changed = True

        if run_changed:
            return True

        # 跨 run 命中时回退段落级替换；保留首个 run 的关键样式。
        first_run = paragraph.runs[0] if paragraph.runs else None
        keep_bold = bool(first_run.bold) if first_run is not None else None
        keep_name = first_run.font.name if first_run is not None else None
        keep_size = first_run.font.size if first_run is not None else None
        keep_color = first_run.font.color.rgb if first_run is not None else None

        paragraph.text = text.replace(find_text, replace_text)
        if paragraph.runs:
            run = paragraph.runs[0]
            if keep_bold is not None:
                run.bold = keep_bold
            if keep_name is not None:
                run.font.name = keep_name
            if keep_size is not None:
                run.font.size = keep_size
            if keep_color is not None:
                run.font.color.rgb = keep_color

        return True

    def _apply_insert_toc(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        if self._has_existing_toc():
            normalized = self._normalize_existing_toc_block()
            return {
                "toc_entries": 0,
                "semantic_heading_used": False,
                "llm_understanding_used": False,
                "contextualized": False,
                "linked_entries": 0,
                "skipped": True,
                "reason": "目录已存在，跳过重复插入",
                "normalized_lines": normalized,
            }

        headings = []
        toc_entries: List[str] = []
        understanding = self._get_document_understanding()
        for paragraph in self.document.paragraphs:
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style_name and paragraph.text.strip():
                headings.append(paragraph.text.strip())

        semantic_used = False
        llm_understanding_used = False
        if not headings:
            if isinstance(understanding.get("toc_entries"), list) and understanding.get("toc_entries"):
                for entry in understanding.get("toc_entries", []):
                    if str(entry).strip():
                        toc_entries.append(str(entry).strip())

            if isinstance(understanding.get("headings"), list) and understanding.get("headings"):
                for item in understanding.get("headings", []):
                    if not isinstance(item, dict):
                        continue
                    title = str(item.get("title", "")).strip()
                    if title:
                        headings.append(title)
                if headings or toc_entries:
                    llm_understanding_used = True

            if not headings:
                semantic_indices = self._get_semantic_heading_indices(level=1)
            else:
                semantic_indices = []

            for idx in semantic_indices:
                if idx < 0 or idx >= len(self.document.paragraphs):
                    continue
                t = self.document.paragraphs[idx].text.strip()
                if t:
                    headings.append(t)
            semantic_used = len(headings) > 0

        if not toc_entries:
            toc_entries = self._build_contextual_toc_entries(understanding, headings)

        # 借助 LLM 识别文档总标题，并作为目录首项（若尚未包含）。
        doc_title = self._get_document_main_title(understanding)
        if doc_title:
            title_text = str(doc_title.get("title", "")).strip()
            if title_text and title_text not in toc_entries:
                toc_entries = [title_text] + toc_entries

        toc_entries = [e for e in toc_entries if not self._is_caption_like_text(e)]

        # 目录条目最终以 toc_entries 为准：
        # 1) 匹配时允许“去序号归一化”，解决“一、综合”与“综合”匹配问题。
        # 2) 按出现顺序稳定匹配重复标题，避免跨年度同名标题链接漂移。
        # 3) 目录显示优先使用正文原文标题，尽量保留原文序号。
        heading_items = self._build_heading_items_from_paragraphs(
            toc_entries,
            strip_numbering=True,
            stable_order=True,
        )

        bookmark_names = self._ensure_heading_bookmarks(heading_items)

        toc_paragraph = self.document.paragraphs[0] if self.document.paragraphs else self.document.add_paragraph()
        title_para = toc_paragraph.insert_paragraph_before("目录")
        self._format_toc_paragraph(title_para)

        linked_count = 0
        for idx, entry in enumerate(toc_entries):
            item = heading_items[idx] if idx < len(heading_items) else {}
            display_entry = str(item.get("display_title") or entry).strip()
            anchor = bookmark_names[idx] if idx < len(bookmark_names) else ""
            if anchor:
                p = self._insert_hyperlink_paragraph_before(toc_paragraph, f"- {display_entry}", anchor)
                self._format_toc_paragraph(p)
                linked_count += 1
            else:
                p = toc_paragraph.insert_paragraph_before(f"- {display_entry}")
                self._format_toc_paragraph(p)

        return {
            "toc_entries": len(toc_entries),
            "semantic_heading_used": semantic_used,
            "llm_understanding_used": llm_understanding_used,
            "contextualized": any("[" in e and "月" in e for e in toc_entries),
            "linked_entries": linked_count,
            "document_title_included": bool(doc_title),
            "skipped": False,
        }

    def _apply_remove_blank_lines(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        removed = 0
        for paragraph in list(self.document.paragraphs):
            if paragraph.text.strip():
                continue
            p = paragraph._element
            p.getparent().remove(p)
            removed += 1
        return {"removed_blank_paragraphs": removed}

    def _apply_set_font_family(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        font_name = str(params.get("font_name", "宋体")).strip() or "宋体"
        updated_runs = 0
        paragraphs = self._select_target_paragraphs(target)
        for paragraph in paragraphs:
            self._ensure_paragraph_runs(paragraph)
            for run in paragraph.runs:
                run.font.name = font_name
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                except Exception:
                    pass
                updated_runs += 1
        return {"scope": str(target.get("scope", "all")), "font_name": font_name, "updated_runs": updated_runs}

    def _apply_set_font_color(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        color_hex = self._normalize_color_hex(str(params.get("color", "000000")))
        rgb = RGBColor.from_string(color_hex)
        updated_runs = 0
        paragraphs = self._select_target_paragraphs(target)
        for paragraph in paragraphs:
            self._ensure_paragraph_runs(paragraph)
            for run in paragraph.runs:
                run.font.color.rgb = rgb
                updated_runs += 1
        return {"scope": str(target.get("scope", "all")), "color": color_hex, "updated_runs": updated_runs}

    def _apply_set_font_size(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        try:
            size_pt = float(params.get("size_pt", 11.0))
        except Exception:
            size_pt = 11.0
        updated_runs = 0
        paragraphs = self._select_target_paragraphs(target)
        for paragraph in paragraphs:
            self._ensure_paragraph_runs(paragraph)
            for run in paragraph.runs:
                run.font.size = Pt(size_pt)
                updated_runs += 1
        return {"scope": str(target.get("scope", "all")), "size_pt": size_pt, "updated_runs": updated_runs}

    def _apply_set_paragraph_alignment(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        scope = str(target.get("scope", "all"))
        heading_level = self._safe_int(target.get("level"))
        alignment_name = str(params.get("alignment", "left")).lower()
        alignment_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        align = alignment_map.get(alignment_name, WD_PARAGRAPH_ALIGNMENT.LEFT)

        updated_paragraphs = 0
        for paragraph in self._iter_target_paragraphs(scope, heading_level=heading_level):
            paragraph.alignment = align
            updated_paragraphs += 1
        return {"scope": scope, "alignment": alignment_name, "updated_paragraphs": updated_paragraphs}

    def _apply_set_line_spacing(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        scope = str(target.get("scope", "all"))
        heading_level = self._safe_int(target.get("level"))
        try:
            spacing = float(params.get("line_spacing", 1.5))
        except Exception:
            spacing = 1.5

        updated_paragraphs = 0
        for paragraph in self._iter_target_paragraphs(scope, heading_level=heading_level):
            paragraph.paragraph_format.line_spacing = spacing
            updated_paragraphs += 1
        return {"scope": scope, "line_spacing": spacing, "updated_paragraphs": updated_paragraphs}

    def _apply_set_first_line_indent(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        scope = str(target.get("scope", "body"))
        heading_level = self._safe_int(target.get("level"))
        try:
            indent_pt = float(params.get("indent_pt", 24.0))
        except Exception:
            indent_pt = 24.0

        updated_paragraphs = 0
        for paragraph in self._iter_target_paragraphs(scope, heading_level=heading_level):
            paragraph.paragraph_format.first_line_indent = Pt(indent_pt)
            updated_paragraphs += 1
        return {"scope": scope, "indent_pt": indent_pt, "updated_paragraphs": updated_paragraphs}

    def _apply_set_highlight(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """应用文本高亮：黄色背景标记修改处"""
        find_text = str(params.get("find", "") or target.get("text", "") or target.get("target_text", ""))
        highlight_color = str(params.get("highlight_color", "yellow")).lower()
        if not find_text:
            return {"highlighted": 0, "color": highlight_color}

        # 将颜色名映射为python-docx支持的高亮颜色(WD_COLOR_INDEX)
        color_map = {
            "yellow": "yellow",
            "green": "green",
            "blue": "blue",
            "red": "red",
            "pink": "pink",
            "cyan": "cyan",
        }
        wd_highlight = color_map.get(highlight_color, "yellow")

        highlighted = 0
        for paragraph in self.document.paragraphs:
            highlighted += self._highlight_text_in_paragraph(paragraph, find_text, wd_highlight)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        highlighted += self._highlight_text_in_paragraph(paragraph, find_text, wd_highlight)

        return {"find": find_text, "highlighted": highlighted, "color": highlight_color}

    @staticmethod
    def _highlight_text_in_paragraph(paragraph, find_text: str, highlight_color: str) -> int:
        """在段落中仅高亮指定文本子串。"""
        text = paragraph.text or ""
        if find_text not in text:
            return 0

        from docx.enum.text import WD_COLOR_INDEX

        color_enum = getattr(WD_COLOR_INDEX, highlight_color.upper(), WD_COLOR_INDEX.YELLOW)
        changed = 0
        for run in list(paragraph.runs):
            run_text = run.text or ""
            if not run_text or find_text not in run_text:
                continue

            parts = run_text.split(find_text)
            run.text = parts[0]
            cursor = run._r
            template = copy.deepcopy(run._r)

            for idx in range(1, len(parts)):
                highlighted_elem = copy.deepcopy(template)
                highlighted_run = Run(highlighted_elem, paragraph)
                highlighted_run.text = find_text
                highlighted_run.font.highlight_color = color_enum
                cursor.addnext(highlighted_elem)
                cursor = highlighted_elem
                changed += 1

                tail = parts[idx]
                if tail:
                    tail_elem = copy.deepcopy(template)
                    tail_run = Run(tail_elem, paragraph)
                    tail_run.text = tail
                    cursor.addnext(tail_elem)
                    cursor = tail_elem

        return changed

    def _apply_insert_table(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """在文档末尾插入表格"""
        try:
            rows = int(params.get("rows", 3))
            cols = int(params.get("cols", 4))
        except Exception:
            rows, cols = 3, 4
        
        bold_header = bool(params.get("bold_header", True))

        if rows <= 0 or cols <= 0:
            return {"inserted": False, "reason": "行数或列数无效"}

        try:
            # 在文档末尾添加表格
            table = self.document.add_table(rows=rows, cols=cols)
            table.style = "Light Grid Accent 1"  # 应用基本表格样式

            # 如果需要加粗表头，对第一行进行处理
            if bold_header and rows > 0:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                        # 添加默认文本（表头）
                        if not any(r.text for r in paragraph.runs):
                            paragraph.add_run(f"列{table.rows[0].cells.index(cell) + 1}")

            return {
                "inserted": True,
                "rows": rows,
                "cols": cols,
                "bold_header": bold_header,
                "position": "end",
            }
        except Exception as e:
            return {"inserted": False, "reason": str(e)}

    def _apply_insert_footer_text(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """在页脚插入自定义文本"""
        footer_text = str(params.get("text", "Page {page_number}"))
        alignment = str(params.get("alignment", "center")).lower()
        include_page_number = bool(params.get("include_page_number", False))
        preserve_existing_page = bool(params.get("preserve_existing_page_number", False))

        try:
            # 获取所有section的footer（通常只有一个）
            for section in self.document.sections:
                footer = section.footer
                existing_page_para = None
                for p in footer.paragraphs:
                    if "PAGE" in p._p.xml:
                        existing_page_para = p
                        break

                if preserve_existing_page and existing_page_para is not None:
                    footer_para = existing_page_para
                else:
                    footer_para = footer.add_paragraph()

                # 设置对齐方式
                alignment_map = {
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                }
                if not (preserve_existing_page and existing_page_para is not None):
                    footer_para.alignment = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)
                
                template_text = footer_text.replace("{page_count}", "{total_pages}")
                if preserve_existing_page:
                    template_text = template_text.replace("{page_number}", "").replace("{total_pages}", "")
                    template_text = re.sub(r"^[，,\s]+|[，,\s]+$", "", template_text)
                    if template_text:
                        if footer_para.text and not footer_para.text.rstrip().endswith(("，", ",")):
                            footer_para.add_run("，")
                        footer_para.add_run(template_text)
                    continue

                if "{page_number}" in template_text or "{total_pages}" in template_text:
                    parts = re.split(r"(\{page_number\}|\{total_pages\})", template_text)
                    for part in parts:
                        if not part:
                            continue
                        if part == "{page_number}":
                            self._insert_page_number_field(footer_para)
                        elif part == "{total_pages}":
                            self._insert_total_pages_field(footer_para)
                        else:
                            footer_para.add_run(part)
                elif include_page_number:
                    footer_para.add_run(template_text)
                    self._insert_page_number_field(footer_para)
                else:
                    footer_para.add_run(template_text)

            return {"inserted": True, "text": footer_text, "alignment": alignment}
        except Exception as e:
            return {"inserted": False, "reason": str(e)}

    def _insert_page_number_field(self, paragraph) -> None:
        """在段落中插入页码域"""
        try:
            fld_simple = OxmlElement("w:fldSimple")
            fld_simple.set(qn("w:instr"), "PAGE")
            run = OxmlElement("w:r")
            fld_simple.append(run)
            paragraph._p.append(fld_simple)
        except Exception as e:
            # 如果xml操作失败，直接添加页码文本
            paragraph.add_run("{page}")

    def _insert_total_pages_field(self, paragraph) -> None:
        """在段落中插入总页数域"""
        try:
            fld_simple = OxmlElement("w:fldSimple")
            fld_simple.set(qn("w:instr"), "NUMPAGES")
            run = OxmlElement("w:r")
            fld_simple.append(run)
            paragraph._p.append(fld_simple)
        except Exception:
            paragraph.add_run("{total_pages}")

    def _apply_set_heading_numbering(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """为标题添加自动编号"""
        level = self._safe_int(target.get("level", 1))
        format_str = str(params.get("format", "第{index}章"))

        heading_indices = self._get_heading_indices_for_title_ops(level=level)
        updated = 0

        for idx, para_idx in enumerate(heading_indices, start=1):
            para_idx_int = int(para_idx) if isinstance(para_idx, (str, float)) else para_idx
            if para_idx_int < 0 or para_idx_int >= len(self.document.paragraphs):
                continue
            
            paragraph = self.document.paragraphs[para_idx_int]
            current_text = (paragraph.text or "").strip()
            
            # 生成编号
            number_str = format_str.replace("{index}", str(idx))
            
            # 检查是否已有编号
            if re.match(r"^(?:第\d+章|\.)", current_text):
                # 已有编号，跳过
                continue
            
            # 添加编号
            if paragraph.runs:
                first_run = paragraph.runs[0]
                first_run.text = number_str + current_text
                updated += 1
            else:
                paragraph.add_run(number_str + current_text)
                updated += 1

        return {"updated": updated, "level": level, "format": format_str}

    def _select_target_paragraphs(self, target: Dict[str, Any]) -> List[Any]:
        """
        根据 target 字典中的信息选择目标段落。
        支持的 target 格式：
        - {"scope": "document"} - 全文（默认）
        - {"scope": "selective", "target_text": "xxx"} - 包含特定文本的段落
        - {"scope": "paragraph", "paragraph_index": 0} - 特定索引的段落
        """
        target = dict(target or {})
        if str(target.get("type", "")).strip().lower() == "paragraph" and "paragraph_index" not in target:
            target["paragraph_index"] = target.get("index", -1)
            target.setdefault("paragraph_index_basis", "body")

        element = str(target.get("element", "")).strip().lower()
        if element in {"body_text", "body", "正文", "paragraph"} and not target.get("scope"):
            target["scope"] = "body"
        elif element in {"heading", "title", "headings"} and not target.get("scope"):
            target["scope"] = "heading"

        scope = str(target.get("scope", "document")).lower()
        target_text = target.get("target_text", "")
        section_title = str(target.get("section_title", "")).strip()
        paragraph_index = target.get("paragraph_index", -1)
        paragraph_index_basis = str(target.get("paragraph_index_basis", "absolute") or "absolute").lower()
        
        all_paragraphs = list(self.document.paragraphs)

        try:
            paragraph_index = int(paragraph_index)
        except Exception:
            paragraph_index = -1

        def _resolve_paragraph_by_index(index: int):
            if index < 0:
                return None
            if paragraph_index_basis == "body":
                body_paragraphs = self._get_content_paragraphs_for_indexing()
                return body_paragraphs[index] if index < len(body_paragraphs) else None
            return all_paragraphs[index] if index < len(all_paragraphs) else None

        if paragraph_index >= 0:
            if scope == "selective" and target_text:
                para = _resolve_paragraph_by_index(paragraph_index)
                if para is None:
                    return []
                return [para] if target_text in (para.text or "") else []

            para = _resolve_paragraph_by_index(paragraph_index)
            if para is not None:
                return [para]
            return []

        if scope == "section_content":
            if section_title:
                selected = self._iter_section_content_paragraphs(section_title)
                if selected:
                    return selected
            return self._get_content_paragraphs_for_indexing()
        
        # 情况 1：选择特定索引的段落
        if scope == "paragraph" and paragraph_index >= 0:
            para = _resolve_paragraph_by_index(paragraph_index)
            if para is not None:
                return [para]
            return []
        
        # 情况 2：选择包含特定文本的段落
        if scope == "selective" and target_text:
            # 指定了段落索引时，仅在该段内匹配，避免扩散到全文同名词。
            if paragraph_index >= 0:
                para = _resolve_paragraph_by_index(paragraph_index)
                if para is not None:
                    return [para] if target_text in (para.text or "") else []
                return []

            selected = []
            for para in all_paragraphs:
                if target_text in (para.text or ""):
                    selected.append(para)
            return selected
        
        # 默认情况：使用原有的 _iter_target_paragraphs 逻辑
        return self._iter_target_paragraphs(scope)

    def _get_content_paragraphs_for_indexing(self) -> List[Any]:
        """用于“第X段”定位的正文段落列表（排除标题/署名/日期/空行）。"""
        paragraphs = list(self.document.paragraphs)
        toc_block = self._get_toc_block_indices()
        content: List[Any] = []

        for idx, para in enumerate(paragraphs):
            if idx in toc_block:
                continue

            text = (para.text or "").strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name:
                continue

            if self._looks_like_structured_heading(text):
                continue

            # 排除主标题、署名机构、日期等元信息行。
            if len(text) <= 40 and ("统计局" in text or "调查队" in text):
                continue
            if re.fullmatch(r"\d{4}年\d{1,2}月", text):
                continue
            if len(text) <= 40 and text.endswith("统计公报"):
                continue

            content.append(para)

        return content

    def _iter_target_paragraphs(self, scope: str, heading_level: Optional[int] = None) -> List[Any]:
        scope_norm = (scope or "all").lower()
        paragraphs = list(self.document.paragraphs)

        if scope_norm in {"heading", "headings", "title"}:
            heading_indices = set(self._get_heading_indices_for_title_ops(level=heading_level))
            return [p for idx, p in enumerate(paragraphs) if idx in heading_indices]

        if scope_norm in {"body", "正文", "paragraph"}:
            return self._get_content_paragraphs_for_indexing()

        return paragraphs

    def _iter_section_content_paragraphs(self, section_title: str) -> List[Any]:
        title_norm = self._normalize_heading_text(section_title)
        if not title_norm:
            return []

        paragraphs = list(self.document.paragraphs)
        heading_indices = sorted(set(self._get_heading_indices_for_title_ops(level=None)))
        toc_block = self._get_toc_block_indices()

        matched_heading_indices: List[int] = []
        for idx in heading_indices:
            if idx in toc_block or idx < 0 or idx >= len(paragraphs):
                continue
            heading_text = self._normalize_heading_text((paragraphs[idx].text or ""))
            if title_norm and title_norm in heading_text:
                matched_heading_indices.append(idx)

        if not matched_heading_indices:
            return []

        heading_set = set(heading_indices)
        selected: List[Any] = []
        for start_idx in matched_heading_indices:
            end_idx = len(paragraphs)
            for idx in heading_indices:
                if idx > start_idx:
                    end_idx = idx
                    break

            for idx in range(start_idx + 1, end_idx):
                if idx in toc_block or idx in heading_set:
                    continue
                paragraph = paragraphs[idx]
                if not (paragraph.text or "").strip():
                    continue
                selected.append(paragraph)

        return selected

    @staticmethod
    def _normalize_color_hex(value: str) -> str:
        s = (value or "").strip().lstrip("#")
        name_map = {
            "red": "FF0000",
            "blue": "0000FF",
            "black": "000000",
            "green": "008000",
            "gray": "808080",
            "grey": "808080",
            "orange": "FFA500",
            "purple": "800080",
            "violet": "8A2BE2",
            "紫": "800080",
            "紫色": "800080",
            "yellow": "FFFF00",
        }
        if s.lower() in name_map:
            return name_map[s.lower()]
        if re.fullmatch(r"[0-9a-fA-F]{6}", s):
            return s.upper()
        return "000000"

    def _get_semantic_heading_indices(self, level: int = 1) -> List[int]:
        # 缓存按级别隔离，避免不同级别互相污染。
        if level in self._semantic_heading_cache:
            return self._semantic_heading_cache[level]

        understanding = self._get_document_understanding()
        if isinstance(understanding.get("headings"), list) and understanding.get("headings"):
            idxs = []
            for item in understanding.get("headings", []):
                if not isinstance(item, dict):
                    continue
                try:
                    lv = int(item.get("level", 1))
                except Exception:
                    lv = 1
                if lv != level:
                    continue
                try:
                    idx = int(item.get("index"))
                except Exception:
                    continue
                if 0 <= idx < len(self.document.paragraphs):
                    if not self._is_semantic_heading_candidate(idx):
                        continue
                    idxs.append(idx)
            if idxs:
                self._semantic_heading_cache[level] = sorted(set(idxs))
                return self._semantic_heading_cache[level]

        paragraphs = [p.text.strip() for p in self.document.paragraphs]
        non_empty_items = [(idx, txt) for idx, txt in enumerate(paragraphs) if txt]
        if not non_empty_items:
            self._semantic_heading_cache[level] = []
            return []

        llm = self._llm_service or get_llm_service()
        if llm and hasattr(llm, "is_available") and llm.is_available():
            try:
                items_text = "\n".join([f"{i}: {txt}" for i, txt in non_empty_items[:300]])
                system_prompt = (
                    "你是文档结构分析器。"
                    "请识别最像标题的段落索引。"
                    "只输出 JSON，不要解释。"
                )
                user_prompt = (
                    f"目标标题级别: {level}\n"
                    "给定段落列表（格式: 原始段落索引: 文本）:\n"
                    f"{items_text}\n"
                    "输出格式: {\"indices\": [0, 5, 12]}"
                )

                original_streaming = None
                can_toggle_streaming = hasattr(llm, "config") and hasattr(llm.config, "streaming")
                if can_toggle_streaming:
                    original_streaming = llm.config.streaming
                    llm.config.streaming = False
                try:
                    raw = llm.chat(
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                        temperature=0,
                    )
                finally:
                    if can_toggle_streaming:
                        llm.config.streaming = original_streaming

                parsed = self._safe_load_json(raw)
                if isinstance(parsed, dict) and isinstance(parsed.get("indices"), list):
                    indices = []
                    for v in parsed["indices"]:
                        try:
                            iv = int(v)
                        except Exception:
                            continue
                        if 0 <= iv < len(paragraphs):
                            if not self._is_semantic_heading_candidate(iv):
                                continue
                            indices.append(iv)
                    if indices:
                        self._semantic_heading_cache[level] = sorted(set(indices))
                        return self._semantic_heading_cache[level]
            except Exception:
                pass

        # LLM 不可用或失败时，使用简单启发式标题识别兜底。
        heuristic = []
        for idx, txt in non_empty_items:
            if len(txt) <= 30 and re.match(r"^(第[一二三四五六七八九十\d]+[章节部分]|[一二三四五六七八九十]+、|\d+\.)", txt):
                if self._is_semantic_heading_candidate(idx):
                    heuristic.append(idx)
                continue

            paragraph = self.document.paragraphs[idx]
            if self._looks_like_non_numbered_heading(paragraph) and self._is_semantic_heading_candidate(idx):
                heuristic.append(idx)
        self._semantic_heading_cache[level] = heuristic
        return heuristic

    def _get_style_heading_indices(self, level: Optional[int] = None) -> List[int]:
        idxs: List[int] = []
        toc_block = self._get_toc_block_indices()
        for idx, paragraph in enumerate(self.document.paragraphs):
            if idx in toc_block:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" not in style_name:
                continue
            if level is not None and level > 0 and f"heading {level}" not in style_name:
                continue
            idxs.append(idx)
        return idxs

    def _map_toc_entries_to_indices(self, entries: List[str], strip_numbering: bool = False) -> List[int]:
        if not entries:
            return []

        indices: List[int] = []
        used: set[int] = set()
        toc_block = self._get_toc_block_indices()
        paragraphs = [
            (i, (p.text or "").strip())
            for i, p in enumerate(self.document.paragraphs)
            if i not in toc_block
        ]

        for entry in entries:
            title = str(entry or "").strip()
            if not title or self._is_caption_like_text(title):
                continue
            norm_title = self._normalize_heading_text(title)
            norm_title_wo_num = self._normalize_heading_text(title, strip_numbering=True)
            for idx, txt in paragraphs:
                if idx in used:
                    continue
                norm_txt = self._normalize_heading_text(txt)
                if norm_txt != norm_title:
                    if not strip_numbering:
                        continue
                    norm_txt_wo_num = self._normalize_heading_text(txt, strip_numbering=True)
                    if norm_txt_wo_num != norm_title_wo_num:
                        continue
                if not self._is_semantic_heading_candidate(idx):
                    continue
                indices.append(idx)
                used.add(idx)
                break

        return indices

    def _get_document_main_title(self, understanding: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        if self._document_title_cache is not None:
            title = str(self._document_title_cache.get("title", "")).strip()
            try:
                idx = int(self._document_title_cache.get("index", -1))
            except Exception:
                idx = -1

            if self._is_document_title_candidate(title, idx):
                return self._document_title_cache

            # 文本被替换后，优先尝试复用同一索引处的最新标题文本。
            if 0 <= idx < len(self.document.paragraphs) and idx not in self._get_toc_block_indices():
                live_text = (self.document.paragraphs[idx].text or "").strip()
                if self._is_document_title_candidate(live_text, idx):
                    self._document_title_cache = {"index": idx, "title": live_text}
                    return self._document_title_cache

            # 目录插入后段落索引可能变化，按标题文本重新定位。
            resolved_idx = self._resolve_title_index_by_text(title)
            if resolved_idx >= 0:
                self._document_title_cache = {"index": resolved_idx, "title": title}
                return self._document_title_cache

        title_obj = understanding.get("document_title") if isinstance(understanding, dict) else None
        if isinstance(title_obj, dict):
            try:
                idx = int(title_obj.get("index", -1))
            except Exception:
                idx = -1
            title = str(title_obj.get("title", "")).strip()
            if self._is_document_title_candidate(title, idx):
                self._document_title_cache = {"index": idx, "title": title}
                return self._document_title_cache

            # 若 understanding 的标题文本过期，但索引仍指向总标题，则更新为实时文本。
            if 0 <= idx < len(self.document.paragraphs) and idx not in self._get_toc_block_indices():
                live_text = (self.document.paragraphs[idx].text or "").strip()
                if self._is_document_title_candidate(live_text, idx):
                    self._document_title_cache = {"index": idx, "title": live_text}
                    return self._document_title_cache

        paragraphs = [(i, (p.text or "").strip()) for i, p in enumerate(self.document.paragraphs) if (p.text or "").strip()]
        if not paragraphs:
            self._document_title_cache = None
            return None

        llm = self._llm_service or get_llm_service()
        if llm and hasattr(llm, "is_available") and llm.is_available():
            try:
                context = "\n".join([f"{i}: {t}" for i, t in paragraphs[:40]])
                system_prompt = (
                    "你是文档标题识别器。"
                    "请从给定段落中识别整篇文档总标题（不是章节标题）。"
                    "仅输出 JSON。"
                )
                user_prompt = (
                    "段落列表(索引:文本):\n"
                    f"{context}\n\n"
                    "输出格式: {\"title\": \"...\", \"index\": 0, \"confidence\": 0.95}。"
                    "如果不存在总标题，输出 {\"title\": \"\", \"index\": -1, \"confidence\": 0}."
                )

                original_streaming = None
                can_toggle_streaming = hasattr(llm, "config") and hasattr(llm.config, "streaming")
                if can_toggle_streaming:
                    original_streaming = llm.config.streaming
                    llm.config.streaming = False
                try:
                    raw = llm.chat(
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                        temperature=0,
                    )
                finally:
                    if can_toggle_streaming:
                        llm.config.streaming = original_streaming

                parsed = self._safe_load_json(raw)
                if isinstance(parsed, dict):
                    title = str(parsed.get("title", "")).strip()
                    try:
                        idx = int(parsed.get("index", -1))
                    except Exception:
                        idx = -1
                    if self._is_document_title_candidate(title, idx):
                        self._document_title_cache = {"index": idx, "title": title}
                        return self._document_title_cache
            except Exception:
                pass

        # 保守兜底：在首段中匹配“公报/报告”等常见文档总标题关键词。
        for idx, txt in paragraphs[:8]:
            if self._is_document_title_candidate(txt, idx):
                self._document_title_cache = {"index": idx, "title": txt}
                return self._document_title_cache

        self._document_title_cache = None
        return None

    def _resolve_title_index_by_text(self, title: str) -> int:
        norm_title = self._normalize_heading_text(title)
        if not norm_title:
            return -1

        toc_block = self._get_toc_block_indices()
        for idx, p in enumerate(self.document.paragraphs):
            if idx in toc_block:
                continue
            txt = (p.text or "").strip()
            if self._normalize_heading_text(txt) == norm_title:
                return idx
        return -1

    def _is_document_title_candidate(self, title: str, idx: int) -> bool:
        t = (title or "").strip()
        if not t:
            return False
        if idx < 0 or idx >= len(self.document.paragraphs):
            return False
        if idx in self._get_toc_block_indices():
            return False
        live_text = (self.document.paragraphs[idx].text or "").strip()
        if self._normalize_heading_text(live_text) != self._normalize_heading_text(t):
            return False
        if len(t) < 8 or len(t) > 80:
            return False
        if self._is_caption_like_text(t):
            return False
        # 总标题通常包含“公报/报告/白皮书/规划”等关键词。
        if not re.search(r"公报|报告|白皮书|规划|综述", t):
            return False
        return True

    def _get_heading_indices_for_title_ops(self, level: Optional[int] = None) -> List[int]:
        # 与目录提取一致：每次实时重算，不复用历史索引缓存。
        # 规则：样式优先；无样式时优先使用 toc_entries 反查；最后回退 headings/语义。
        style_indices = self._get_style_heading_indices(level=level)
        if style_indices:
            return sorted(set(style_indices))

        # 文档已含目录时，以文档中的目录条目为唯一事实来源，避免使用插入目录前的旧索引。
        if level is None and self._has_existing_toc():
            live_toc_entries = self._get_toc_entries_from_document()
            live_indices = self._map_toc_entries_to_indices(live_toc_entries, strip_numbering=True)
            doc_title = self._get_document_main_title(understanding={})
            if isinstance(doc_title, dict):
                try:
                    doc_idx = int(doc_title.get("index", -1))
                except Exception:
                    doc_idx = -1
                if 0 <= doc_idx < len(self.document.paragraphs) and doc_idx not in self._get_toc_block_indices():
                    live_indices = sorted(set(list(live_indices) + [doc_idx]))
            if live_indices:
                return sorted(set(live_indices))
            return []

        understanding = self._get_document_understanding()

        if level is None:
            toc_entries = understanding.get("toc_entries", []) if isinstance(understanding, dict) else []
            if isinstance(toc_entries, list) and toc_entries:
                toc_indices = self._map_toc_entries_to_indices(
                    [str(e).strip() for e in toc_entries if str(e).strip()],
                    strip_numbering=True,
                )
                if toc_indices:
                    return sorted(set(toc_indices))

        heading_indices: List[int] = []
        headings = understanding.get("headings", []) if isinstance(understanding, dict) else []
        if isinstance(headings, list) and headings:
            for item in headings:
                if not isinstance(item, dict):
                    continue
                try:
                    idx = int(item.get("index"))
                except Exception:
                    continue
                if idx < 0 or idx >= len(self.document.paragraphs):
                    continue
                if idx in self._get_toc_block_indices():
                    continue
                if level is not None:
                    try:
                        lv = int(item.get("level", 1))
                    except Exception:
                        lv = 1
                    if lv != level:
                        continue
                if not self._is_semantic_heading_candidate(idx):
                    continue
                heading_indices.append(idx)

        if heading_indices:
            return sorted(set(heading_indices))

        if level is not None:
            return sorted(set(self._get_semantic_heading_indices(level=level)))

        result = sorted(set(self._get_semantic_heading_indices(level=1)))
        # 对标题样式范围（level=None）补入总标题，避免“总标题不变色/不加粗”。
        doc_title = self._get_document_main_title(understanding)
        if isinstance(doc_title, dict):
            try:
                doc_idx = int(doc_title.get("index", -1))
            except Exception:
                doc_idx = -1
            if 0 <= doc_idx < len(self.document.paragraphs) and doc_idx not in self._get_toc_block_indices():
                result = sorted(set(result + [doc_idx]))
        return result

    def _get_toc_entries_from_document(self) -> List[str]:
        entries: List[str] = []
        in_toc = False
        for paragraph in self.document.paragraphs:
            text = (paragraph.text or "").strip()
            xml = paragraph._p.xml
            if text == "目录":
                in_toc = True
                continue
            if not in_toc:
                continue

            is_toc_line = text.startswith("-") or ("w:hyperlink" in xml and "toc_" in xml)
            if is_toc_line:
                if text.startswith("-"):
                    entry = text[1:].strip()
                    if entry:
                        entries.append(entry)
                continue

            if text:
                break

        return entries

    def _get_document_understanding(self) -> Dict[str, Any]:
        """
        Let LLM understand the whole document structure once, then reuse.

        Expected JSON shape:
        {
                    "headings": [{"index": 0, "title": "第一章 总则", "level": 1}],
                    "toc_entries": ["第一章 总则"],
                    "field_values": {"金额": "1000万元"}
        }
        """
        if self._document_understanding_cache is not None:
            return self._document_understanding_cache

        paragraphs = [p.text.strip() for p in self.document.paragraphs]
        non_empty = [(idx, txt) for idx, txt in enumerate(paragraphs) if txt]
        tables_preview: List[str] = []
        for t_idx, table in enumerate(self.document.tables[:20]):
            row_text = []
            for row in table.rows[:8]:
                row_text.append(" | ".join(cell.text.strip() for cell in row.cells[:8]))
            if row_text:
                tables_preview.append(f"table#{t_idx}: " + " || ".join(row_text))

        if not non_empty and not tables_preview:
            self._document_understanding_cache = {"headings": [], "toc_entries": [], "field_values": {}}
            return self._document_understanding_cache

        llm = self._llm_service or get_llm_service()
        if llm and hasattr(llm, "is_available") and llm.is_available():
            try:
                para_text = "\n".join([f"{i}: {txt}" for i, txt in non_empty[:400]])
                table_text = "\n".join(tables_preview[:60])
                system_prompt = (
                    "你是文档语义理解器。"
                    "请基于段落和表格，识别章节标题、目录展示项与关键字段值。"
                    "要特别识别没有一、二、三或1,2,3等编号，但具有标题性质的短句。"
                    "目录展示项不需要月份/时间上下文前缀，只输出标题本身。"
                    "图表说明（如图1、表2、Figure、Table等）不要放入目录。"
                    "只输出 JSON。"
                )
                user_prompt = (
                    "文档段落(格式: 段落索引: 文本):\n"
                    f"{para_text}\n\n"
                    "文档表格摘要:\n"
                    f"{table_text}\n\n"
                    "识别规则补充:\n"
                    "1) 标题可能没有编号；\n"
                    "2) 典型特征：短句、概括性强、后续常跟解释段落；\n"
                    "3) 不要因为缺少编号就漏掉标题；\n"
                    "4) 图表说明不要作为目录项。\n\n"
                    "输出格式:\n"
                    "{\"headings\": [{\"index\": 0, \"title\": \"第一章 总则\", \"level\": 1}], "
                    "\"toc_entries\": [\"第一章 总则\"], "
                    "\"field_values\": {\"金额\": \"1000万元\", \"地区生产总值\": \"18500亿元\"}}"
                )

                original_streaming = None
                can_toggle_streaming = hasattr(llm, "config") and hasattr(llm.config, "streaming")
                if can_toggle_streaming:
                    original_streaming = llm.config.streaming
                    llm.config.streaming = False
                try:
                    raw = llm.chat(
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                        temperature=0,
                    )
                finally:
                    if can_toggle_streaming:
                        llm.config.streaming = original_streaming

                parsed = self._safe_load_json(raw)
                if isinstance(parsed, dict):
                    headings = parsed.get("headings", [])
                    toc_entries = parsed.get("toc_entries", [])
                    field_values = parsed.get("field_values", {})
                    document_title = parsed.get("document_title", {})
                    if not isinstance(headings, list):
                        headings = []
                    if not isinstance(toc_entries, list):
                        toc_entries = []
                    if not isinstance(field_values, dict):
                        field_values = {}
                    if not isinstance(document_title, dict):
                        document_title = {}

                    if not toc_entries:
                        toc_entries = self._build_contextual_toc_entries(
                            {"headings": headings},
                            [str(h.get("title", "")).strip() for h in headings if isinstance(h, dict)],
                        )

                    self._document_understanding_cache = {
                        "headings": headings,
                        "toc_entries": toc_entries,
                        "field_values": field_values,
                        "document_title": document_title,
                    }
                    return self._document_understanding_cache
            except Exception:
                pass

        self._document_understanding_cache = {"headings": [], "toc_entries": [], "field_values": {}}
        return self._document_understanding_cache

    def _build_contextual_toc_entries(self, understanding: Dict[str, Any], headings: List[str]) -> List[str]:
        entries: List[str] = []

        u_heads = understanding.get("headings", []) if isinstance(understanding, dict) else []
        if isinstance(u_heads, list) and u_heads:
            for item in u_heads:
                if not isinstance(item, dict):
                    continue
                title = str(item.get("title", "")).strip()
                if not title:
                    continue
                entries.append(title)

            if entries:
                return entries

        for title in headings:
            if title:
                entries.append(title)
        return entries

    def _find_nearest_month_label(self, index: int) -> str:
        if index < 0 or index >= len(self.document.paragraphs):
            return ""

        window = 8
        lo = max(0, index - window)
        hi = min(len(self.document.paragraphs), index + window + 1)

        for i in range(index, lo - 1, -1):
            text = self.document.paragraphs[i].text.strip()
            label = self._extract_month_label(text)
            if label:
                return label
        for i in range(index + 1, hi):
            text = self.document.paragraphs[i].text.strip()
            label = self._extract_month_label(text)
            if label:
                return label
        return ""

    @staticmethod
    def _extract_month_label(text: str) -> str:
        if not text:
            return ""

        m = re.search(r"((?:20\d{2}|19\d{2})年\s*(?:1[0-2]|0?[1-9])月)", text)
        if m:
            return re.sub(r"\s+", "", m.group(1))

        m = re.search(r"((?:1[0-2]|0?[1-9])月)", text)
        if m:
            return m.group(1)

        return ""

    @staticmethod
    def _looks_like_non_numbered_heading(paragraph) -> bool:
        text = (paragraph.text or "").strip()
        if not text:
            return False

        # 过长内容一般是正文，不是标题。
        if len(text) > 28:
            return False

        # 正文句末与密集分隔符降低标题概率。
        if re.search(r"[。！？；;，,]", text):
            return False

        title_keywords = [
            "概况",
            "情况",
            "指标",
            "目标",
            "问题",
            "措施",
            "建议",
            "分析",
            "总结",
            "说明",
            "综述",
            "附录",
            "附件",
            "工作",
            "进展",
            "完成",
        ]

        score = 0
        if len(text) <= 14:
            score += 1
        if any(k in text for k in title_keywords):
            score += 2
        if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            score += 1
        if paragraph.runs and any(bool(run.bold) for run in paragraph.runs):
            score += 1

        # 兜底：短且无句末标点时，给出基础分，防止漏掉简短语义标题。
        if len(text) <= 10:
            score += 1

        return score >= 2

    def _is_semantic_heading_candidate(self, idx: int) -> bool:
        if idx < 0 or idx >= len(self.document.paragraphs):
            return False

        paragraph = self.document.paragraphs[idx]
        text = (paragraph.text or "").strip()
        if not text:
            return False

        heading_like_prefix = bool(
            re.match(
                r"^(?:\d{4}年\d{1,2}月|第[一二三四五六七八九十\d]+[章节部分]|[一二三四五六七八九十]+、|\d+[\.、]|注释[:：]?|资料来源[:：]?|附录[:：]?)",
                text,
            )
        )

        if ("：" in text or ":" in text) and len(text) >= 8:
            return False
        if re.search(r"[。！？；;]$", text):
            return False
        if re.search(r"[，,。；;]", text) and not heading_like_prefix:
            return False
        if self._is_caption_like_text(text):
            return False
        if len(text) > 42:
            return False

        return True

    @staticmethod
    def _safe_int(value: Any) -> Optional[int]:
        try:
            iv = int(value)
        except Exception:
            return None
        return iv if iv > 0 else None

    @staticmethod
    def _ensure_paragraph_runs(paragraph) -> None:
        if paragraph.runs:
            return
        text = (paragraph.text or "")
        if not text:
            return
        paragraph.text = ""
        paragraph.add_run(text)

    def _has_existing_toc(self) -> bool:
        for paragraph in self.document.paragraphs:
            text = (paragraph.text or "").strip()
            if text == "目录":
                return True
            xml = paragraph._p.xml
            if "TOC" in xml or "w:hyperlink" in xml and "toc_" in xml:
                return True
        return False

    def _build_heading_items_from_paragraphs(
        self,
        headings: List[str],
        strip_numbering: bool = False,
        stable_order: bool = True,
    ) -> List[Dict[str, Any]]:
        items: List[Dict[str, Any]] = []
        used: set[int] = set()
        toc_block = self._get_toc_block_indices()

        candidates: List[Dict[str, Any]] = []
        for idx, paragraph in enumerate(self.document.paragraphs):
            if idx in toc_block:
                continue
            raw_text = (paragraph.text or "").strip()
            if not raw_text:
                continue
            candidates.append(
                {
                    "idx": idx,
                    "raw": raw_text,
                    "norm": self._normalize_heading_text(raw_text),
                    "norm_wo_num": self._normalize_heading_text(raw_text, strip_numbering=True),
                }
            )

        cursor = 0
        for title in headings:
            title_text = str(title or "").strip()
            norm_title = self._normalize_heading_text(title_text)
            norm_title_wo_num = self._normalize_heading_text(title_text, strip_numbering=True)
            if not norm_title:
                items.append(
                    {
                        "index": -1,
                        "title": title_text,
                        "display_title": title_text,
                        "level": 1,
                        "matched": False,
                    }
                )
                continue

            def _is_match(candidate: Dict[str, Any]) -> bool:
                if candidate["idx"] in used:
                    return False
                if candidate["norm"] == norm_title:
                    return True
                if strip_numbering and candidate["norm_wo_num"] == norm_title_wo_num:
                    return True
                return False

            selected = None
            selected_pos = -1

            if stable_order:
                for pos in range(cursor, len(candidates)):
                    c = candidates[pos]
                    if _is_match(c):
                        selected = c
                        selected_pos = pos
                        break

            if selected is None:
                for pos, c in enumerate(candidates):
                    if _is_match(c):
                        selected = c
                        selected_pos = pos
                        break

            if selected is None:
                items.append(
                    {
                        "index": -1,
                        "title": title_text,
                        "display_title": title_text,
                        "level": 1,
                        "matched": False,
                    }
                )
                continue

            used.add(int(selected["idx"]))
            if stable_order and selected_pos >= 0:
                cursor = selected_pos + 1

            items.append(
                {
                    "index": int(selected["idx"]),
                    "title": title_text,
                    "display_title": str(selected["raw"]),
                    "level": 1,
                    "matched": True,
                }
            )

        return items

    def _get_toc_block_indices(self) -> set[int]:
        # 识别目录块：从“目录”开始，后续连续的目录行（- 开头或含内部链接）都排除。
        indices: set[int] = set()
        in_toc = False
        for idx, paragraph in enumerate(self.document.paragraphs):
            text = (paragraph.text or "").strip()
            xml = paragraph._p.xml
            if text == "目录":
                indices.add(idx)
                in_toc = True
                continue

            if not in_toc:
                continue

            is_toc_line = text.startswith("-") or ("w:hyperlink" in xml and "toc_" in xml)
            if is_toc_line:
                indices.add(idx)
                continue

            if text:
                break

            # 目录条目之间的空行也作为目录块处理。
            indices.add(idx)

        return indices

    def _ensure_heading_bookmarks(self, heading_items: List[Dict[str, Any]]) -> List[str]:
        names: List[str] = []
        bookmark_id = 1000
        for i, item in enumerate(heading_items):
            try:
                idx = int(item.get("index"))
            except Exception:
                idx = -1
            if idx < 0 or idx >= len(self.document.paragraphs):
                names.append("")
                continue

            paragraph = self.document.paragraphs[idx]
            name = f"toc_{i + 1}_{idx}"
            if self._paragraph_has_bookmark(paragraph, name):
                names.append(name)
                continue
            self._add_bookmark_to_paragraph(paragraph, name=name, bookmark_id=bookmark_id + i)
            names.append(name)
        return names

    @staticmethod
    def _paragraph_has_bookmark(paragraph, name: str) -> bool:
        return f'w:name="{name}"' in paragraph._p.xml

    @staticmethod
    def _add_bookmark_to_paragraph(paragraph, name: str, bookmark_id: int) -> None:
        p = paragraph._p
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), name)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))

        p.insert(0, start)
        p.append(end)

    def _insert_hyperlink_paragraph_before(self, before_paragraph, text: str, anchor: str):
        paragraph = before_paragraph.insert_paragraph_before("")

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)
        hyperlink.set(qn("w:history"), "1")

        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        # 保持可点击，同时固定为黑色，避免默认 Hyperlink 样式显示为蓝色。
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rpr.append(color)
        run.append(rpr)

        t = OxmlElement("w:t")
        t.text = text
        run.append(t)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        return paragraph

    @staticmethod
    def _format_toc_paragraph(paragraph) -> None:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    def _normalize_existing_toc_block(self) -> int:
        normalized = 0
        in_toc = False
        for paragraph in self.document.paragraphs:
            text = (paragraph.text or "").strip()
            xml = paragraph._p.xml

            if text == "目录":
                in_toc = True
                self._format_toc_paragraph(paragraph)
                normalized += 1
                continue

            if not in_toc:
                continue

            is_toc_line = text.startswith("-") or ("w:hyperlink" in xml and "toc_" in xml)
            if is_toc_line:
                self._format_toc_paragraph(paragraph)
                normalized += 1
                continue

            if text:
                break

        return normalized

    @staticmethod
    def _is_caption_like_text(text: str) -> bool:
        s = (text or "").strip()
        if not s:
            return False

        patterns = [
            r"^图\s*\d+",
            r"^表\s*\d+",
            r"^图\s*[一二三四五六七八九十\d]+",
            r"^表\s*[一二三四五六七八九十\d]+",
            r"^figure\s*\d+",
            r"^table\s*\d+",
            r"^fig\.\s*\d+",
            r"^来源[:：]",
            r"^作者[:：]",
            r"^编制[:：]",
        ]
        lowered = s.lower()
        if any(re.match(p, s, flags=re.IGNORECASE) for p in patterns) or lowered.startswith("figure ") or lowered.startswith("table "):
            return True

        # 排除署名/来源机构行，例如“南京市统计局、国家统计局南京调查队”。
        institution_hits = sum(1 for k in ["统计局", "调查队", "研究院", "委员会", "办公室"] if k in s)
        if institution_hits >= 2 and len(s) <= 40:
            return True
        if ("统计局" in s or "调查队" in s) and ("、" in s or "," in s or "，" in s) and len(s) <= 40:
            return True
        if len(s) <= 20 and (s.endswith("统计局") or s.endswith("调查队") or s.endswith("统计局南京调查队")):
            return True

        return False

    @staticmethod
    def _extract_fields_heuristically(body: List[str], fields: List[Any]) -> Dict[str, str]:
        text = "\n".join(body)
        result: Dict[str, str] = {}
        for f in fields:
            key = str(f).strip()
            if not key:
                continue
            pattern = rf"{re.escape(key)}\s*[:：]?\s*([^\n，。,;；]+)"
            m = re.search(pattern, text)
            if m:
                result[key] = m.group(1).strip()
        return result

    @staticmethod
    def _safe_load_json(text: str):
        raw = (text or "").strip()
        if not raw:
            return None
        try:
            return json.loads(raw)
        except Exception:
            pass
        match = re.search(r"\{[\s\S]*\}", raw)
        if not match:
            return None
        try:
            return json.loads(match.group(0))
        except Exception:
            return None

    @staticmethod
    def _normalize_heading_text(text: str, strip_numbering: bool = False) -> str:
        s = (text or "").strip()
        if strip_numbering:
            s = DocxAdapter._strip_heading_numbering_prefix(s)
        s = re.sub(r"\s+", "", s)
        s = s.replace("（", "(").replace("）", ")")
        s = s.strip("：:")
        return s

    @staticmethod
    def _strip_heading_numbering_prefix(text: str) -> str:
        s = (text or "").strip()
        if not s:
            return s

        patterns = [
            r"^第[一二三四五六七八九十百零〇\d]+[章节部分篇]\s*",
            r"^[（(]?[一二三四五六七八九十百零〇]+[)）][、.．]?\s*",
            r"^[一二三四五六七八九十百零〇]+[、.．]\s*",
            r"^\d+(?:\.\d+){0,3}[、.．]\s*",
        ]

        changed = True
        while changed:
            changed = False
            for p in patterns:
                new_s = re.sub(p, "", s)
                if new_s != s:
                    s = new_s.strip()
                    changed = True
        return s


    def _apply_set_italic(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """应用文本斜体格式"""
        paragraphs = self._select_target_paragraphs(target)
        target_text = str(target.get("target_text", "") or "").strip()
        updated_runs = 0
        for paragraph in paragraphs:
            self._ensure_paragraph_runs(paragraph)
            if target_text:
                updated_runs += self._apply_inline_text_style(paragraph, target_text, "italic")
                continue
            for run in paragraph.runs:
                run.italic = True
                updated_runs += 1
        return {"italic": True, "updated_runs": updated_runs}
    
    def _apply_set_underline(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """应用文本下划线格式"""
        paragraphs = self._select_target_paragraphs(target)
        target_text = str(target.get("target_text", "") or "").strip()
        updated_runs = 0
        for paragraph in paragraphs:
            self._ensure_paragraph_runs(paragraph)
            if target_text:
                updated_runs += self._apply_inline_text_style(paragraph, target_text, "underline")
                continue
            for run in paragraph.runs:
                run.underline = True
                updated_runs += 1
        return {"underline": True, "updated_runs": updated_runs}

    def _apply_inline_text_style(self, paragraph: Any, target_text: str, style_attr: str) -> int:
        """仅对 run 内命中的子串应用样式，避免整段误改。"""
        if not target_text:
            return 0

        updated = 0
        for run in list(paragraph.runs):
            run_text = run.text or ""
            if not run_text or target_text not in run_text:
                continue

            parts = run_text.split(target_text)
            run.text = parts[0]
            cursor = run._r
            template = copy.deepcopy(run._r)

            for idx in range(1, len(parts)):
                styled_elem = copy.deepcopy(template)
                styled_run = Run(styled_elem, paragraph)
                styled_run.text = target_text
                setattr(styled_run, style_attr, True)
                cursor.addnext(styled_elem)
                cursor = styled_elem
                updated += 1

                tail = parts[idx]
                if tail:
                    tail_elem = copy.deepcopy(template)
                    tail_run = Run(tail_elem, paragraph)
                    tail_run.text = tail
                    cursor.addnext(tail_elem)
                    cursor = tail_elem

        return updated
    
    def _apply_set_paragraph_spacing(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """设置段落间距"""
        paragraphs = self._select_target_paragraphs(target)
        before_spacing = float(params.get("before_spacing", 0.0))
        after_spacing = float(params.get("after_spacing", 6.0))
        updated = 0
        for paragraph in paragraphs:
            paragraph.paragraph_format.space_before = Pt(before_spacing)
            paragraph.paragraph_format.space_after = Pt(after_spacing)
            updated += 1
        return {"before_spacing": before_spacing, "after_spacing": after_spacing, "updated_paragraphs": updated}
    
    def _apply_set_bullet_list(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """应用项目符号列表"""
        paragraphs = self._select_target_paragraphs(target)
        updated = 0
        for paragraph in paragraphs:
            try:
                # 尝试使用样式 "List Bullet"
                paragraph.style = "List Bullet"
                updated += 1
            except Exception:
                # 如果样式不存在，使用 XML 直接添加列表格式
                try:
                    pPr = paragraph._element.get_or_add_pPr()
                    numPr = OxmlElement('w:numPr')
                    numPr.attrib[qn('w:ilvl')] = '0'
                    numPr.attrib[qn('w:numId')] = '1'
                    pPr.append(numPr)
                    updated += 1
                except Exception:
                    pass
        return {"bullet_type": "bullet", "updated_paragraphs": updated}
    
    def _apply_set_numbered_list(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """应用编号列表"""
        paragraphs = self._select_target_paragraphs(target)
        updated = 0
        for paragraph in paragraphs:
            try:
                # 尝试使用样式 "List Number"
                paragraph.style = "List Number"
                updated += 1
            except Exception:
                # 如果样式不存在，使用 XML 直接添加编号列表格式
                try:
                    pPr = paragraph._element.get_or_add_pPr()
                    numPr = OxmlElement('w:numPr')
                    numPr.attrib[qn('w:ilvl')] = '0'
                    numPr.attrib[qn('w:numId')] = '2'
                    pPr.append(numPr)
                    updated += 1
                except Exception:
                    pass
        return {"number_format": "decimal", "updated_paragraphs": updated}
    
    def _apply_set_paragraph_shading(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """应用段落底纹"""
        paragraphs = self._select_target_paragraphs(target)
        shading_color = str(params.get("shading_color", "FFFF00")).upper()
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        updated = 0
        for paragraph in paragraphs:
            try:
                pPr = paragraph._element.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
                pPr.append(shd)
                updated += 1
            except Exception:
                pass
        return {"shading_color": shading_color, "updated_paragraphs": updated}
    
    def _apply_set_paragraph_border(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """应用段落边框"""
        paragraphs = self._select_target_paragraphs(target)
        from docx.oxml.ns import qn
        updated = 0
        for paragraph in paragraphs:
            try:
                pPr = paragraph._element.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                for border_name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "12")  # 1.5pt
                    border.set(qn("w:space"), "0")
                    border.set(qn("w:color"), "000000")
                    pBdr.append(border)
                pPr.append(pBdr)
                updated += 1
            except Exception:
                pass
        return {"border_type": "all", "updated_paragraphs": updated}
    
    def _apply_add_hyperlink(self, target: Dict[str, Any], params: Dict[str, Any]) -> Dict[str, Any]:
        """添加超链接"""
        url = str(params.get("url", ""))
        display_text = str(params.get("display_text", url))
        if not url:
            return {"added": False, "reason": "URL不能为空"}
        try:
            from docx.oxml import OxmlElement, parse_xml
            from docx.oxml.ns import qn, nsdecls
            
            # 在文档末尾添加新段落
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(display_text)
            
            # 获取 run 元素
            r = run._element
            
            # 添加颜色和下划线格式
            rPr = r.get_or_add_rPr()
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0000FF")
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(color)
            rPr.append(u)
            
            # 为 run 元素添加超链接属性（如果支持外部链接）
            # 注：python-docx 的超链接支持有限，这里实现简化版本
            # 真正的 OOXML 超链接需要在 relationships 中定义
            
            return {"added": True, "url": url, "display_text": display_text}
        except Exception as e:
            return {"added": False, "reason": str(e)}
    