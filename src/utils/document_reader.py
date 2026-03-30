"""
Document Reader - 本地文档读取器
修复了 venv 包中的 Excel inlineStr 读取问题
支持Excel数据的统计分析（均值、方差、分布等）
"""
import os
import json
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from collections import Counter

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader as PyPdfReader

# 配置日志
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# 如果没有handler，添加一个
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter('[%(asctime)s] %(levelname)s - %(name)s - %(message)s', datefmt='%H:%M:%S'))
    logger.addHandler(handler)


class DocumentReader(ABC):
    """Abstract base class for document readers"""

    @abstractmethod
    def read(self, file_path: str) -> str:
        """Read and extract text from a document"""
        pass


class DocxReader(DocumentReader):
    """DOCX document reader implementation"""

    def read(self, file_path: str) -> str:
        """Read and extract text from DOCX file"""
        logger.info(f"开始读取DOCX文件: {file_path}")
        try:
            doc = DocxDocument(file_path)
            text = []

            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = " ".join([p.text for p in cell.paragraphs]).strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text.append("\t".join(row_text))

            extracted_text = "\n".join(text)
            logger.info(f"DOCX文件读取完成: {file_path}, 字符数: {len(extracted_text)}")
            return extracted_text if extracted_text else "No text found in the DOCX."
        except Exception as e:
            logger.error(f"读取DOCX文件失败: {file_path}, 错误: {str(e)}")
            return f"Error reading DOCX: {str(e)}"


class PdfReader(DocumentReader):
    """PDF document reader implementation"""

    def read(self, file_path: str) -> str:
        """Read and extract text from PDF file"""
        logger.info(f"开始读取PDF文件: {file_path}")
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPdfReader(file)
                page_count = len(pdf_reader.pages)
                logger.info(f"PDF文件共 {page_count} 页")
                text = []

                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text.strip())
                    if (i + 1) % 10 == 0:
                        logger.debug(f"已处理 {i + 1}/{page_count} 页")

                extracted_text = "\n\n".join(text)
                logger.info(f"PDF文件读取完成: {file_path}, 字符数: {len(extracted_text)}")
                return extracted_text if extracted_text else "No text found in the PDF."
        except Exception as e:
            logger.error(f"读取PDF文件失败: {file_path}, 错误: {str(e)}")
            return f"Error reading PDF: {str(e)}"


class MdReader(DocumentReader):
    """Markdown document reader implementation"""

    def read(self, file_path: str) -> str:
        """Read and extract text from Markdown file with encoding handling"""
        logger.info(f"开始读取Markdown文件: {file_path}")
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

        for i, encoding in enumerate(encodings):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                logger.info(f"Markdown文件读取成功: {file_path}, 使用编码: {encoding}, 字符数: {len(text)}")
                return text if text else "No text found in the Markdown file."
            except UnicodeDecodeError:
                logger.debug(f"编码 {encoding} 尝试失败，尝试下一个...")
                continue
            except Exception as e:
                logger.error(f"读取Markdown文件失败: {file_path}, 错误: {str(e)}")
                return f"Error reading Markdown: {str(e)}"

        logger.error(f"无法解码Markdown文件: {file_path}, 所有编码都失败")
        return "Error reading Markdown: Could not decode file with any supported encoding."


class TxtReader(DocumentReader):
    """TXT document reader implementation"""

    def read(self, file_path: str) -> str:
        """Read and extract text from TXT file with encoding handling"""
        logger.info(f"开始读取TXT文件: {file_path}")
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

        for i, encoding in enumerate(encodings):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                logger.info(f"TXT文件读取成功: {file_path}, 使用编码: {encoding}, 字符数: {len(text)}")
                return text if text else "No text found in the TXT file."
            except UnicodeDecodeError:
                logger.debug(f"编码 {encoding} 尝试失败，尝试下一个...")
                continue
            except Exception as e:
                logger.error(f"读取TXT文件失败: {file_path}, 错误: {str(e)}")
                return f"Error reading TXT: {str(e)}"

        logger.error(f"无法解码TXT文件: {file_path}, 所有编码都失败")
        return "Error reading TXT: Could not decode file with any supported encoding."


class ExcelReader(DocumentReader):
    """Excel document reader implementation

    修复说明：
    1. 不使用 read_only=True，因为该模式无法处理 inlineStr 类型的单元格
    2. 添加 max_rows 参数，防止数据量太大超出模型上下文限制
    3. 统计计算默认分析全文件（max_rows_stats），只展示部分样本（max_rows）
    """

    def read(self, file_path: str, max_rows: int = 300, compute_stats: bool = True, max_rows_stats: int = 100000) -> str:
        """
        Read and extract text from Excel file with optional statistical analysis

        :param file_path: Path to the Excel file
        :param max_rows: Maximum number of rows to display/sample (default: 300)
        :param compute_stats: Whether to compute statistics for numeric columns (default: True)
        :param max_rows_stats: Maximum rows to analyze for statistics (default: 100000)
        """
        logger.info(f"开始读取Excel文件: {file_path}")
        logger.info(f"参数: max_rows={max_rows} (展示), max_rows_stats={max_rows_stats} (统计), compute_stats={compute_stats}")

        try:
            wb = load_workbook(file_path)
            sheet_names = wb.sheetnames
            logger.info(f"工作簿包含 {len(sheet_names)} 个工作表: {sheet_names}")

            text = []

            for sheet_name in wb.sheetnames:
                logger.info(f"正在处理工作表: {sheet_name}")
                sheet = wb[sheet_name]
                text.append(f"=== Sheet: {sheet_name} ===")

                # 先读取全部数据用于统计（使用更大的行数限制）
                all_rows_data = []
                headers = None
                total_rows_in_sheet = 0
                
                for row in sheet.iter_rows(values_only=True):
                    total_rows_in_sheet += 1
                    if total_rows_in_sheet <= max_rows_stats:
                        row_text = [str(cell) if cell is not None else "" for cell in row]
                        if any(row_text):
                            if headers is None:
                                headers = row_text
                            else:
                                all_rows_data.append(row_text)

                logger.info(f"工作表 '{sheet_name}': 共有 {total_rows_in_sheet} 行数据，用于统计分析 {len(all_rows_data)} 行")

                # 计算统计信息（基于全文件或大样本）
                if compute_stats and headers and all_rows_data:
                    logger.info(f"工作表 '{sheet_name}': 开始计算统计信息...")
                    stats = self._compute_statistics(headers, all_rows_data)
                    text.append(f"\n【统计摘要】（基于 {len(all_rows_data)} 行数据分析）")
                    text.append(stats)
                    logger.info(f"工作表 '{sheet_name}': 统计计算完成")
                elif compute_stats:
                    logger.warning(f"工作表 '{sheet_name}': 没有有效数据，跳过统计计算")

                # 展示样本数据（截取前 max_rows 行）
                if total_rows_in_sheet > max_rows:
                    logger.info(f"工作表 '{sheet_name}': 显示前 {max_rows} 行作为样本")
                    text.append(f"\n【数据样本】（共 {total_rows_in_sheet} 行，显示前 {max_rows} 行）")
                else:
                    text.append(f"\n【完整数据】（共 {len(all_rows_data)} 行）")
                
                text.append("\t".join(headers) if headers else "")
                for i, row in enumerate(all_rows_data[:max_rows]):
                    text.append("\t".join(row))
                
                if total_rows_in_sheet > max_rows:
                    text.append(f"[... 还有 {total_rows_in_sheet - max_rows} 行数据未显示]")

                text.append("")

            extracted_text = "\n".join(text)
            wb.close()
            logger.info(f"Excel文件读取完成: {file_path}，总行数: {sum(1 for _ in open(file_path, 'rb')) if False else 'N/A'}，输出长度: {len(extracted_text)} 字符")
            return extracted_text if extracted_text else "No text found in the Excel file."
        except Exception as e:
            logger.error(f"读取Excel文件失败: {file_path}, 错误: {str(e)}")
            return f"Error reading Excel: {str(e)}"

    def _compute_statistics(self, headers: List[str], rows_data: List[List[str]]) -> str:
        """
        Compute statistics for numeric columns
        
        :param headers: Column headers
        :param rows_data: List of data rows
        :return: Formatted statistics string
        """
        import statistics
        
        stats_lines = []
        total_rows = len(rows_data)
        stats_lines.append(f"数据总量: {total_rows} 行")
        logger.debug(f"开始计算 {len(headers)} 列的统计信息")

        # 分析每一列
        numeric_cols = 0
        text_cols = 0
        
        for col_idx, header in enumerate(headers):
            if not header or header.strip() == "":
                continue
                
            # 提取该列的所有值
            col_values = []
            for row in rows_data:
                if col_idx < len(row):
                    val = row[col_idx]
                    if val is not None and str(val).strip() != "":
                        col_values.append(val)

            if not col_values:
                continue

            # 判断是否为数值列
            numeric_values = []
            for val in col_values:
                try:
                    numeric_values.append(float(val))
                except (ValueError, TypeError):
                    numeric_values = None
                    break

            if numeric_values is not None and len(numeric_values) > 0:
                # 数值列统计
                numeric_cols += 1
                logger.debug(f"列 '{header}': 检测为数值型，包含 {len(numeric_values)} 个有效值")
                stats_lines.append(f"\n【{header}】(数值型, {len(numeric_values)} 个有效值)")
                
                try:
                    # 基本统计
                    mean_val = statistics.mean(numeric_values)
                    median_val = statistics.median(numeric_values)
                    stats_lines.append(f"  均值: {mean_val:.4f}")
                    stats_lines.append(f"  中位数: {median_val:.4f}")
                    
                    # 标准差和方差
                    if len(numeric_values) > 1:
                        stdev = statistics.stdev(numeric_values)
                        variance = statistics.variance(numeric_values)
                        stats_lines.append(f"  标准差: {stdev:.4f}")
                        stats_lines.append(f"  方差: {variance:.4f}")
                    
                    # 极值
                    min_val = min(numeric_values)
                    max_val = max(numeric_values)
                    stats_lines.append(f"  最小值: {min_val:.4f}")
                    stats_lines.append(f"  最大值: {max_val:.4f}")
                    stats_lines.append(f"  极差: {max_val - min_val:.4f}")
                    
                    # 分位数
                    sorted_values = sorted(numeric_values)
                    q1_idx = len(sorted_values) // 4
                    q3_idx = 3 * len(sorted_values) // 4
                    stats_lines.append(f"  25%分位数(Q1): {sorted_values[q1_idx]:.4f}")
                    stats_lines.append(f"  75%分位数(Q3): {sorted_values[q3_idx]:.4f}")
                    
                    # 分布分析 - 创建直方图(分箱)
                    bins = self._create_histogram(numeric_values, num_bins=5)
                    stats_lines.append("  分布(分5箱):")
                    for bin_label, count in bins.items():
                        bar = "█" * int(count / max(len(numeric_values), 1) * 20)
                        stats_lines.append(f"    {bin_label}: {bar} ({count})")
                    
                except Exception as e:
                    logger.warning(f"列 '{header}': 统计计算错误: {str(e)}")
                    stats_lines.append(f"  [统计计算错误: {str(e)}]")
            else:
                # 文本列统计
                text_cols += 1
                unique_count = len(set(str(v) for v in col_values))
                logger.debug(f"列 '{header}': 检测为文本型，包含 {len(col_values)} 个值，{unique_count} 个唯一值")
                stats_lines.append(f"\n【{header}】(文本型, {len(col_values)} 个值, {unique_count} 个唯一值)")
                
                # 最常见的值
                value_counts = Counter(str(v) for v in col_values)
                top_values = value_counts.most_common(5)
                stats_lines.append("  最常见值:")
                for val, count in top_values:
                    pct = count / len(col_values) * 100
                    stats_lines.append(f"    \"{val[:30]}\": {count} ({pct:.1f}%)")

        logger.info(f"统计计算完成: {numeric_cols} 个数值列, {text_cols} 个文本列")
        return "\n".join(stats_lines)

    def _create_histogram(self, values: List[float], num_bins: int = 5) -> Dict[str, int]:
        """
        Create histogram bins for distribution analysis
        
        :param values: List of numeric values
        :param num_bins: Number of bins
        :return: Dictionary mapping bin labels to counts
        """
        if not values:
            return {}
            
        min_val = min(values)
        max_val = max(values)
        
        if min_val == max_val:
            return {f"{min_val:.2f}": len(values)}
        
        bin_width = (max_val - min_val) / num_bins
        bins = {i: 0 for i in range(num_bins)}
        bin_labels = []
        
        for i in range(num_bins):
            lower = min_val + i * bin_width
            upper = min_val + (i + 1) * bin_width
            if i == num_bins - 1:
                bin_labels.append(f"{lower:.2f}-{upper:.2f}")
            else:
                bin_labels.append(f"{lower:.2f}-{upper:.2f}")
        
        for val in values:
            bin_idx = min(int((val - min_val) / bin_width), num_bins - 1)
            bins[bin_idx] += 1
        
        return {bin_labels[k]: v for k, v in bins.items()}

    def get_statistics_only(self, file_path: str, max_rows: int = 10000) -> Dict[str, Any]:
        """
        Get only statistics without raw data (useful for large files)
        
        :param file_path: Path to the Excel file
        :param max_rows: Maximum rows to analyze
        :return: Dictionary with statistics for each sheet
        """
        try:
            wb = load_workbook(file_path)
            result = {}
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows_data = []
                headers = None
                rows_read = 0
                
                for row in sheet.iter_rows(values_only=True):
                    if rows_read >= max_rows:
                        break
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):
                        if headers is None:
                            headers = row_text
                        else:
                            rows_data.append(row_text)
                        rows_read += 1
                
                if headers and rows_data:
                    result[sheet_name] = self._compute_stats_dict(headers, rows_data)
            
            wb.close()
            return result
        except Exception as e:
            return {"error": str(e)}

    def _compute_stats_dict(self, headers: List[str], rows_data: List[List[str]]) -> Dict[str, Any]:
        """Compute statistics as dictionary (for programmatic access)"""
        import statistics
        
        col_stats = {}
        
        for col_idx, header in enumerate(headers):
            if not header or header.strip() == "":
                continue
                
            col_values = []
            for row in rows_data:
                if col_idx < len(row):
                    val = row[col_idx]
                    if val is not None and str(val).strip() != "":
                        col_values.append(val)
            
            if not col_values:
                continue
            
            # 检查是否为数值列
            numeric_values = None
            for val in col_values:
                try:
                    if numeric_values is None:
                        numeric_values = []
                    numeric_values.append(float(val))
                except (ValueError, TypeError):
                    numeric_values = None
                    break
            
            if numeric_values and len(numeric_values) > 0:
                try:
                    col_stats[header] = {
                        "type": "numeric",
                        "count": len(numeric_values),
                        "mean": statistics.mean(numeric_values),
                        "median": statistics.median(numeric_values),
                        "stdev": statistics.stdev(numeric_values) if len(numeric_values) > 1 else 0,
                        "variance": statistics.variance(numeric_values) if len(numeric_values) > 1 else 0,
                        "min": min(numeric_values),
                        "max": max(numeric_values),
                        "range": max(numeric_values) - min(numeric_values),
                    }
                except Exception:
                    pass
            else:
                value_counts = Counter(str(v) for v in col_values)
                col_stats[header] = {
                    "type": "text",
                    "count": len(col_values),
                    "unique_count": len(value_counts),
                    "top_values": dict(value_counts.most_common(5))
                }
        
        return col_stats


class DocumentReaderFactory:
    """Factory for creating document readers based on file extension"""

    _readers: dict[str, type[DocumentReader]] = {
        ".txt": TxtReader,
        ".md": MdReader,
        ".docx": DocxReader,
        ".pdf": PdfReader,
        ".xlsx": ExcelReader,
        ".xls": ExcelReader,
    }

    @classmethod
    def get_reader(cls, file_path: str) -> DocumentReader:
        """Get appropriate reader for the given file"""
        _, ext = os.path.splitext(file_path.lower())
        if ext not in cls._readers:
            raise ValueError(f"Unsupported document type: {ext}")
        return cls._readers[ext]()

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if the file type is supported"""
        _, ext = os.path.splitext(file_path.lower())
        return ext in cls._readers


def read_document(filename: str, max_rows: int = 100, compute_stats: bool = True, max_rows_stats: int = 100000) -> str:
    """
    Reads and extracts text from a specified document file.
    Supports multiple document types: TXT, DOCX, PDF, Excel (XLSX, XLS).

    :param filename: Path to the document file to read
    :param max_rows: Maximum number of rows to display/sample (default: 100)
    :param compute_stats: Whether to compute statistics for Excel numeric columns (default: True)
    :param max_rows_stats: Maximum rows to analyze for statistics (default: 100000)
    :return: Extracted text from the document
    """
    logger.info(f"[read_document] 请求读取文件: {filename}")
    file_path = Path(filename)

    if not file_path.exists():
        logger.error(f"[read_document] 文件不存在: {filename}")
        return f"Error: File '{filename}' not found."

    if not DocumentReaderFactory.is_supported(str(file_path)):
        logger.error(f"[read_document] 不支持的文件类型: {filename}")
        return f"Error: Unsupported document type for file '{filename}'."

    try:
        reader = DocumentReaderFactory.get_reader(str(file_path))

        # 将参数传递给支持该参数的 reader
        ext = file_path.suffix.lower()
        if ext in ['.xlsx', '.xls'] and isinstance(reader, ExcelReader):
            logger.info(f"[read_document] 调用 ExcelReader 读取: {filename}")
            return reader.read(str(file_path), max_rows=max_rows, compute_stats=compute_stats, max_rows_stats=max_rows_stats)

        result = reader.read(str(file_path))
        logger.info(f"[read_document] 读取完成: {filename}, 长度: {len(result)} 字符")
        return result
    except Exception as e:
        logger.error(f"[read_document] 读取文件失败: {filename}, 错误: {str(e)}")
        return f"Error reading document: {str(e)}"


def get_excel_statistics(filename: str, max_rows: int = 10000) -> Dict[str, Any]:
    """
    Get only statistics for an Excel file without raw data.
    Useful for large files where raw data would exceed context limits.

    :param filename: Path to the Excel file
    :param max_rows: Maximum rows to analyze (default: 10000)
    :return: Dictionary with statistics for each sheet
    """
    logger.info(f"[get_excel_statistics] 请求获取统计: {filename}")
    file_path = Path(filename)
    
    if not file_path.exists():
        logger.error(f"[get_excel_statistics] 文件不存在: {filename}")
        return {"error": f"File '{filename}' not found."}
    
    ext = file_path.suffix.lower()
    if ext not in ['.xlsx', '.xls']:
        logger.error(f"[get_excel_statistics] 非Excel文件: {filename}")
        return {"error": "Only Excel files are supported."}
    
    try:
        reader = ExcelReader()
        result = reader.get_statistics_only(str(file_path), max_rows=max_rows)
        logger.info(f"[get_excel_statistics] 统计计算完成: {filename}")
        return result
    except Exception as e:
        logger.error(f"[get_excel_statistics] 获取统计失败: {filename}, 错误: {str(e)}")
        return {"error": str(e)}
